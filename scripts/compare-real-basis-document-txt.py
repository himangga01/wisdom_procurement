from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from zipfile import ZipFile
from xml.etree import ElementTree as ET


ROOT_DIR = Path(__file__).resolve().parents[1]
BACKEND_DIR = ROOT_DIR / "backend"
DEFAULT_SAMPLE_DIR = BACKEND_DIR / "tests" / "real-basis-document-samples"
DEFAULT_MANIFEST = DEFAULT_SAMPLE_DIR / "manifest.json"
DEFAULT_OUTPUT = DEFAULT_SAMPLE_DIR / "text-comparison-report.json"
TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp949", "euc-kr")
REFERENCE_EXTENSIONS = {".txt", ".docx", ".md", ".markdown"}
REQUIRED_TERMS = [
    "직접생산",
    "확인기준",
    "중소기업자간",
    "경쟁제품",
    "세부품명",
    "생산시설",
    "검사설비",
]


def import_service_extractors():
    sys.path.insert(0, str(BACKEND_DIR))
    from app.pipelines.basis_document import normalize_basis_text  # noqa: PLC0415
    from app.pipelines.parser import extract_document  # noqa: PLC0415

    return extract_document, normalize_basis_text


def read_text_auto(path: Path) -> tuple[str, str]:
    last_error: Exception | None = None
    for encoding in TEXT_ENCODINGS:
        try:
            return path.read_text(encoding=encoding), encoding
        except UnicodeDecodeError as exc:
            last_error = exc
    raise UnicodeDecodeError("unknown", b"", 0, 1, f"Unable to decode {path}: {last_error}")


def extract_docx_reference_text(path: Path) -> tuple[str, dict[str, Any]]:
    from docx import Document  # noqa: PLC0415

    document = Document(str(path))
    python_docx_parts: list[str] = []
    paragraph_count = 0
    table_count = len(document.tables)
    table_cell_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            python_docx_parts.append(text)
            paragraph_count += 1

    for table_index, table in enumerate(document.tables, start=1):
        python_docx_parts.append(f"[TABLE {table_index}]")
        for row_index, row in enumerate(table.rows, start=1):
            cell_texts: list[str] = []
            for cell in row.cells:
                table_cell_count += 1
                cell_lines = [paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()]
                cell_text = " ".join(cell_lines).strip()
                cell_texts.append(cell_text)
            if any(cell_texts):
                python_docx_parts.append(" | ".join(cell_texts))

    xml_parts = extract_docx_xml_text_parts(path)
    python_docx_text = "\n".join(python_docx_parts)
    xml_text = "\n".join(xml_parts)
    selected_text = xml_text.strip() or python_docx_text

    return selected_text, {
        "type": "docx",
        "encoding": "",
        "engine": "docx-package-xml-wt" if xml_text.strip() else "python-docx-paragraphs-and-tables",
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "table_cell_count": table_cell_count,
        "python_docx_char_count": len(normalize_reference_text(python_docx_text)),
        "xml_text_part_count": len(xml_parts),
        "xml_char_count": len(normalize_reference_text(xml_text)),
    }


def extract_docx_xml_text_parts(path: Path) -> list[str]:
    parts: list[str] = []
    with ZipFile(path) as archive:
        xml_names = [
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]
        for name in xml_names:
            try:
                root = ET.fromstring(archive.read(name))
            except ET.ParseError:
                continue
            text_parts: list[str] = []
            for node in root.iter():
                tag = node.tag.rsplit("}", 1)[-1]
                if tag == "t" and node.text:
                    text_parts.append(node.text)
                elif tag == "tab":
                    text_parts.append(" ")
                elif tag in {"br", "cr", "p"}:
                    text_parts.append("\n")
            text = "".join(text_parts).strip()
            if text:
                parts.append(text)
    return parts


def read_reference_file(path: Path) -> tuple[str, dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix not in REFERENCE_EXTENSIONS:
        raise ValueError("Reference file must be .txt, .docx, .md, or .markdown")
    if suffix == ".docx":
        return extract_docx_reference_text(path)

    text, encoding = read_text_auto(path)
    file_type = "md" if suffix in {".md", ".markdown"} else "txt"
    return text, {
        "type": file_type,
        "encoding": encoding,
        "paragraph_count": 0,
        "table_count": 0,
        "table_cell_count": 0,
    }


def normalize_reference_text(text: str) -> str:
    text = (text or "").replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u2024", ".")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def tokens(text: str) -> list[str]:
    return [token.lower() for token in re.findall(r"[0-9A-Za-z가-힣]{2,}", text or "")]


def numeric_tokens(text: str) -> list[str]:
    return re.findall(r"\d[\d.,:/-]*", text or "")


def ngram_counter(text: str, n: int) -> Counter[str]:
    compacted = compact_text(text)
    if not compacted:
        return Counter()
    if len(compacted) < n:
        return Counter([compacted])
    return Counter(compacted[index : index + n] for index in range(len(compacted) - n + 1))


def counter_recall(base: Counter[str], target: Counter[str]) -> float:
    total = sum(base.values())
    if total <= 0:
        return 1.0
    matched = sum(min(count, target.get(item, 0)) for item, count in base.items())
    return matched / total


def unique_recall(base_items: list[str], target_items: list[str]) -> float:
    base_set = set(base_items)
    if not base_set:
        return 1.0
    return len(base_set & set(target_items)) / len(base_set)


def line_items(text: str, *, min_length: int = 12, limit: int = 8000) -> list[str]:
    items: list[str] = []
    for raw_line in (text or "").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if len(line) >= min_length:
            items.append(line)
        if len(items) >= limit:
            break
    return items


def substring_coverage(base_lines: list[str], target_text: str) -> tuple[float, list[str]]:
    if not base_lines:
        return 1.0, []
    target_compact = compact_text(target_text)
    missing: list[str] = []
    hit_count = 0
    for line in base_lines:
        if compact_text(line) in target_compact:
            hit_count += 1
        elif len(missing) < 10:
            missing.append(line)
    return hit_count / len(base_lines), missing


def required_term_coverage(service_text: str, reference_text: str) -> dict[str, dict[str, bool]]:
    return {
        term: {
            "in_service": term in service_text,
            "in_reference": term in reference_text,
        }
        for term in REQUIRED_TERMS
    }


def load_pdf_from_manifest(manifest_path: Path) -> Path:
    if not manifest_path.exists():
        pdfs = sorted(DEFAULT_SAMPLE_DIR.glob("*.pdf"))
        if pdfs:
            return pdfs[0]
        raise FileNotFoundError(f"Manifest not found and no PDF sample exists: {manifest_path}")

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    sample = manifest.get("sample") if isinstance(manifest.get("sample"), dict) else manifest
    saved_path = Path(sample.get("saved_path", ""))
    pdf_path = saved_path if saved_path.is_absolute() else ROOT_DIR / saved_path
    if not pdf_path.exists():
        pdfs = sorted(DEFAULT_SAMPLE_DIR.glob("*.pdf"))
        if pdfs:
            return pdfs[0]
        raise FileNotFoundError(f"PDF sample not found from manifest: {pdf_path}")
    return pdf_path


def compare_texts(service_text: str, reference_text: str, *, ngram_size: int) -> dict[str, Any]:
    service_norm = normalize_reference_text(service_text)
    reference_norm = normalize_reference_text(reference_text)
    service_tokens = tokens(service_norm)
    reference_tokens = tokens(reference_norm)
    service_token_counter = Counter(service_tokens)
    reference_token_counter = Counter(reference_tokens)
    service_numbers = Counter(numeric_tokens(service_norm))
    reference_numbers = Counter(numeric_tokens(reference_norm))
    service_ngrams = ngram_counter(service_norm, ngram_size)
    reference_ngrams = ngram_counter(reference_norm, ngram_size)
    service_lines = line_items(service_norm)
    reference_lines = line_items(reference_norm)
    service_line_coverage, service_missing_lines = substring_coverage(service_lines, reference_norm)
    reference_line_coverage, reference_missing_lines = substring_coverage(reference_lines, service_norm)

    return {
        "service_char_count": len(service_norm),
        "reference_char_count": len(reference_norm),
        "service_compact_char_count": len(compact_text(service_norm)),
        "reference_compact_char_count": len(compact_text(reference_norm)),
        "service_token_count": len(service_tokens),
        "reference_token_count": len(reference_tokens),
        "service_unique_token_count": len(set(service_tokens)),
        "reference_unique_token_count": len(set(reference_tokens)),
        "service_token_multiset_recall_in_reference": round(counter_recall(service_token_counter, reference_token_counter), 4),
        "reference_token_multiset_recall_in_service": round(counter_recall(reference_token_counter, service_token_counter), 4),
        "service_unique_token_recall_in_reference": round(unique_recall(service_tokens, reference_tokens), 4),
        "reference_unique_token_recall_in_service": round(unique_recall(reference_tokens, service_tokens), 4),
        "service_numeric_recall_in_reference": round(counter_recall(service_numbers, reference_numbers), 4),
        "reference_numeric_recall_in_service": round(counter_recall(reference_numbers, service_numbers), 4),
        f"service_char_{ngram_size}gram_recall_in_reference": round(counter_recall(service_ngrams, reference_ngrams), 4),
        f"reference_char_{ngram_size}gram_recall_in_service": round(counter_recall(reference_ngrams, service_ngrams), 4),
        "service_line_coverage_in_reference": round(service_line_coverage, 4),
        "reference_line_coverage_in_service": round(reference_line_coverage, 4),
        "service_missing_line_examples": service_missing_lines,
        "reference_missing_line_examples": reference_missing_lines,
        "required_terms": required_term_coverage(service_norm, reference_norm),
    }


def build_report(args: argparse.Namespace) -> dict[str, Any]:
    extract_document, normalize_basis_text = import_service_extractors()
    manifest_path = Path(args.manifest)
    pdf_path = Path(args.pdf) if args.pdf else load_pdf_from_manifest(manifest_path)
    reference_value = args.reference_file or args.reference_txt
    if not reference_value:
        raise ValueError("Provide --reference-file or --reference-txt.")
    reference_path = Path(reference_value)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    if not reference_path.exists():
        raise FileNotFoundError(f"Reference file not found: {reference_path}")

    parsed = extract_document(pdf_path)
    service_text = normalize_basis_text(parsed.text)
    reference_text, reference_metadata = read_reference_file(reference_path)
    metrics = compare_texts(service_text, reference_text, ngram_size=args.ngram_size)
    qa_errors: list[str] = []
    if metrics["service_token_multiset_recall_in_reference"] < args.min_service_token_recall:
        qa_errors.append("Service token recall in reference is below threshold.")
    if metrics[f"service_char_{args.ngram_size}gram_recall_in_reference"] < args.min_service_ngram_recall:
        qa_errors.append("Service char ngram recall in reference is below threshold.")
    if not all(item["in_service"] and item["in_reference"] for item in metrics["required_terms"].values()):
        qa_errors.append("Required term coverage failed.")

    return {
        "schema_version": "real_basis_text_comparison_report_v1",
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "inputs": {
            "pdf_path": pdf_path.resolve().as_posix(),
            "reference_file_path": reference_path.resolve().as_posix(),
            "reference_file_type": reference_metadata["type"],
            "reference_txt_path": reference_path.resolve().as_posix() if reference_metadata["type"] == "txt" else "",
            "reference_txt_encoding": reference_metadata["encoding"],
            "ngram_size": args.ngram_size,
        },
        "reference_extraction": reference_metadata,
        "service_parser": {
            "kind": parsed.kind,
            "metadata": parsed.metadata,
        },
        "metrics": metrics,
        "interpretation": {
            "service_in_reference": "How much of the service PDF extraction is also present in the external reference file.",
            "reference_in_service": "How much of the external reference file is recovered by the service PDF extraction.",
            "note": "Asymmetric values are expected when the external reference includes additional table layout, page headers, or whitespace artifacts.",
        },
        "qa": {
            "passed": not qa_errors,
            "errors": qa_errors,
            "thresholds": {
                "min_service_token_recall": args.min_service_token_recall,
                "min_service_ngram_recall": args.min_service_ngram_recall,
            },
        },
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Compare service PDF extraction against an external TXT, DOCX, or MD extraction.")
    parser.add_argument("--reference-file", default="", help="External TXT, DOCX, or MD extraction path.")
    parser.add_argument("--reference-txt", default="", help="External TXT extraction path. Kept for backward compatibility.")
    parser.add_argument("--pdf", default="", help="PDF path. Defaults to the real-basis manifest PDF.")
    parser.add_argument("--manifest", default=str(DEFAULT_MANIFEST), help="Real-basis sample manifest path.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="Comparison report output path.")
    parser.add_argument("--ngram-size", type=int, default=5)
    parser.add_argument("--min-service-token-recall", type=float, default=0.75)
    parser.add_argument("--min-service-ngram-recall", type=float, default=0.55)
    parser.add_argument("--strict", action="store_true", help="Exit non-zero when QA thresholds fail.")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    report = build_report(args)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    metrics = report["metrics"]
    ngram_key = f"service_char_{args.ngram_size}gram_recall_in_reference"
    reverse_ngram_key = f"reference_char_{args.ngram_size}gram_recall_in_service"
    print(
        json.dumps(
            {
                "status": "passed" if report["qa"]["passed"] else "failed",
                "output": output_path.as_posix(),
                "service_char_count": metrics["service_char_count"],
                "reference_char_count": metrics["reference_char_count"],
                "service_token_multiset_recall_in_reference": metrics["service_token_multiset_recall_in_reference"],
                "reference_token_multiset_recall_in_service": metrics["reference_token_multiset_recall_in_service"],
                ngram_key: metrics[ngram_key],
                reverse_ngram_key: metrics[reverse_ngram_key],
                "service_line_coverage_in_reference": metrics["service_line_coverage_in_reference"],
                "reference_line_coverage_in_service": metrics["reference_line_coverage_in_service"],
                "errors": report["qa"]["errors"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    if args.strict and not report["qa"]["passed"]:
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
