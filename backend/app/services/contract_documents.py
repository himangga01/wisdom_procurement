import json
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

from app.core.json_utils import parse_json_dict
from app.core.logging import get_logger, log_event, log_exception
from app.core.text import clean_text, parse_int
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

CONTRACT_STORAGE_DIR = "contracts"
CONTRACT_TEMPLATE_VERSION = "contract_docx_template_v1"
CONTRACT_TYPE_STANDARD_SERVICE = "standard_service_contract"
CONTRACT_STATUS_VALUES = {"generated", "failed", "deleted"}
CONTRACT_REVIEW_STATUS_VALUES = {"draft", "needs_review", "approved", "rejected", "archived"}
KST = timezone(timedelta(hours=9))
LOGGER = get_logger("services.contract_documents")

NOTICE_SNAPSHOT_FIELDS = [
    "id",
    "bid_ntce_no",
    "bid_ntce_ord",
    "bid_ntce_nm",
    "ntce_instt_nm",
    "dminstt_nm",
    "bid_ntce_dt",
    "bid_begin_dt",
    "bid_clse_dt",
    "openg_dt",
    "presmpt_prce",
    "bdgt_amt",
    "bssamt",
    "region_text",
    "license_text",
    "source_url",
    "save_status",
    "download_status",
    "analysis_status",
]
CORPORATION_SNAPSHOT_FIELDS = [
    "id",
    "name",
    "management_group_name",
    "business_registration_number",
    "representative_name",
    "corporate_registration_number",
    "business_address",
    "headquarters_address",
    "region",
    "business_category",
    "business_type",
    "business_item",
    "company_size_classification",
    "license_summary",
    "procurement_registration_status",
    "evidence_verification_status",
]
CONTRACT_CUSTOM_FIELD_KEYS = [
    "contract_number",
    "contract_amount",
    "total_service_amount",
    "contract_deposit",
    "delay_penalty_rate",
    "contract_period",
    "delivery_location",
    "other_terms",
    "attachment_notes",
    "buyer_name",
    "buyer_representative_name",
    "buyer_address",
    "corporation_phone",
]
SERVICE_NOTICE_HINTS = ["용역", "위탁", "운영", "유지관리", "관리", "대행", "컨설팅", "설계"]


class ContractInputError(ValueError):
    def __init__(self, detail: str, *, status_code: int = 400, code: str = "contract_input_invalid") -> None:
        super().__init__(detail)
        self.detail = detail
        self.status_code = status_code
        self.code = code


def _now_iso() -> str:
    return datetime.now(KST).isoformat(timespec="seconds")


def _row_allowlist(row: Any, fields: list[str]) -> dict[str, Any]:
    data = dict(row)
    return {field: data.get(field) for field in fields}


def _json_object(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        return parse_json_dict(value)
    return {}


def normalize_contract_custom_fields(value: Any) -> dict[str, str]:
    source = value if isinstance(value, dict) else {}
    return {key: clean_text(source.get(key)) for key in CONTRACT_CUSTOM_FIELD_KEYS}


def _notice_number(notice: dict[str, Any]) -> str:
    number = clean_text(notice.get("bid_ntce_no"))
    order = clean_text(notice.get("bid_ntce_ord"))
    if number and order and order not in {"0", "00", "000"}:
        return f"{number}-{order}"
    return number


def _reference_amount(notice: dict[str, Any]) -> str:
    return clean_text(notice.get("bssamt")) or clean_text(notice.get("presmpt_prce")) or clean_text(notice.get("bdgt_amt"))


def _looks_like_service_notice(notice: dict[str, Any]) -> bool:
    text = " ".join(
        clean_text(notice.get(key))
        for key in ["bid_ntce_nm", "license_text", "analysis_summary_markdown"]
        if clean_text(notice.get(key))
    )
    return any(hint in text for hint in SERVICE_NOTICE_HINTS)


def build_generated_contract_fields(
    notice: dict[str, Any],
    corporation: dict[str, Any],
    custom_fields: dict[str, str],
) -> dict[str, str]:
    buyer_name = custom_fields.get("buyer_name") or clean_text(notice.get("ntce_instt_nm")) or clean_text(notice.get("dminstt_nm"))
    corporation_address = clean_text(corporation.get("business_address")) or clean_text(corporation.get("headquarters_address"))
    return {
        "contract_number": custom_fields.get("contract_number", ""),
        "notice_number": _notice_number(notice),
        "buyer_name": buyer_name,
        "buyer_representative_name": custom_fields.get("buyer_representative_name", ""),
        "buyer_address": custom_fields.get("buyer_address", ""),
        "corporation_name": clean_text(corporation.get("name")),
        "corporation_address": corporation_address,
        "corporation_representative_name": clean_text(corporation.get("representative_name")),
        "corporate_registration_number": clean_text(corporation.get("corporate_registration_number")),
        "corporation_phone": custom_fields.get("corporation_phone", ""),
        "service_name": clean_text(notice.get("bid_ntce_nm")),
        "contract_amount": custom_fields.get("contract_amount") or _reference_amount(notice),
        "total_service_amount": custom_fields.get("total_service_amount", ""),
        "contract_deposit": custom_fields.get("contract_deposit", ""),
        "delay_penalty_rate": custom_fields.get("delay_penalty_rate", ""),
        "contract_period": custom_fields.get("contract_period", ""),
        "delivery_location": custom_fields.get("delivery_location", ""),
        "other_terms": custom_fields.get("other_terms", ""),
        "attachment_notes": custom_fields.get("attachment_notes", ""),
    }


def _judgment_snapshot(row: Any | None) -> dict[str, Any] | None:
    if row is None:
        return None
    data = dict(row)
    result = _json_object(data.get("result_json"))
    return {
        "id": data.get("id"),
        "status": data.get("status"),
        "review_status": data.get("review_status"),
        "summary": _json_object(data.get("summary_json")),
        "preparation_guide": result.get("preparation_guide", {}) if isinstance(result, dict) else {},
    }


def build_contract_input_snapshot(
    conn,
    notice_id: int,
    corporation_id: int,
    judgment_run_id: int | None = None,
    custom_fields: dict[str, Any] | None = None,
) -> dict[str, Any]:
    notice = conn.execute("SELECT * FROM nara_notices WHERE id=?", (parse_int(notice_id),)).fetchone()
    if not notice:
        raise ContractInputError("Saved Nara notice was not found.", status_code=404, code="notice_not_found")
    corporation = conn.execute("SELECT * FROM corporations WHERE id=?", (parse_int(corporation_id),)).fetchone()
    if not corporation:
        raise ContractInputError("Corporation was not found.", status_code=404, code="corporation_not_found")

    judgment = None
    parsed_judgment_run_id = parse_int(judgment_run_id)
    if parsed_judgment_run_id:
        judgment = conn.execute("SELECT * FROM judgment_runs WHERE id=?", (parsed_judgment_run_id,)).fetchone()
        if not judgment:
            raise ContractInputError("Judgment run was not found.", status_code=404, code="judgment_run_not_found")
        if judgment["nara_notice_id"] != notice["id"] or judgment["corporation_id"] != corporation["id"]:
            raise ContractInputError(
                "Judgment run does not belong to the selected notice and corporation.",
                status_code=400,
                code="judgment_run_mismatch",
            )

    notice_payload = _row_allowlist(notice, NOTICE_SNAPSHOT_FIELDS)
    corporation_payload = _row_allowlist(corporation, CORPORATION_SNAPSHOT_FIELDS)
    normalized_custom_fields = normalize_contract_custom_fields(custom_fields)
    generated_fields = build_generated_contract_fields(notice_payload, corporation_payload, normalized_custom_fields)
    warnings: list[str] = []
    if not _looks_like_service_notice({**notice_payload, "analysis_summary_markdown": dict(notice).get("analysis_summary_markdown")}):
        warnings.append("선택한 공고가 용역 공고인지 명확하지 않습니다. 표준계약서 초안 적용 전 원문 검토가 필요합니다.")

    snapshot = {
        "snapshot_version": "contract_input_snapshot_v1",
        "template_version": CONTRACT_TEMPLATE_VERSION,
        "contract_type": CONTRACT_TYPE_STANDARD_SERVICE,
        "created_at": _now_iso(),
        "source_ids": {
            "nara_notice_id": notice["id"],
            "corporation_id": corporation["id"],
            "judgment_run_id": parsed_judgment_run_id or None,
        },
        "notice": notice_payload,
        "corporation": corporation_payload,
        "judgment_run": _judgment_snapshot(judgment),
        "custom_fields": normalized_custom_fields,
        "generated_fields": generated_fields,
        "warnings": warnings,
    }
    snapshot["validation"] = validate_contract_generation_input(snapshot)
    return snapshot


def validate_contract_generation_input(snapshot: dict[str, Any]) -> dict[str, Any]:
    errors: list[str] = []
    warnings = list(snapshot.get("warnings") or [])
    notice = snapshot.get("notice") if isinstance(snapshot.get("notice"), dict) else {}
    corporation = snapshot.get("corporation") if isinstance(snapshot.get("corporation"), dict) else {}
    generated_fields = snapshot.get("generated_fields") if isinstance(snapshot.get("generated_fields"), dict) else {}

    if not clean_text(notice.get("bid_ntce_nm")):
        errors.append("공고명이 없어 계약서 초안을 생성할 수 없습니다.")
    if not clean_text(corporation.get("name")):
        errors.append("법인명이 없어 계약서 초안을 생성할 수 없습니다.")
    for field, label in [
        ("corporation_representative_name", "대표자"),
        ("corporation_address", "주소"),
        ("corporate_registration_number", "법인등록번호"),
        ("corporation_phone", "전화번호"),
    ]:
        if not clean_text(generated_fields.get(field)):
            warnings.append(f"계약상대자 {label} 값이 비어 있습니다.")

    return {
        "valid": not errors,
        "errors": errors,
        "warnings": list(dict.fromkeys(warnings)),
    }


def _set_run_font(run, *, size: int = 10, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _set_cell_text(cell, text: Any, *, bold: bool = False, align: int | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(clean_text(text))
    _set_run_font(run, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_cell(cell, fill: str = "F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_width(cell, width_mm: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_mm * 56.7)))
    tc_w.set(qn("w:type"), "dxa")


def _add_label_row(table, label: str, value: Any, *, label_width_mm: int = 34) -> None:
    row = table.add_row()
    _set_cell_text(row.cells[0], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(row.cells[0])
    _set_cell_width(row.cells[0], label_width_mm)
    _set_cell_text(row.cells[1], value)


def _add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    _set_run_font(run, size=11, bold=True)


def _add_contract_parties_table(document: Document, fields: dict[str, str]) -> None:
    rows = [
        ("발주처", fields.get("buyer_name")),
        ("계약상대자", fields.get("corporation_name")),
        ("상호 또는 법인명", fields.get("corporation_name")),
        ("주소", fields.get("corporation_address")),
        ("대표자", fields.get("corporation_representative_name")),
        ("법인등록번호", fields.get("corporate_registration_number")),
        ("전화번호", fields.get("corporation_phone")),
    ]
    table = document.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    vertical = table.cell(0, 0)
    vertical.merge(table.cell(len(rows) - 1, 0))
    _set_cell_text(vertical, "계약서", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(vertical, "E8EDF5")
    _set_cell_width(vertical, 18)

    for index, (label, value) in enumerate(rows):
        _set_cell_text(table.cell(index, 1), label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(table.cell(index, 1))
        _set_cell_text(table.cell(index, 2), value)


def _add_contract_details_table(document: Document, fields: dict[str, str]) -> None:
    rows = [
        ("용역명", fields.get("service_name")),
        ("계약금액", fields.get("contract_amount")),
        ("총용역부기금액", fields.get("total_service_amount")),
        ("계약보증금", fields.get("contract_deposit")),
        ("지연배상금률", fields.get("delay_penalty_rate")),
        ("계약기간", fields.get("contract_period")),
        ("위치", fields.get("delivery_location")),
        ("그 밖의 사항", fields.get("other_terms")),
    ]
    table = document.add_table(rows=len(rows), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    vertical = table.cell(0, 0)
    vertical.merge(table.cell(len(rows) - 1, 0))
    _set_cell_text(vertical, "계약내용", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(vertical, "E8EDF5")
    _set_cell_width(vertical, 18)

    for index, (label, value) in enumerate(rows):
        _set_cell_text(table.cell(index, 1), label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(table.cell(index, 1))
        _set_cell_text(table.cell(index, 2), value)


def _add_paragraph(document: Document, text: str, *, size: int = 10, bold: bool = False, align: int | None = None) -> None:
    paragraph = document.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold)


def _render_contract_docx_to_path(snapshot: dict[str, Any], output_path: Path) -> None:
    validation = snapshot.get("validation") if isinstance(snapshot.get("validation"), dict) else {}
    if validation and not validation.get("valid", False):
        raise ContractInputError("Contract input snapshot is not valid.", code="contract_snapshot_invalid")

    fields = snapshot.get("generated_fields") if isinstance(snapshot.get("generated_fields"), dict) else {}
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)

    styles = document.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("[별지 제9호서식] <개정 2024. 12. 6.>                                      (앞쪽)")
    _set_run_font(run, size=9)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("용역표준계약서")
    _set_run_font(run, size=18, bold=True)

    meta_table = document.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    meta_table.style = "Table Grid"
    _set_cell_text(meta_table.cell(0, 0), "계약번호", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(meta_table.cell(0, 0))
    _set_cell_text(meta_table.cell(0, 1), fields.get("contract_number"))
    _set_cell_text(meta_table.cell(1, 0), "공고번호", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(meta_table.cell(1, 0))
    _set_cell_text(meta_table.cell(1, 1), fields.get("notice_number"))

    _add_contract_parties_table(document, fields)
    _add_contract_details_table(document, fields)

    _add_paragraph(
        document,
        "위 계약에 관하여 발주처와 계약상대자는 붙임 서류 및 관계 법령에 따라 계약을 체결하며, "
        "본 문서는 시스템에서 자동 생성한 계약서 초안이므로 관리자 검토 후 사용해야 합니다.",
    )
    _add_section_heading(document, "붙임서류")
    attachment_notes = fields.get("attachment_notes") or "계약 관련 산출물, 공고문, 과업지시서, 필요 증빙서류"
    _add_paragraph(document, f"1. {attachment_notes}")

    warnings = snapshot.get("warnings") if isinstance(snapshot.get("warnings"), list) else []
    validation_warnings = validation.get("warnings") if isinstance(validation.get("warnings"), list) else []
    review_notes = list(dict.fromkeys([clean_text(item) for item in [*warnings, *validation_warnings] if clean_text(item)]))
    if review_notes:
        _add_section_heading(document, "자동 생성 검토 메모")
        for index, note in enumerate(review_notes, start=1):
            _add_paragraph(document, f"{index}. {note}")

    signature_table = document.add_table(rows=2, cols=2)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    signature_table.style = "Table Grid"
    _set_cell_text(signature_table.cell(0, 0), "자치단체의 장 또는 계약담당자", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(signature_table.cell(0, 1), "계약상대자", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(signature_table.cell(1, 0), "서명 또는 날인")
    _set_cell_text(signature_table.cell(1, 1), "서명 또는 날인")

    _add_paragraph(document, "210㎜×297㎜", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
    document.save(output_path)


def render_standard_service_contract_docx(snapshot: dict[str, Any], output_path: Path) -> None:
    output_path = Path(output_path)
    tmp_path = output_path.with_suffix(f"{output_path.suffix}.tmp")
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        if tmp_path.exists():
            tmp_path.unlink()
        _render_contract_docx_to_path(snapshot, tmp_path)
        tmp_path.replace(output_path)
    except Exception:
        try:
            tmp_path.unlink(missing_ok=True)
        except OSError:
            pass
        raise


def json_dumps(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False)


def _json_payload(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        return parse_json_dict(value)
    return {}


def _relative_contract_path(storage_root: Path, path: Path) -> str:
    root = storage_root.resolve()
    resolved = path.resolve()
    if not resolved.is_relative_to(root):
        raise ValueError("Contract path is outside storage root.")
    return resolved.relative_to(root).as_posix()


def resolve_contract_stored_path(storage_root: Path, stored_file_path: str) -> Path | None:
    cleaned = clean_text(stored_file_path)
    if not cleaned:
        return None
    root = (storage_root / CONTRACT_STORAGE_DIR).resolve()
    candidate = (storage_root / cleaned).resolve()
    if not candidate.is_relative_to(root):
        raise ValueError("Stored contract path is outside storage/contracts.")
    return candidate


def contract_document_payload(conn, row: Any | None) -> dict[str, Any] | None:
    if row is None:
        return None
    payload = dict(row)
    payload["input_snapshot"] = _json_payload(payload.pop("input_snapshot_json", "{}"))
    payload["generated_fields"] = _json_payload(payload.pop("generated_fields_json", "{}"))
    payload["validation"] = _json_payload(payload.pop("validation_json", "{}"))
    payload["download_url"] = f"/api/contracts/{payload['id']}/download" if payload.get("file_name") else ""
    notice = conn.execute("SELECT id, bid_ntce_no, bid_ntce_ord, bid_ntce_nm FROM nara_notices WHERE id=?", (payload["nara_notice_id"],)).fetchone()
    corporation = conn.execute("SELECT id, name FROM corporations WHERE id=?", (payload["corporation_id"],)).fetchone()
    payload["notice"] = dict(notice) if notice else None
    payload["corporation"] = dict(corporation) if corporation else None
    return payload


def list_contract_documents_payload(conn, filters: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    filters = filters or {}
    where: list[str] = []
    params: list[Any] = []
    for key, column in [
        ("notice_id", "cd.nara_notice_id"),
        ("corporation_id", "cd.corporation_id"),
    ]:
        value = parse_int(filters.get(key))
        if value:
            where.append(f"{column}=?")
            params.append(value)
    status = clean_text(filters.get("status"))
    if status:
        where.append("cd.status=?")
        params.append(status)
    review_status = clean_text(filters.get("review_status"))
    if review_status:
        where.append("cd.review_status=?")
        params.append(review_status)
    keyword = clean_text(filters.get("keyword"))
    if keyword:
        like = f"%{keyword}%"
        where.append("(cd.title LIKE ? OR cd.file_name LIKE ? OR n.bid_ntce_nm LIKE ? OR c.name LIKE ?)")
        params.extend([like, like, like, like])
    query = """
        SELECT cd.*
        FROM contract_documents cd
        LEFT JOIN nara_notices n ON n.id=cd.nara_notice_id
        LEFT JOIN corporations c ON c.id=cd.corporation_id
    """
    if where:
        query += " WHERE " + " AND ".join(where)
    query += " ORDER BY cd.id DESC"
    rows = conn.execute(query, tuple(params)).fetchall()
    return [payload for row in rows if (payload := contract_document_payload(conn, row))]


def get_contract_document_payload(conn, contract_document_id: int) -> dict[str, Any] | None:
    row = conn.execute("SELECT * FROM contract_documents WHERE id=?", (parse_int(contract_document_id),)).fetchone()
    return contract_document_payload(conn, row)


def create_contract_document(
    conn,
    payload: dict[str, Any],
    storage_root: Path,
    renderer=render_standard_service_contract_docx,
) -> dict[str, Any]:
    notice_id = parse_int(payload.get("nara_notice_id"))
    corporation_id = parse_int(payload.get("corporation_id"))
    judgment_run_id = parse_int(payload.get("judgment_run_id")) or None
    log_event(
        LOGGER,
        "contract.create.started",
        notice_id=notice_id,
        corporation_id=corporation_id,
        judgment_run_id=judgment_run_id,
    )
    snapshot = build_contract_input_snapshot(
        conn,
        notice_id,
        corporation_id,
        judgment_run_id=judgment_run_id,
        custom_fields=payload.get("custom_fields") if isinstance(payload.get("custom_fields"), dict) else {},
    )
    generated_fields = snapshot["generated_fields"]
    title = clean_text(payload.get("title")) or f"{generated_fields.get('service_name') or '공고'} 계약서 초안"
    validation = snapshot.get("validation") if isinstance(snapshot.get("validation"), dict) else {"valid": False, "errors": ["Invalid snapshot."], "warnings": []}
    now = _now_iso()
    initial_error_message = "; ".join(validation.get("errors", [])) if not validation.get("valid") else ""
    cur = conn.execute(
        """
        INSERT INTO contract_documents (
          nara_notice_id, corporation_id, judgment_run_id, status, review_status,
          contract_type, template_version, title, input_snapshot_json,
          generated_fields_json, validation_json, error_message, created_at, updated_at
        ) VALUES (?, ?, ?, ?, 'draft', ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            notice_id,
            corporation_id,
            judgment_run_id,
            "generated" if validation.get("valid") else "failed",
            CONTRACT_TYPE_STANDARD_SERVICE,
            CONTRACT_TEMPLATE_VERSION,
            title,
            json_dumps(snapshot),
            json_dumps(generated_fields),
            json_dumps(validation),
            initial_error_message,
            now,
            now,
        ),
    )
    contract_id = cur.lastrowid
    if not validation.get("valid"):
        log_event(
            LOGGER,
            "contract.validation.failed",
            level="warning",
            contract_id=contract_id,
            notice_id=notice_id,
            corporation_id=corporation_id,
            error_count=len(validation.get("errors", [])),
            warning_count=len(validation.get("warnings", [])),
        )
        row = conn.execute("SELECT * FROM contract_documents WHERE id=?", (contract_id,)).fetchone()
        return contract_document_payload(conn, row)

    file_name = safe_contract_file_name(contract_id, title)
    output_path = contract_output_path(storage_root, contract_id, file_name)
    try:
        renderer(snapshot, output_path)
        completed_at = _now_iso()
        conn.execute(
            """
            UPDATE contract_documents
            SET file_name=?, stored_file_path=?, file_size_bytes=?,
                status='generated', error_message='', updated_at=?
            WHERE id=?
            """,
            (
                file_name,
                _relative_contract_path(storage_root, output_path),
                output_path.stat().st_size,
                completed_at,
                contract_id,
            ),
        )
        log_event(
            LOGGER,
            "contract.create.completed",
            contract_id=contract_id,
            notice_id=notice_id,
            corporation_id=corporation_id,
            file_size_bytes=output_path.stat().st_size,
        )
    except Exception as exc:
        log_exception(
            LOGGER,
            "contract.render.failed",
            exc,
            contract_id=contract_id,
            notice_id=notice_id,
            corporation_id=corporation_id,
            template_version=CONTRACT_TEMPLATE_VERSION,
        )
        failed_at = _now_iso()
        failed_validation = {
            "valid": False,
            "errors": [str(exc)],
            "warnings": validation.get("warnings", []),
        }
        conn.execute(
            """
            UPDATE contract_documents
            SET status='failed', validation_json=?, error_message=?, updated_at=?
            WHERE id=?
            """,
            (json_dumps(failed_validation), str(exc), failed_at, contract_id),
        )
    row = conn.execute("SELECT * FROM contract_documents WHERE id=?", (contract_id,)).fetchone()
    return contract_document_payload(conn, row)


def delete_contract_document(conn, contract_document_id: int, storage_root: Path) -> dict[str, Any] | None:
    row = conn.execute("SELECT * FROM contract_documents WHERE id=?", (parse_int(contract_document_id),)).fetchone()
    if not row:
        return None
    payload = contract_document_payload(conn, row)
    path = resolve_contract_stored_path(storage_root, row["stored_file_path"])
    if path and path.exists():
        path.unlink()
        try:
            path.parent.rmdir()
        except OSError:
            pass
    conn.execute("DELETE FROM contract_documents WHERE id=?", (row["id"],))
    return payload


def safe_contract_file_name(contract_id: int, title: str = "") -> str:
    stem = "".join(ch if ch.isalnum() else "-" for ch in title.strip()).strip("-")
    stem = "-".join(part for part in stem.split("-") if part)
    if stem:
        stem = stem[:80]
        return f"contract-{contract_id}-{stem}.docx"
    return f"contract-{contract_id}.docx"


def contract_storage_dir(storage_root: Path, contract_document_id: int) -> Path:
    root = (storage_root / CONTRACT_STORAGE_DIR).resolve()
    target = (root / str(contract_document_id)).resolve()
    if not target.is_relative_to(root):
        raise ValueError("Contract storage path is outside storage/contracts.")
    return target


def contract_output_path(storage_root: Path, contract_document_id: int, file_name: str) -> Path:
    safe_name = Path(file_name).name
    if not safe_name.lower().endswith(".docx"):
        raise ValueError("Contract file name must end with .docx.")
    target_dir = contract_storage_dir(storage_root, contract_document_id)
    target = (target_dir / safe_name).resolve()
    if not target.is_relative_to(target_dir):
        raise ValueError("Contract output path is outside its contract directory.")
    return target
