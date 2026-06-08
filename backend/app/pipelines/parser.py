from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

from app.core.logging import get_logger, log_event, log_exception
from app.pipelines.pdf_readers import read_pdf_document

LOGGER = get_logger("pipelines.parser")


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    kind: str
    metadata: dict[str, Any]


def extract_text_from_file(file_path: str | Path) -> tuple[str, str]:
    parsed = extract_document(file_path)
    return parsed.text, parsed.kind


def extract_document(file_path: str | Path) -> ParsedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()
    log_event(
        LOGGER,
        "document.extract.started",
        domain="document",
        file_name=path.name,
        file_extension=suffix,
        file_size_bytes=path.stat().st_size if path.exists() else 0,
    )

    try:
        if suffix == ".pdf":
            result = read_pdf_document(path)
            log_event(
                LOGGER,
                "document.extract.completed",
                domain="document",
                file_name=path.name,
                file_extension=suffix,
                kind="pdf",
                char_count=len(result.text),
                metadata={
                    "engine": result.metadata.get("engine"),
                    "page_count": result.metadata.get("page_count"),
                    "table_count": result.metadata.get("table_count"),
                    "needs_ocr": result.metadata.get("needs_ocr"),
                },
            )
            return ParsedDocument(text=result.text, kind="pdf", metadata=result.metadata)
        if suffix == ".docx":
            text = _extract_docx_text(path)
            log_event(
                LOGGER,
                "document.extract.completed",
                domain="document",
                file_name=path.name,
                file_extension=suffix,
                kind="docx",
                char_count=len(text),
            )
            return ParsedDocument(
                text=text,
                kind="docx",
                metadata={"engine": "python-docx", "char_count": len(text), "needs_ocr": False},
            )
    except Exception as exc:
        log_exception(
            LOGGER,
            "document.extract.failed",
            exc,
            domain="document",
            file_name=path.name,
            file_extension=suffix,
            file_size_bytes=path.stat().st_size if path.exists() else 0,
        )
        raise

    error = ValueError("Unsupported file type. Only PDF and DOCX are allowed.")
    log_exception(
        LOGGER,
        "document.extract.failed",
        error,
        domain="document",
        file_name=path.name,
        file_extension=suffix,
        file_size_bytes=path.stat().st_size if path.exists() else 0,
    )
    raise error


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).replace("\u2024", ".")
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
