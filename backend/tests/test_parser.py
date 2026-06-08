import os
import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document

from app.pipelines.parser import extract_document, extract_text_from_file


class PdfReaderEnv:
    def __init__(self, engine: str) -> None:
        self.engine = engine
        self.previous: str | None = None

    def __enter__(self) -> None:
        self.previous = os.environ.get("PDF_READER_ENGINE")
        os.environ["PDF_READER_ENGINE"] = self.engine

    def __exit__(self, *_args: object) -> None:
        if self.previous is None:
            os.environ.pop("PDF_READER_ENGINE", None)
        else:
            os.environ["PDF_READER_ENGINE"] = self.previous


class ParserTests(unittest.TestCase):
    def test_pdf_uses_pymupdf_and_returns_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / "notice.pdf"
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)
            page.insert_text((72, 72), "1. Procurement notice")
            page.insert_text((72, 100), "A. Project name : Smoke Test Forest Project")
            page.insert_text((72, 128), "BidStartBidEndOpenDate2026. 4. 7. 09:00")
            doc.save(pdf_path)
            doc.close()

            with PdfReaderEnv("pymupdf"):
                parsed = extract_document(pdf_path)

            self.assertEqual(parsed.kind, "pdf")
            self.assertEqual(parsed.metadata["engine"], "PyMuPDF")
            self.assertEqual(parsed.metadata["page_count"], 1)
            self.assertFalse(parsed.metadata["needs_ocr"])
            self.assertIn("Procurement notice", parsed.text)
            self.assertIn("Smoke Test Forest Project", parsed.text)

    def test_blank_pdf_is_marked_as_ocr_candidate(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / "blank.pdf"
            doc = fitz.open()
            doc.new_page(width=300, height=200)
            doc.save(pdf_path)
            doc.close()

            with PdfReaderEnv("pymupdf"):
                parsed = extract_document(pdf_path)

            self.assertEqual(parsed.kind, "pdf")
            self.assertTrue(parsed.metadata["needs_ocr"])
            self.assertEqual(parsed.text, "")

    def test_docx_extraction_still_works(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "notice.docx"
            doc = Document()
            doc.add_paragraph("공 사 명 : 조달 테스트 사업")
            doc.add_paragraph("기초금액 : 금61,864,000원")
            doc.save(docx_path)

            text, kind = extract_text_from_file(docx_path)

            self.assertEqual(kind, "docx")
            self.assertIn("조달 테스트 사업", text)
            self.assertIn("61,864,000", text)

    def test_docx_extraction_includes_table_cells(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "notice-table.docx"
            doc = Document()
            doc.add_paragraph("Procurement notice")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Requirement"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "License"
            table.cell(1, 1).text = "Forest business license"
            doc.save(docx_path)

            text, kind = extract_text_from_file(docx_path)

        self.assertEqual(kind, "docx")
        self.assertIn("Procurement notice", text)
        self.assertIn("Requirement | Value", text)
        self.assertIn("License | Forest business license", text)


if __name__ == "__main__":
    unittest.main()
