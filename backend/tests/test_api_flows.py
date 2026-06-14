import io
import json
import os
import shutil
import sqlite3
import tempfile
import time
import unittest
import zipfile
from pathlib import Path

import fitz
from docx import Document

_TEST_TMP_ROOT = Path(os.environ.get("WISDOM_TEST_TMPDIR", Path(__file__).resolve().parents[2] / "temp" / "api-tests"))
_TEST_TMP_ROOT.mkdir(parents=True, exist_ok=True)
_TMP_DIR = Path(tempfile.mkdtemp(prefix="wisdom_api_tests_", dir=_TEST_TMP_ROOT))

os.environ["SQLITE_PATH"] = str(_TMP_DIR / "test.db")
os.environ["STORAGE_ROOT"] = str(_TMP_DIR / "storage")
os.environ["OCR_ENGINE"] = "noop"
os.environ["OPENAI_API_KEY"] = ""
os.environ["GEMINI_API_KEY"] = ""
os.environ["AI_PROVIDER_DEFAULT"] = "gemini"
os.environ["AI_MODEL_DEFAULT"] = "gemini-2.5-flash"
os.environ.pop("NARA_API_SERVICE_KEY", None)

from app import main as runtime  # noqa: E402

runtime.NARA_API_SERVICE_KEY = ""
runtime.OPENAI_API_KEY = ""
runtime.GEMINI_API_KEY = ""


def make_pdf_bytes(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=300, height=200)
    page.insert_text((40, 80), text)
    payload = doc.tobytes()
    doc.close()
    return payload


def make_blank_pdf_bytes() -> bytes:
    doc = fitz.open()
    doc.new_page(width=300, height=200)
    payload = doc.tobytes()
    doc.close()
    return payload


def make_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.strip().splitlines():
        doc.add_paragraph(line.strip())
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class ApiFlowTests(unittest.TestCase):
    def test_decode_http_body_falls_back_to_korean_legacy_encodings(self) -> None:
        encoded = "공고명: 미세먼지 저감숲".encode("cp949")

        decoded = runtime.decode_http_body(encoded, "utf-8")

        self.assertEqual(decoded, "공고명: 미세먼지 저감숲")

    def test_json_responses_declare_utf8_for_korean_text(self) -> None:
        response = self.client.get("/api/operations/summary")

        self.assertEqual(response.status_code, 200)
        self.assertIn("charset=utf-8", response.content_type.lower())

    def test_api_responses_include_request_id_for_debug_logs(self) -> None:
        response = self.client.get("/api/operations/summary")

        self.assertEqual(response.status_code, 200)
        self.assertRegex(response.headers.get("X-Request-ID", ""), r"^[0-9a-f]{32}$")

    def test_fallback_user_summary_removes_legacy_reinforcement_copy(self) -> None:
        summary = runtime.deterministic_user_summary(
            mode="judgment",
            summary={"missing_count": 1, "matched_count": 0, "needs_review_count": 0, "uncertain_count": 0},
            items=[
                {
                    "requirement_input_id": "money:1",
                    "match_status": "missing",
                    "label": "금액 조건",
                    "required_value": "추정가격: 36,190,000원",
                    "gap_reason": "보강 필요 항목입니다.",
                    "recommended_action": "법인 프로필을 보강하세요.",
                }
            ],
            evidence_links=[],
        )
        payload = json.dumps(summary, ensure_ascii=False)

        self.assertNotIn("보강 필요", payload)
        self.assertNotIn("보강하세요", payload)
        self.assertIn("준비 필요", payload)
        self.assertIn("자료를 보완하세요", payload)

    def test_user_summary_sanitizer_falls_back_when_action_shape_is_invalid(self) -> None:
        fallback = {
            "generated_by": "fallback",
            "headline_status": "준비 필요",
            "plain_summary": "기본 요약",
            "top_priority_actions": [
                {
                    "title": "기본 준비 항목",
                    "reason": "기본 사유",
                    "next_step": "기본 다음 행동",
                    "related_requirement_ids": ["notice_requirement:1"],
                    "documents": ["증빙서류"],
                }
            ],
            "missing_groups": [],
            "item_explanations": {},
            "risk_notes": [],
            "evidence_links": [],
        }

        sanitized = runtime.sanitize_user_summary_payload(
            {
                "headline_status": "보강 필요",
                "plain_summary": "요약",
                "top_priority_actions": [None, {"title": 123, "reason": None, "next_step": None}],
                "missing_groups": "bad-shape",
                "item_explanations": {"notice_requirement:1": None},
                "risk_notes": [None, "citation 확인"],
            },
            fallback,
            [],
        )

        self.assertEqual(sanitized["top_priority_actions"], fallback["top_priority_actions"])
        self.assertEqual(sanitized["headline_status"], "준비 필요")
        self.assertNotIn("citation", json.dumps(sanitized, ensure_ascii=False))

    def test_fallback_user_summary_keeps_all_related_requirement_ids(self) -> None:
        items = [
            {
                "requirement_input_id": f"notice_requirement:{index}",
                "match_status": "missing",
                "requirement_type": "license",
                "label": f"면허 조건 {index}",
                "required_value": f"면허 {index}",
                "source_text": f"면허 {index} 필요",
            }
            for index in range(1, 8)
        ]

        summary = runtime.deterministic_user_summary(
            mode="judgment",
            summary={"missing_count": 7, "matched_count": 0, "needs_review_count": 0, "uncertain_count": 0},
            items=items,
            evidence_links=[],
        )

        self.assertEqual(summary["top_priority_actions"][0]["related_requirement_ids"], [f"notice_requirement:{index}" for index in range(1, 8)])

    def test_phase4a_operations_summary_handles_empty_database(self) -> None:
        response = self.client.get("/api/operations/summary")

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertIn(payload["overall_status"], {"ok", "warning", "action_required"})
        self.assertEqual(payload["counts"]["failed_jobs_24h"], 0)
        self.assertEqual(payload["counts"]["pending_reviews"], 0)
        self.assertEqual(payload["health"]["database"]["status"], "ok")
        self.assertEqual(payload["health"]["ocr"]["status"], "unavailable")
        self.assertEqual(payload["health"]["nara_api"]["configured"], False)
        self.assertEqual(payload["last_backup"]["status"], "not_available")

    def test_phase4e_external_access_status_is_secret_safe(self) -> None:
        status_path = runtime.BASE_DIR / "temp" / "ngrok.status.json"
        status_path.parent.mkdir(parents=True, exist_ok=True)
        status_path.unlink(missing_ok=True)

        missing_response = self.client.get("/api/external-access/status")
        self.assertEqual(missing_response.status_code, 200)
        self.assertFalse(missing_response.get_json()["enabled"])

        try:
            status_path.write_text(
                json.dumps(
                    {
                        "enabled": True,
                        "provider": "ngrok",
                        "frontend_public_url": "https://front.ngrok-free.app",
                        "backend_public_url": "https://back.ngrok-free.app",
                        "frontend_local_url": "http://127.0.0.1:5199",
                        "backend_local_url": "http://127.0.0.1:18111",
                        "updated_at": "2026-06-07 10:00:00 +0900",
                        "NGROK_AUTHTOKEN": "SHOULD_NOT_LEAK",
                        "API_KEY": "SHOULD_NOT_LEAK",
                        "raw_env": "SHOULD_NOT_LEAK",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            response = self.client.get("/api/external-access/status")
        finally:
            status_path.unlink(missing_ok=True)

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertTrue(payload["enabled"])
        self.assertEqual(payload["frontend_public_url"], "https://front.ngrok-free.app")
        serialized = json.dumps(payload, ensure_ascii=False)
        self.assertNotIn("SHOULD_NOT_LEAK", serialized)
        self.assertNotIn("NGROK_AUTHTOKEN", serialized)
        self.assertNotIn("API_KEY", serialized)

    def test_phase4a_operations_summary_reports_failures_and_review_queues(self) -> None:
        now = runtime.now_iso()
        with runtime.db_conn() as conn:
            conn.execute(
                """
                INSERT INTO nara_collection_runs (
                  status, mode, keyword, start_date, end_date, searched_count,
                  saved_count, skipped_count, error_message, criteria_json,
                  result_json, created_at, updated_at
                ) VALUES ('failed', 'api', '조경', '2026-05-23', '2026-05-24',
                  0, 0, 0, 'network timeout', '{}', '{}', ?, ?)
                """,
                (now, now),
            )
            conn.execute(
                """
                INSERT INTO basis_documents (
                  title, original_file_name, stored_file_path, processing_status,
                  parse_status, ocr_status, chunk_status, index_status,
                  error_message, created_at, updated_at
                ) VALUES ('기준문서 실패 샘플', 'basis.pdf', 'missing.pdf',
                  'failed', 'failed', 'pending', 'failed', 'failed',
                  'Stored file not found', ?, ?)
                """,
                (now, now),
            )
            conn.execute(
                """
                INSERT INTO corporation_evidence_documents (
                  original_file_name, stored_file_path, extraction_status,
                  ocr_status, review_status, created_at, updated_at
                ) VALUES ('증빙.pdf', 'evidence.pdf', 'completed', 'skipped',
                  'pending', ?, ?)
                """,
                (now, now),
            )
            conn.execute(
                """
                INSERT INTO backup_runs (
                  backup_type, status, file_name, file_path, file_size_bytes,
                  manifest_json, validation_json, error_message,
                  created_at, completed_at, updated_at
                ) VALUES ('full_local', 'failed', 'failed-backup.zip', '',
                  0, '{}', '{}', 'checksum mismatch', ?, ?, ?)
                """,
                (now, now, now),
            )
            conn.commit()

        response = self.client.get("/api/operations/summary")

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["overall_status"], "action_required")
        self.assertGreaterEqual(payload["counts"]["failed_jobs_24h"], 2)
        self.assertGreaterEqual(payload["counts"]["pending_reviews"], 1)
        failure_types = {item["operation_type"] for item in payload["recent_failures"]}
        self.assertIn("nara_collection", failure_types)
        self.assertIn("basis_document_processing", failure_types)
        self.assertIn("backup_create", failure_types)
        self.assertEqual(payload["last_backup"]["status"], "failed")
        self.assertEqual(payload["last_backup"]["file_name"], "failed-backup.zip")
        queue_types = {item["queue_type"] for item in payload["review_queues"]}
        self.assertIn("corporation_evidence", queue_types)

    def test_phase4a_operations_summary_never_exposes_raw_api_keys(self) -> None:
        original_nara_key = runtime.NARA_API_SERVICE_KEY
        original_gemini_key = runtime.GEMINI_API_KEY
        runtime.NARA_API_SERVICE_KEY = "PHASE4_RAW_NARA_SECRET_123456"
        runtime.GEMINI_API_KEY = "PHASE4_RAW_GEMINI_SECRET_123456"
        try:
            response = self.client.get("/api/operations/summary")
        finally:
            runtime.NARA_API_SERVICE_KEY = original_nara_key
            runtime.GEMINI_API_KEY = original_gemini_key

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        body = json.dumps(payload, ensure_ascii=False)
        self.assertTrue(payload["health"]["nara_api"]["configured"])
        self.assertTrue(payload["health"]["ai_provider"]["configured"])
        self.assertNotIn("PHASE4_RAW_NARA_SECRET_123456", body)
        self.assertNotIn("PHASE4_RAW_GEMINI_SECRET_123456", body)

    def test_phase4b_operation_runs_record_nara_collection(self) -> None:
        response = self.client.post(
            "/api/nara/collection-runs",
            json={
                "keyword": "조경",
                "start_date": "2026-05-23",
                "end_date": "2026-05-24",
                "dry_run": True,
                "notices": [
                    {
                        "bidNtceNo": "PHASE4B-001",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "Phase 4B 운영 이력 테스트",
                        "ntceInsttNm": "테스트기관",
                        "ntceSpecFileNm1": "공고문.pdf",
                        "ntceSpecDocUrl1": "https://example.com/notice.pdf",
                    }
                ],
            },
        )
        self.assertEqual(response.status_code, 201)
        collection_run = response.get_json()

        list_response = self.client.get("/api/operation-runs?operation_type=nara_collection")

        self.assertEqual(list_response.status_code, 200)
        runs = list_response.get_json()
        self.assertEqual(len(runs), 1)
        self.assertEqual(runs[0]["operation_type"], "nara_collection")
        self.assertEqual(runs[0]["target_id"], collection_run["id"])
        self.assertEqual(runs[0]["status"], "completed")
        self.assertEqual(runs[0]["request"]["keyword"], "조경")

    def test_phase4c_operation_run_retry_creates_new_linked_run(self) -> None:
        response = self.client.post(
            "/api/nara/collection-runs",
            json={
                "keyword": "재시도",
                "start_date": "2026-05-23",
                "end_date": "2026-05-24",
                "dry_run": False,
            },
        )
        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.get_json()["status"], "not_configured")
        runs = self.client.get("/api/operation-runs?operation_type=nara_collection").get_json()
        original_run = runs[0]

        retry_response = self.client.post(f"/api/operation-runs/{original_run['id']}/retry")

        self.assertEqual(retry_response.status_code, 201)
        retry_run = retry_response.get_json()
        self.assertEqual(retry_run["retry_of_run_id"], original_run["id"])
        self.assertEqual(retry_run["retry_count"], 1)
        self.assertEqual(retry_run["operation_type"], "nara_collection")
        self.assertEqual(retry_run["status"], "not_configured")
        self.assertEqual(retry_run["error_code"], "missing_api_key")

        second_retry_response = self.client.post(f"/api/operation-runs/{original_run['id']}/retry")
        self.assertEqual(second_retry_response.status_code, 201)
        self.assertEqual(second_retry_response.get_json()["retry_count"], 2)

    def test_phase4d_backup_create_validate_and_exclude_env(self) -> None:
        storage_root = Path(os.environ["STORAGE_ROOT"])
        upload_dir = storage_root / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        (upload_dir / "sample.txt").write_text("backup sample", encoding="utf-8")
        (upload_dir / ".env").write_text("UPLOAD_SECRET=1", encoding="utf-8")
        nested_dir = upload_dir / "nested"
        nested_dir.mkdir()
        (nested_dir / ".env.local").write_text("NESTED_SECRET=1", encoding="utf-8")
        (storage_root / ".env").write_text("SHOULD_NOT_BE_INCLUDED=1", encoding="utf-8")
        self.upload_basis_document(
            "Direct production confirmation certificate is a required document candidate.",
            file_name="backup-basis-index.pdf",
        )

        response = self.client.post("/api/backups", json={})

        self.assertEqual(response.status_code, 201)
        backup = response.get_json()
        self.assertEqual(backup["status"], "completed")
        backup_path = Path(backup["file_path"])
        self.assertTrue(backup_path.exists())
        with zipfile.ZipFile(backup_path, "r") as zip_file:
            names = set(zip_file.namelist())
        self.assertIn("manifest.json", names)
        self.assertIn("database/app.db", names)
        self.assertIn("storage/uploads/sample.txt", names)
        self.assertIn("storage/basis-index/basis-index.json", names)
        self.assertNotIn("storage/uploads/.env", names)
        self.assertNotIn("storage/uploads/nested/.env.local", names)
        self.assertNotIn(".env", names)
        self.assertFalse(any(name.endswith(".env") or "/.env" in name for name in names))
        self.assertEqual(backup["manifest"]["basis_index"]["status"], "included")
        self.assertTrue(backup["manifest"]["basis_index"]["sha256"])

        validate_response = self.client.post("/api/backups/validate", json={"backup_id": backup["id"]})
        self.assertEqual(validate_response.status_code, 200)
        validation = validate_response.get_json()
        self.assertTrue(validation["valid"])

    def test_phase4d_backup_uses_sqlite_snapshot_and_removes_temp_file(self) -> None:
        corporation = self.create_corporation()

        response = self.client.post("/api/backups", json={})

        self.assertEqual(response.status_code, 201)
        backup = response.get_json()
        backup_path = Path(backup["file_path"])
        self.assertTrue(backup_path.exists())
        self.assertFalse(list(backup_path.parent.glob("*.db.snapshot.tmp")))

        with zipfile.ZipFile(backup_path, "r") as zip_file:
            db_bytes = zip_file.read("database/app.db")
        snapshot_db = Path(tempfile.mkdtemp(prefix="wisdom_backup_snapshot_")) / "snapshot.db"
        snapshot_db.write_bytes(db_bytes)
        try:
            with sqlite3.connect(snapshot_db) as conn:
                row = conn.execute("SELECT name FROM corporations WHERE id=?", (corporation["id"],)).fetchone()
            self.assertEqual(row[0], corporation["name"])
        finally:
            shutil.rmtree(snapshot_db.parent, ignore_errors=True)

    def test_phase4d_failed_backup_returns_payload_and_records_operation(self) -> None:
        previous_create_backup_run = runtime.create_backup_run

        def fake_create_backup_run(conn, *, sqlite_path, storage_root):
            now = runtime.now_iso()
            validation = {"valid": False, "errors": ["forced validation failure"], "warnings": []}
            cur = conn.execute(
                """
                INSERT INTO backup_runs (
                  backup_type, status, file_name, file_path, file_size_bytes,
                  manifest_json, validation_json, error_message,
                  created_at, completed_at, updated_at
                ) VALUES ('full_local', 'failed', 'failed.zip', '', 0,
                  '{}', ?, 'forced validation failure', ?, ?, ?)
                """,
                (json.dumps(validation, ensure_ascii=False), now, now, now),
            )
            return runtime.get_backup_run_payload(conn, cur.lastrowid)

        try:
            runtime.create_backup_run = fake_create_backup_run
            response = self.client.post("/api/backups", json={})
        finally:
            runtime.create_backup_run = previous_create_backup_run

        self.assertEqual(response.status_code, 201)
        payload = response.get_json()
        self.assertEqual(payload["status"], "failed")
        self.assertEqual(payload["validation"]["errors"], ["forced validation failure"])

        operation_runs = self.client.get("/api/operation-runs?operation_type=backup_create").get_json()
        self.assertEqual(operation_runs[0]["status"], "failed")

    def test_phase4d_exception_backup_failure_returns_validation_shape(self) -> None:
        previous_create_backup_run = runtime.create_backup_run

        def fake_create_backup_run(conn, *, sqlite_path, storage_root):
            return previous_create_backup_run(
                conn,
                sqlite_path=Path(storage_root) / "missing-app.db",
                storage_root=storage_root,
            )

        try:
            runtime.create_backup_run = fake_create_backup_run
            response = self.client.post("/api/backups", json={})
        finally:
            runtime.create_backup_run = previous_create_backup_run

        self.assertEqual(response.status_code, 201)
        payload = response.get_json()
        self.assertEqual(payload["status"], "failed")
        self.assertFalse(payload["validation"]["valid"])
        self.assertTrue(payload["validation"]["errors"])
        self.assertIn("SQLite DB not found", payload["validation"]["errors"][0])
        self.assertIsInstance(payload["validation"]["warnings"], list)
        self.assertIsInstance(payload["validation"]["manifest"], dict)

    def test_phase4d_backup_file_path_is_limited_to_backup_directory(self) -> None:
        outside_zip = Path(tempfile.mkdtemp(prefix="wisdom_external_backup_")) / "external.zip"
        with zipfile.ZipFile(outside_zip, "w") as zip_file:
            zip_file.writestr("manifest.json", "{}")
        try:
            response = self.client.post("/api/backups/validate", json={"file_path": str(outside_zip)})
        finally:
            shutil.rmtree(outside_zip.parent, ignore_errors=True)

        self.assertEqual(response.status_code, 400)

    def test_phase4d_restore_dry_run_rejects_backup_path_outside_backup_directory(self) -> None:
        outside_zip = Path(tempfile.mkdtemp(prefix="wisdom_external_restore_")) / "external.zip"
        with zipfile.ZipFile(outside_zip, "w") as zip_file:
            zip_file.writestr("manifest.json", "{}")
        now = runtime.now_iso()
        try:
            with runtime.db_conn() as conn:
                cur = conn.execute(
                    """
                    INSERT INTO backup_runs (
                      backup_type, status, file_name, file_path, file_size_bytes,
                      manifest_json, validation_json, error_message,
                      created_at, completed_at, updated_at
                    ) VALUES ('full_local', 'completed', 'external.zip', ?, 0,
                      '{}', '{}', '', ?, ?, ?)
                    """,
                    (str(outside_zip), now, now, now),
                )
                backup_id = cur.lastrowid

            response = self.client.post(f"/api/backups/{backup_id}/restore", json={"dry_run": True})
        finally:
            shutil.rmtree(outside_zip.parent, ignore_errors=True)

        self.assertEqual(response.status_code, 400)
        self.assertIn("outside the allowed backup directory", response.get_json()["detail"])

    def test_phase4d_restore_plan_is_dry_run_and_blocks_direct_restore(self) -> None:
        create_response = self.client.post("/api/backups", json={})
        self.assertEqual(create_response.status_code, 201)
        backup = create_response.get_json()

        plan_response = self.client.post("/api/backups/restore-plan", json={"backup_id": backup["id"]})

        self.assertEqual(plan_response.status_code, 200)
        plan = plan_response.get_json()
        self.assertTrue(plan["dry_run"])
        self.assertTrue(plan["can_restore"])

        direct_response = self.client.post(f"/api/backups/{backup['id']}/restore", json={"dry_run": False})
        self.assertEqual(direct_response.status_code, 400)

        dry_run_response = self.client.post(f"/api/backups/{backup['id']}/restore", json={"dry_run": True})
        self.assertEqual(dry_run_response.status_code, 200)
        self.assertTrue(dry_run_response.get_json()["dry_run"])

    def test_phase5a_contract_document_table_indexes_and_status_constraints(self) -> None:
        with runtime.db_conn() as conn:
            table = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='contract_documents'"
            ).fetchone()
            self.assertIsNotNone(table)

            indexes = {
                row["name"]
                for row in conn.execute(
                    "SELECT name FROM sqlite_master WHERE type='index' AND tbl_name='contract_documents'"
                ).fetchall()
            }
            self.assertIn("idx_contract_documents_notice_corporation", indexes)
            self.assertIn("idx_contract_documents_created_at", indexes)
            self.assertIn("idx_contract_documents_status", indexes)

            now = runtime.now_iso()
            conn.execute(
                """
                INSERT INTO contract_documents (
                  nara_notice_id, corporation_id, title, file_name,
                  stored_file_path, created_at, updated_at
                ) VALUES (1, 1, '계약서 초안', 'contract-1.docx',
                  'contracts/1/contract-1.docx', ?, ?)
                """,
                (now, now),
            )

            with self.assertRaises(sqlite3.IntegrityError):
                conn.execute(
                    """
                    INSERT INTO contract_documents (
                      nara_notice_id, corporation_id, status, created_at, updated_at
                    ) VALUES (1, 1, 'unknown', ?, ?)
                    """,
                    (now, now),
                )

            with self.assertRaises(sqlite3.IntegrityError):
                conn.execute(
                    """
                    INSERT INTO contract_documents (
                      nara_notice_id, corporation_id, review_status, created_at, updated_at
                    ) VALUES (1, 1, 'unknown', ?, ?)
                    """,
                    (now, now),
                )

    def test_phase5a_contract_storage_path_and_backup_policy(self) -> None:
        from app.services.contract_documents import contract_output_path

        storage_root = Path(os.environ["STORAGE_ROOT"])
        output_path = contract_output_path(storage_root, 10, "계약서 초안.docx")
        self.assertTrue(output_path.is_relative_to(storage_root / "contracts" / "10"))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(make_docx_bytes("계약서 백업 샘플"))
        (output_path.parent / ".env").write_text("SECRET=contract-secret", encoding="utf-8")

        response = self.client.post("/api/backups", json={})

        self.assertEqual(response.status_code, 201)
        backup = response.get_json()
        self.assertEqual(backup["status"], "completed")
        backup_path = Path(backup["file_path"])
        with zipfile.ZipFile(backup_path, "r") as zip_file:
            names = set(zip_file.namelist())
        self.assertIn("storage/contracts/10/계약서 초안.docx", names)
        self.assertNotIn("storage/contracts/10/.env", names)

    def setUp(self) -> None:
        runtime.app.config["TESTING"] = True
        self.client = runtime.app.test_client()
        with runtime.db_conn() as conn:
            conn.executescript(
                """
                DELETE FROM analyses;
                DELETE FROM backup_runs;
                DELETE FROM operation_runs;
                DELETE FROM nara_collection_runs;
                DELETE FROM contract_documents;
                DELETE FROM judgment_runs;
                DELETE FROM basis_retrieval_evaluations;
                DELETE FROM basis_rule_candidates;
                DELETE FROM basis_document_chunks;
                DELETE FROM basis_documents;
                DELETE FROM corporation_profile_update_candidates;
                DELETE FROM corporation_evidence_documents;
                DELETE FROM nara_notice_attachments;
                DELETE FROM notice_corporation_comparisons;
                DELETE FROM notice_requirement_candidates;
                DELETE FROM nara_notices;
                DELETE FROM integration_test_results;
                DELETE FROM project_documents;
                DELETE FROM projects;
                DELETE FROM corporations;
                """
            )
            conn.commit()

        storage_root = Path(os.environ["STORAGE_ROOT"])
        if storage_root.exists():
            shutil.rmtree(storage_root)
        runtime.init_db()

    def create_corporation(self, name: str = "테스트 법인") -> dict:
        response = self.client.post(
            "/api/corporations",
            json={
                "name": name,
                "business_category": "IT",
                "region": "서울",
                "certifications_json": ["ISO9001"],
                "company_size_classification": "중소기업",
                "internal_notes": "테스트",
            },
        )
        self.assertEqual(response.status_code, 201)
        return response.get_json()

    def create_project(self, corporation_id: int) -> dict:
        response = self.client.post(
            "/api/projects",
            json={
                "name": "테스트 프로젝트",
                "corporation_id": corporation_id,
                "status": "active",
                "notes": "초기 메모",
            },
        )
        self.assertEqual(response.status_code, 201)
        return response.get_json()

    def create_saved_notice_row(self, *, title: str = "테스트 용역 공고", order: str = "00") -> int:
        now = runtime.now_iso()
        with runtime.db_conn() as conn:
            cur = conn.execute(
                """
                INSERT INTO nara_notices (
                  bid_ntce_no, bid_ntce_ord, bid_ntce_nm, ntce_instt_nm, dminstt_nm,
                  bid_ntce_dt, bid_begin_dt, bid_clse_dt, openg_dt, presmpt_prce,
                  bdgt_amt, bssamt, region_text, license_text, source_url, raw_json,
                  detail_json, save_status, download_status, analysis_status,
                  analysis_summary_json, analysis_summary_markdown, error_message,
                  created_at, updated_at
                ) VALUES ('20260607001', ?, ?, '서울시청', '서울시 수요기관',
                  '2026-06-01', '2026-06-02', '2026-06-10', '2026-06-11',
                  '1000000', '1200000', '900000', '서울', '소프트웨어사업자',
                  'https://example.test/notice', '{"ServiceKey":"SHOULD_NOT_LEAK"}',
                  '{}', 'saved', 'completed', 'completed', '{}',
                  '용역 수행 조건과 제출 서류 안내', '', ?, ?)
                """,
                (order, title, now, now),
            )
            return int(cur.lastrowid)

    def test_phase5a_contract_input_snapshot_success_and_secret_allowlist(self) -> None:
        from app.services.contract_documents import build_contract_input_snapshot

        corporation = self.create_corporation()
        notice_id = self.create_saved_notice_row()
        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE corporations
                SET representative_name='홍길동',
                    corporate_registration_number='110111-1234567',
                    business_address='서울시 중구 테스트로 1'
                WHERE id=?
                """,
                (corporation["id"],),
            )
            snapshot = build_contract_input_snapshot(
                conn,
                notice_id,
                corporation["id"],
                custom_fields={
                    "contract_amount": "1,100,000원",
                    "corporation_phone": "02-1234-5678",
                    "API_KEY": "SHOULD_NOT_LEAK",
                    "TOKEN": "SHOULD_NOT_LEAK",
                },
            )

        self.assertTrue(snapshot["validation"]["valid"])
        self.assertEqual(snapshot["template_version"], "contract_docx_template_v1")
        self.assertEqual(snapshot["notice"]["bid_ntce_nm"], "테스트 용역 공고")
        self.assertEqual(snapshot["generated_fields"]["buyer_name"], "서울시청")
        self.assertEqual(snapshot["generated_fields"]["corporation_representative_name"], "홍길동")
        self.assertEqual(snapshot["generated_fields"]["contract_amount"], "1,100,000원")
        self.assertNotIn("raw_json", snapshot["notice"])

        serialized = json.dumps(snapshot, ensure_ascii=False)
        self.assertNotIn("SHOULD_NOT_LEAK", serialized)
        self.assertNotIn("ServiceKey", serialized)
        self.assertNotIn("API_KEY", serialized)
        self.assertNotIn("TOKEN", serialized)

    def test_phase5a_contract_input_snapshot_validates_missing_and_mismatched_sources(self) -> None:
        from app.services.contract_documents import ContractInputError, build_contract_input_snapshot

        corporation = self.create_corporation()
        other_corporation = self.create_corporation("다른 테스트 법인")
        notice_id = self.create_saved_notice_row()
        now = runtime.now_iso()
        with runtime.db_conn() as conn:
            with self.assertRaises(ContractInputError) as missing_notice:
                build_contract_input_snapshot(conn, 999999, corporation["id"])
            self.assertEqual(missing_notice.exception.code, "notice_not_found")
            self.assertEqual(missing_notice.exception.status_code, 404)

            cur = conn.execute(
                """
                INSERT INTO judgment_runs (
                  nara_notice_id, corporation_id, status, review_status,
                  input_snapshot_json, result_json, summary_json,
                  created_at, updated_at
                ) VALUES (?, ?, 'completed', 'pending', '{}', '{}', '{}', ?, ?)
                """,
                (notice_id, other_corporation["id"], now, now),
            )

            with self.assertRaises(ContractInputError) as mismatch:
                build_contract_input_snapshot(conn, notice_id, corporation["id"], judgment_run_id=cur.lastrowid)
            self.assertEqual(mismatch.exception.code, "judgment_run_mismatch")
            self.assertEqual(mismatch.exception.status_code, 400)

    def test_phase5a_contract_input_snapshot_keeps_existing_object_immutable(self) -> None:
        from app.services.contract_documents import build_contract_input_snapshot

        corporation = self.create_corporation()
        notice_id = self.create_saved_notice_row(title="초기 용역 공고")
        with runtime.db_conn() as conn:
            snapshot = build_contract_input_snapshot(conn, notice_id, corporation["id"])
            conn.execute(
                "UPDATE nara_notices SET bid_ntce_nm='수정된 용역 공고' WHERE id=?",
                (notice_id,),
            )
            rebuilt = build_contract_input_snapshot(conn, notice_id, corporation["id"])

        self.assertEqual(snapshot["notice"]["bid_ntce_nm"], "초기 용역 공고")
        self.assertEqual(snapshot["generated_fields"]["service_name"], "초기 용역 공고")
        self.assertEqual(rebuilt["notice"]["bid_ntce_nm"], "수정된 용역 공고")

    def test_phase5a_contract_docx_builder_generates_standard_form_labels(self) -> None:
        from app.services.contract_documents import build_contract_input_snapshot, render_standard_service_contract_docx

        corporation = self.create_corporation()
        notice_id = self.create_saved_notice_row(title="2026년 테스트 용역")
        output_path = Path(os.environ["STORAGE_ROOT"]) / "contracts" / "docx-builder" / "contract.docx"
        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE corporations
                SET representative_name='홍길동',
                    corporate_registration_number='110111-1234567',
                    business_address='서울시 중구 테스트로 1'
                WHERE id=?
                """,
                (corporation["id"],),
            )
            snapshot = build_contract_input_snapshot(
                conn,
                notice_id,
                corporation["id"],
                custom_fields={"contract_number": "C-2026-001", "corporation_phone": "02-1234-5678"},
            )

        render_standard_service_contract_docx(snapshot, output_path)

        self.assertTrue(output_path.exists())
        doc = Document(output_path)
        text_parts = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        text = "\n".join(text_parts)
        for expected in ["용역표준계약서", "계약서", "계약내용", "발주처", "계약상대자", "붙임서류"]:
            self.assertIn(expected, text)
        self.assertIn("2026년 테스트 용역", text)
        self.assertIn("20260607001", text)
        self.assertIn("테스트 법인", text)
        self.assertGreaterEqual(len(doc.tables), 2)
        section = doc.sections[0]
        self.assertAlmostEqual(section.page_width.mm, 210, delta=1)
        self.assertAlmostEqual(section.page_height.mm, 297, delta=1)

    def test_phase5a_contract_docx_builder_does_not_leave_partial_file_on_failure(self) -> None:
        from app.services.contract_documents import build_contract_input_snapshot, render_standard_service_contract_docx

        corporation = self.create_corporation()
        notice_id = self.create_saved_notice_row()
        blocking_parent = Path(os.environ["STORAGE_ROOT"]) / "contracts" / "blocked"
        blocking_parent.parent.mkdir(parents=True, exist_ok=True)
        blocking_parent.write_text("not a directory", encoding="utf-8")
        output_path = blocking_parent / "contract.docx"
        with runtime.db_conn() as conn:
            snapshot = build_contract_input_snapshot(conn, notice_id, corporation["id"])

        with self.assertRaises(OSError):
            render_standard_service_contract_docx(snapshot, output_path)

        self.assertFalse(output_path.exists())

    def test_phase5a_contract_api_create_list_download_review_and_delete(self) -> None:
        corporation = self.create_corporation()
        notice_id = self.create_saved_notice_row(title="계약 API 테스트 용역")

        preview_response = self.client.post(
            "/api/contracts/preview",
            json={
                "nara_notice_id": notice_id,
                "corporation_id": corporation["id"],
                "custom_fields": {"contract_number": "API-001", "corporation_phone": "02-0000-0000"},
            },
        )
        self.assertEqual(preview_response.status_code, 200)
        self.assertTrue(preview_response.get_json()["valid"])

        create_response = self.client.post(
            "/api/contracts",
            json={
                "nara_notice_id": notice_id,
                "corporation_id": corporation["id"],
                "title": "계약 API 테스트 계약서 초안",
                "custom_fields": {"contract_number": "API-001", "corporation_phone": "02-0000-0000"},
            },
        )
        self.assertEqual(create_response.status_code, 201)
        created = create_response.get_json()
        self.assertEqual(created["status"], "generated")
        self.assertEqual(created["review_status"], "draft")
        self.assertTrue(created["download_url"])

        list_response = self.client.get(f"/api/contracts?notice_id={notice_id}&review_status=draft&keyword=API")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(len(list_response.get_json()), 1)

        detail_response = self.client.get(f"/api/contracts/{created['id']}")
        self.assertEqual(detail_response.status_code, 200)
        self.assertEqual(detail_response.get_json()["id"], created["id"])

        download_response = self.client.get(f"/api/contracts/{created['id']}/download")
        self.assertEqual(download_response.status_code, 200)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            download_response.content_type,
        )
        self.assertIn("filename*=UTF-8''", download_response.headers["Content-Disposition"])
        downloaded = Document(io.BytesIO(download_response.data))
        downloaded_text = "\n".join(paragraph.text for paragraph in downloaded.paragraphs)
        self.assertIn("용역표준계약서", downloaded_text)
        download_response.close()

        review_response = self.client.patch(
            f"/api/contracts/{created['id']}/review",
            json={"review_status": "approved", "review_note": "검토 완료"},
        )
        self.assertEqual(review_response.status_code, 200)
        self.assertEqual(review_response.get_json()["review_status"], "approved")
        self.assertEqual(review_response.get_json()["review_note"], "검토 완료")

        operation_runs = self.client.get("/api/operation-runs?operation_type=contract_create").get_json()
        self.assertEqual(operation_runs[0]["status"], "completed")

        delete_response = self.client.delete(f"/api/contracts/{created['id']}")
        self.assertEqual(delete_response.status_code, 200)
        self.assertEqual(delete_response.get_json()["status"], "deleted")
        missing_response = self.client.get(f"/api/contracts/{created['id']}")
        self.assertEqual(missing_response.status_code, 404)

    def test_phase5a_contract_api_failed_create_records_failure_reason(self) -> None:
        corporation = self.create_corporation()
        notice_id = self.create_saved_notice_row(title="")

        response = self.client.post(
            "/api/contracts",
            json={"nara_notice_id": notice_id, "corporation_id": corporation["id"]},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.get_json()
        self.assertEqual(payload["status"], "failed")
        self.assertFalse(payload["validation"]["valid"])
        self.assertTrue(payload["validation"]["errors"])
        self.assertTrue(payload["error_message"])

        operation_runs = self.client.get("/api/operation-runs?operation_type=contract_create").get_json()
        self.assertEqual(operation_runs[0]["status"], "failed")
        self.assertEqual(operation_runs[0]["error_code"], "contract_generation_failed")

    def upload_document(self, project_id: int) -> dict:
        response = self.client.post(
            "/api/documents",
            data={
                "project_id": str(project_id),
                "document_type": "notice",
                "memo": "업로드 메모",
                "revision_note": "r1",
                "file": (io.BytesIO(make_pdf_bytes("Phase 1 API flow notice")), "notice.pdf"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 201)
        return response.get_json()

    def upload_business_registration_evidence(
        self,
        corporation_id: int | None = None,
        management_group_name: str = "기본 관리그룹",
    ) -> dict:
        evidence_text = """
        사업자등록증
        등록번호 : 142-81-28387
        법인명(단체명) : 주식회사 온세이엔씨
        대표자 : 안영식
        개업연월일 : 2010 년 05 월 25 일
        법인등록번호 : 134511-0154712
        사업장 소재지 : 경기도 성남시 수정구 청계산로 686, 8층 820호
        사업의 종류 : 업태 건설업
        종목 전기공사
        """
        data = {
            "document_type": "auto",
            "management_group_name": management_group_name,
            "memo": "사업자등록증 자동 추출 테스트",
            "file": (io.BytesIO(make_docx_bytes(evidence_text)), "business-registration.docx"),
        }
        if corporation_id:
            data["corporation_id"] = str(corporation_id)

        response = self.client.post(
            "/api/corporation-evidence-documents",
            data=data,
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 201)
        return response.get_json()

    def upload_basis_document(self, text: str, file_name: str = "basis.pdf", force_ocr: bool | None = None) -> dict:
        data = {
            "title": "지방계약 기준문서",
            "category": "local_contract",
            "document_version": "2026.05",
            "issuing_agency": "테스트 기관",
            "effective_date": "2026-05-22",
            "memo": "Phase 2 테스트 기준문서",
            "file": (io.BytesIO(make_pdf_bytes(text)), file_name),
        }
        if force_ocr is not None:
            data["force_ocr"] = "true" if force_ocr else "false"
        response = self.client.post(
            "/api/basis-documents",
            data=data,
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 201)
        return response.get_json()

    def wait_for_saved_nara_notice(self, notice_id: int, timeout: float = 5.0) -> dict:
        deadline = time.time() + timeout
        last_payload: dict | None = None
        while time.time() < deadline:
            response = self.client.get(f"/api/nara/saved-notices/{notice_id}")
            self.assertEqual(response.status_code, 200)
            last_payload = response.get_json()
            if last_payload["analysis_status"] not in {"pending", "saving", "queued"}:
                return last_payload
            time.sleep(0.05)

        self.fail(f"Nara notice job did not finish in time. last_payload={last_payload}")

    def upload_small_business_evidence(self, corporation_id: int) -> dict:
        evidence_text = """
        중소기업확인서
        기업명 : 테스트 법인
        사업자등록번호 : 142-81-28387
        대표자 : 안영식
        기업규모 : 소기업
        유효기간 : 2026.01.01 ~ 2026.12.31
        """
        response = self.client.post(
            "/api/corporation-evidence-documents",
            data={
                "corporation_id": str(corporation_id),
                "document_type": "auto",
                "memo": "중소기업확인서 자동 추출 테스트",
                "file": (io.BytesIO(make_docx_bytes(evidence_text)), "small-business.docx"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 201)
        return response.get_json()

    def test_service_rocket_pitch_demo_pipeline_flow(self) -> None:
        business_evidence = self.upload_business_registration_evidence()
        self.assertEqual(business_evidence["document_type"], "business_registration_certificate")
        self.assertEqual(business_evidence["classification_status"], "classified")

        business_approval = self.client.post(
            f"/api/corporation-evidence-documents/{business_evidence['id']}/approve",
            json={},
        )
        self.assertEqual(business_approval.status_code, 200)
        corporation = business_approval.get_json()["corporation"]
        self.assertEqual(corporation["name"], "주식회사 온세이엔씨")
        self.assertEqual(corporation["evidence_verification_status"], "evidence_reviewed")

        small_business_evidence = self.upload_small_business_evidence(corporation["id"])
        small_business_approval = self.client.post(
            f"/api/corporation-evidence-documents/{small_business_evidence['id']}/approve",
            json={},
        )
        self.assertEqual(small_business_approval.status_code, 200)
        corporation = small_business_approval.get_json()["corporation"]
        self.assertEqual(corporation["company_size_classification"], "소기업")
        self.assertIn("중소기업확인서", json.loads(corporation["certifications_json"]))

        dashboard_response = self.client.get("/api/dashboard/summary")
        self.assertEqual(dashboard_response.status_code, 200)
        self.assertGreaterEqual(dashboard_response.get_json()["corporation_count"], 1)

        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "DEMO20260614001",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "Rocket Pitch 시연용 정보통신 용역 공고",
                    "ntceInsttNm": "데모 발주기관",
                    "dminsttNm": "데모 수요기관",
                    "bidNtceDt": "2026-06-14 09:00",
                    "bidBeginDt": "2026-06-15 09:00",
                    "bidClseDt": "2026-06-30 17:00",
                    "opengDt": "2026-07-01 10:00",
                    "presmptPrce": "25000000",
                    "bssamt": "23000000",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "information communication construction license",
                    "bidNtceDtlUrl": "https://example.go.kr/demo-notice",
                    "business_type": "service",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])
        self.assertEqual(notice["bid_ntce_no"], "DEMO20260614001")
        self.assertEqual(notice["business_type"], "service")

        requirements_response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/requirements")
        self.assertEqual(requirements_response.status_code, 200)
        requirements_payload = requirements_response.get_json()
        requirement_types = {item["requirement_type"] for item in requirements_payload["requirements"]}
        self.assertIn("region", requirement_types)
        self.assertIn("license", requirement_types)
        self.assertIn("date", requirement_types)

        structured_response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/requirements/structured")
        self.assertEqual(structured_response.status_code, 200)
        structured_payload = structured_response.get_json()
        self.assertEqual(structured_payload["contract_version"], "phase3_gap_judgment_contract_v1")
        self.assertGreaterEqual(structured_payload["requirement_count"], 3)

        basis = self.upload_basis_document(
            """
            Information communication construction license is a citation candidate for bidder qualification review.
            Small business certificate is a required document candidate for company type review.
            Business registration certificate must be submitted before contract review.
            Contract officers should verify missing documents before deciding readiness.
            """,
            file_name="rocket-pitch-demo-basis.pdf",
        )
        self.assertEqual(basis["processing_status"], "completed")
        self.assertEqual(basis["index_status"], "completed")
        self.assertGreaterEqual(basis["chunk_count"], 1)

        basis_search_response = self.client.post(
            "/api/basis-search",
            json={"query": "information communication construction license", "top_k": 3},
        )
        self.assertEqual(basis_search_response.status_code, 200)
        basis_search = basis_search_response.get_json()
        self.assertEqual(basis_search["index_source"], "json_basis_index")
        self.assertGreaterEqual(basis_search["result_count"], 1)

        comparison_response = self.client.post(
            "/api/notice-comparisons",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"]},
        )
        self.assertEqual(comparison_response.status_code, 201)
        comparison = comparison_response.get_json()
        self.assertEqual(comparison["status"], "preview")
        self.assertGreaterEqual(comparison["summary"]["prepared_count"], 1)
        self.assertGreaterEqual(comparison["summary"]["possibly_missing_count"], 1)

        judgment_response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
        )
        self.assertEqual(judgment_response.status_code, 201)
        judgment = judgment_response.get_json()
        self.assertEqual(judgment["status"], "completed")
        self.assertEqual(judgment["review_status"], "pending")
        self.assertGreaterEqual(judgment["summary"]["missing_count"], 1)
        self.assertTrue(judgment["result"]["preparation_guide"])
        self.assertTrue(any(item["citation_candidates"] for item in judgment["result"]["items"]))

        contract_preview_response = self.client.post(
            "/api/contracts/preview",
            json={
                "nara_notice_id": notice["id"],
                "corporation_id": corporation["id"],
                "judgment_run_id": judgment["id"],
                "custom_fields": {
                    "contract_number": "DEMO-2026-001",
                    "contract_amount": "25,000,000원",
                    "contract_period": "2026.07.01 ~ 2026.12.31",
                },
            },
        )
        self.assertEqual(contract_preview_response.status_code, 200)
        self.assertTrue(contract_preview_response.get_json()["valid"])

        contract_response = self.client.post(
            "/api/contracts",
            json={
                "nara_notice_id": notice["id"],
                "corporation_id": corporation["id"],
                "judgment_run_id": judgment["id"],
                "title": "Rocket Pitch 시연 계약서 초안",
                "custom_fields": {
                    "contract_number": "DEMO-2026-001",
                    "contract_amount": "25,000,000원",
                    "contract_period": "2026.07.01 ~ 2026.12.31",
                },
            },
        )
        self.assertEqual(contract_response.status_code, 201)
        contract = contract_response.get_json()
        self.assertEqual(contract["status"], "generated")
        self.assertEqual(contract["review_status"], "draft")
        self.assertTrue(contract["download_url"])

        download_response = self.client.get(f"/api/contracts/{contract['id']}/download")
        self.assertEqual(download_response.status_code, 200)
        try:
            downloaded = Document(io.BytesIO(download_response.data))
            downloaded_text_parts = [paragraph.text for paragraph in downloaded.paragraphs]
            for table in downloaded.tables:
                for row in table.rows:
                    for cell in row.cells:
                        downloaded_text_parts.append(cell.text)
            downloaded_text = "\n".join(downloaded_text_parts)
            self.assertIn("용역표준계약서", downloaded_text)
            self.assertIn("Rocket Pitch 시연용 정보통신 용역 공고", downloaded_text)
        finally:
            download_response.close()

        operation_runs_response = self.client.get("/api/operation-runs")
        self.assertEqual(operation_runs_response.status_code, 200)
        operation_types = {item["operation_type"] for item in operation_runs_response.get_json()}
        self.assertIn("basis_document_processing", operation_types)
        self.assertIn("judgment_run", operation_types)
        self.assertIn("contract_create", operation_types)

        operations_response = self.client.get("/api/operations/summary")
        self.assertEqual(operations_response.status_code, 200)
        self.assertIn(operations_response.get_json()["overall_status"], {"ok", "warning", "action_required"})

        combined_payload = json.dumps(
            {
                "business_evidence": business_evidence,
                "small_business_evidence": small_business_evidence,
                "notice": notice,
                "requirements": requirements_payload,
                "basis_search": basis_search,
                "comparison": comparison,
                "judgment": judgment,
                "contract": contract,
            },
            ensure_ascii=False,
        )
        self.assertNotIn("eligible", combined_payload.lower())
        for forbidden_verdict in ['"match_status": "eligible"', '"status_label": "지원 가능"', '"eligible": true', '"eligibility"']:
            self.assertNotIn(forbidden_verdict, combined_payload)

    def test_crud_update_delete_flow(self) -> None:
        corporation = self.create_corporation()
        corp_patch = self.client.patch(
            f"/api/corporations/{corporation['id']}",
            json={"region": "부산", "certifications_json": "직접생산확인, ISO14001"},
        )
        self.assertEqual(corp_patch.status_code, 200)
        self.assertEqual(corp_patch.get_json()["region"], "부산")

        project = self.create_project(corporation["id"])
        project_patch = self.client.patch(
            f"/api/projects/{project['id']}",
            json={"name": "수정된 프로젝트", "status": "paused"},
        )
        self.assertEqual(project_patch.status_code, 200)
        self.assertEqual(project_patch.get_json()["status"], "paused")

        document = self.upload_document(project["id"])
        document_patch = self.client.patch(
            f"/api/documents/{document['id']}",
            json={"document_type": "spec", "memo": "수정된 메모", "revision_note": "r2"},
        )
        self.assertEqual(document_patch.status_code, 200)
        self.assertEqual(document_patch.get_json()["document_type"], "spec")

        document_delete = self.client.delete(f"/api/documents/{document['id']}")
        self.assertEqual(document_delete.status_code, 200)

        project_delete = self.client.delete(f"/api/projects/{project['id']}")
        self.assertEqual(project_delete.status_code, 200)

        corp_delete = self.client.delete(f"/api/corporations/{corporation['id']}")
        self.assertEqual(corp_delete.status_code, 200)

    def test_project_delete_removes_documents_analyses_and_files(self) -> None:
        corporation = self.create_corporation()
        project = self.create_project(corporation["id"])
        document = self.upload_document(project["id"])
        stored_path = Path(document["stored_file_path"])
        self.assertTrue(stored_path.exists())

        analysis = self.client.post(f"/api/documents/{document['id']}/analyze")
        self.assertEqual(analysis.status_code, 200)

        project_delete = self.client.delete(f"/api/projects/{project['id']}")
        self.assertEqual(project_delete.status_code, 200)
        self.assertFalse(stored_path.exists())

        with runtime.db_conn() as conn:
            doc_count = conn.execute("SELECT COUNT(*) c FROM project_documents").fetchone()["c"]
            analysis_count = conn.execute("SELECT COUNT(*) c FROM analyses").fetchone()["c"]

        self.assertEqual(doc_count, 0)
        self.assertEqual(analysis_count, 0)

    def test_phase2a_2b_basis_document_crud_and_pdf_guard(self) -> None:
        basis = self.upload_basis_document(
            """
            제1조 목적
            이 기준문서는 입찰 참가자격과 제출서류 관리 절차를 설명한다.
            사업자등록증, 납세증명서, 직접생산확인증명서는 검토 대상 자료이다.
            """
        )
        stored_path = Path(basis["stored_file_path"])
        self.assertTrue(stored_path.exists())
        self.assertEqual(basis["title"], "지방계약 기준문서")
        self.assertEqual(basis["category"], "local_contract")
        self.assertIn(basis["processing_status"], {"completed", "needs_ocr_setup"})

        invalid_response = self.client.post(
            "/api/basis-documents",
            data={
                "title": "잘못된 기준문서",
                "file": (io.BytesIO(make_docx_bytes("PDF가 아닌 기준문서")), "basis.docx"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(invalid_response.status_code, 400)

        list_response = self.client.get("/api/basis-documents")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(list_response.get_json()[0]["id"], basis["id"])

        patch_response = self.client.patch(
            f"/api/basis-documents/{basis['id']}",
            json={"title": "수정된 기준문서", "memo": "메모 수정"},
        )
        self.assertEqual(patch_response.status_code, 200)
        self.assertEqual(patch_response.get_json()["title"], "수정된 기준문서")

        delete_response = self.client.delete(f"/api/basis-documents/{basis['id']}")
        self.assertEqual(delete_response.status_code, 200)
        self.assertFalse(stored_path.exists())

        with runtime.db_conn() as conn:
            doc_count = conn.execute("SELECT COUNT(*) c FROM basis_documents").fetchone()["c"]
            chunk_count = conn.execute("SELECT COUNT(*) c FROM basis_document_chunks").fetchone()["c"]
        self.assertEqual(doc_count, 0)
        self.assertEqual(chunk_count, 0)

    def test_basis_document_force_ocr_option_is_stored_and_reprocessable(self) -> None:
        long_text = "\n".join(
            [
                "Article 1 Purpose. Extractable basis text is present.",
                "Article 2 Required documents include registration.",
                "Article 3 Direct production confirmation is required.",
                "Article 4 Operators may still force OCR during processing.",
            ]
        )

        basis = self.upload_basis_document(long_text, file_name="basis-force-ocr.pdf", force_ocr=True)

        self.assertTrue(basis["metadata"]["options"]["force_ocr"])
        self.assertEqual(basis["metadata"]["ocr"]["status"], "needs_ocr_setup")

        reprocess_response = self.client.post(
            f"/api/basis-documents/{basis['id']}/reprocess",
            json={"force_ocr": False},
        )

        self.assertEqual(reprocess_response.status_code, 200)
        reprocessed = reprocess_response.get_json()
        self.assertFalse(reprocessed["metadata"]["options"]["force_ocr"])
        self.assertEqual(reprocessed["metadata"]["ocr"]["status"], "skipped")

        with runtime.db_conn() as conn:
            operation = conn.execute(
                """
                SELECT request_json, result_json
                FROM operation_runs
                WHERE operation_type='basis_document_processing' AND target_id=?
                ORDER BY id DESC
                LIMIT 1
                """,
                (basis["id"],),
            ).fetchone()
        self.assertIsNotNone(operation)
        operation_request = json.loads(operation["request_json"])
        operation_result = json.loads(operation["result_json"])
        self.assertEqual(operation_request["options"]["force_ocr"], False)
        self.assertEqual(operation_result["processing_status"], "completed")

    def test_phase2c_2d_basis_processing_extracts_normalizes_and_chunks(self) -> None:
        basis = self.upload_basis_document(
            """
            Article 1 Purpose
            This basis document manages public procurement standards as reusable knowledge assets.

            Article 2 Required Documents
            Bidders should prepare business registration, national tax certificate, local tax certificate,
            and direct production confirmation documents.
            """
        )

        self.assertEqual(basis["processing_status"], "completed")
        self.assertEqual(basis["parse_status"], "completed")
        self.assertGreaterEqual(basis["chunk_count"], 1)
        self.assertIn("parser", basis["metadata"])
        self.assertIn("ocr", basis["metadata"])
        self.assertIn("chunker", basis["metadata"])
        self.assertIn("business registration", basis["extracted_text_preview"])

        chunk_response = self.client.get(f"/api/basis-documents/{basis['id']}/chunks")
        self.assertEqual(chunk_response.status_code, 200)
        chunks = chunk_response.get_json()
        self.assertGreaterEqual(len(chunks), 1)
        self.assertIn("chunker", chunks[0]["metadata"])
        self.assertTrue(chunks[0]["chunk_hash"])

        chunk_detail_response = self.client.get(f"/api/basis-documents/{basis['id']}/chunks/{chunks[0]['id']}")
        self.assertEqual(chunk_detail_response.status_code, 200)
        chunk_detail = chunk_detail_response.get_json()
        self.assertEqual(chunk_detail["detail_type"], "basis_chunk")
        self.assertEqual(chunk_detail["basis_document"]["id"], basis["id"])
        self.assertIn("business registration", chunk_detail["chunk_text"].lower())

        missing_chunk_response = self.client.get(f"/api/basis-documents/{basis['id']}/chunks/999999")
        self.assertEqual(missing_chunk_response.status_code, 404)

        reprocess_response = self.client.post(f"/api/basis-documents/{basis['id']}/reprocess")
        self.assertEqual(reprocess_response.status_code, 200)
        self.assertEqual(reprocess_response.get_json()["chunk_count"], len(chunks))

    def test_basis_reprocess_failure_preserves_existing_chunks_and_rule_candidates(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        basis = self.upload_basis_document(
            """
            Required document: business registration certificate must be submitted.
            License requirement: landscape construction business license must be held.
            """,
            file_name="basis-reprocess-safe.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        extracted_payload = extract_response.get_json()
        original_candidate_count = extracted_payload["candidate_count"]
        candidate = extracted_payload["candidates"][0]
        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 200)

        with runtime.db_conn() as conn:
            old_chunk_ids = [
                row["id"]
                for row in conn.execute(
                    "SELECT id FROM basis_document_chunks WHERE basis_document_id=? ORDER BY id",
                    (basis["id"],),
                ).fetchall()
            ]
            old_candidate_ids = [
                row["id"]
                for row in conn.execute(
                    "SELECT id FROM basis_rule_candidates WHERE basis_document_id=? ORDER BY id",
                    (basis["id"],),
                ).fetchall()
            ]

        previous_indexer = basis_pipeline.index_basis_chunks

        def failing_indexer(conn, basis_document_id, processing_run_id=None):
            raise RuntimeError("forced index failure")

        try:
            basis_pipeline.index_basis_chunks = failing_indexer
            response = self.client.post(f"/api/basis-documents/{basis['id']}/reprocess")
        finally:
            basis_pipeline.index_basis_chunks = previous_indexer

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.get_json()["processing_status"], "failed")
        with runtime.db_conn() as conn:
            new_chunk_ids = [
                row["id"]
                for row in conn.execute(
                    "SELECT id FROM basis_document_chunks WHERE basis_document_id=? ORDER BY id",
                    (basis["id"],),
                ).fetchall()
            ]
            new_candidate_ids = [
                row["id"]
                for row in conn.execute(
                    "SELECT id FROM basis_rule_candidates WHERE basis_document_id=? ORDER BY id",
                    (basis["id"],),
                ).fetchall()
            ]
        self.assertEqual(new_chunk_ids, old_chunk_ids)
        self.assertEqual(new_candidate_ids, old_candidate_ids)

    def test_basis_reprocess_missing_file_preserves_existing_index_and_search(self) -> None:
        basis = self.upload_basis_document(
            "Small business certificate submission is required for bidder qualification.",
            file_name="basis-reprocess-missing-file.pdf",
        )
        stored_path = Path(basis["stored_file_path"])
        stored_path.unlink()

        response = self.client.post(f"/api/basis-documents/{basis['id']}/reprocess")

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["processing_status"], "completed")
        self.assertEqual(payload["index_status"], "completed")
        self.assertIn("preserved", payload["error_message"])

        index_status_response = self.client.get("/api/basis-index/status")
        self.assertEqual(index_status_response.status_code, 200)
        index_status = index_status_response.get_json()
        self.assertTrue(index_status["valid"])
        self.assertTrue(index_status["can_search"])

        search_response = self.client.post(
            "/api/basis-search",
            json={"query": "small business certificate", "top_k": 3},
        )
        self.assertEqual(search_response.status_code, 200)
        self.assertGreaterEqual(search_response.get_json()["result_count"], 1)

    def test_basis_reprocess_swap_failure_preserves_existing_chunks_index_and_candidate_status(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        basis = self.upload_basis_document(
            """
            Required document: national tax certificate must be submitted.
            License requirement: information communication construction license must be held.
            """,
            file_name="basis-reprocess-swap-safe.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        extracted_payload = extract_response.get_json()
        original_candidate_count = extracted_payload["candidate_count"]
        candidate = extracted_payload["candidates"][0]
        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 200)

        with runtime.db_conn() as conn:
            old_chunk_ids = [
                row["id"]
                for row in conn.execute(
                    "SELECT id FROM basis_document_chunks WHERE basis_document_id=? ORDER BY id",
                    (basis["id"],),
                ).fetchall()
            ]
            old_candidate = conn.execute(
                "SELECT id, status FROM basis_rule_candidates WHERE id=?",
                (candidate["id"],),
            ).fetchone()
        old_index = basis_pipeline.load_basis_index()

        previous_marker = basis_pipeline.mark_basis_rule_candidates_for_revalidation

        def failing_marker(conn, basis_document_id):
            raise RuntimeError("forced swap failure after old chunk cleanup")

        try:
            basis_pipeline.mark_basis_rule_candidates_for_revalidation = failing_marker
            response = self.client.post(f"/api/basis-documents/{basis['id']}/reprocess")
        finally:
            basis_pipeline.mark_basis_rule_candidates_for_revalidation = previous_marker

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.get_json()["processing_status"], "failed")
        with runtime.db_conn() as conn:
            new_chunk_ids = [
                row["id"]
                for row in conn.execute(
                    "SELECT id FROM basis_document_chunks WHERE basis_document_id=? ORDER BY id",
                    (basis["id"],),
                ).fetchall()
            ]
            new_candidate = conn.execute(
                "SELECT id, status FROM basis_rule_candidates WHERE id=?",
                (candidate["id"],),
            ).fetchone()

        self.assertEqual(new_chunk_ids, old_chunk_ids)
        self.assertEqual(dict(new_candidate), dict(old_candidate))
        self.assertEqual(basis_pipeline.load_basis_index(), old_index)

    def test_phase2c_basis_blank_pdf_degrades_when_ocr_is_unavailable(self) -> None:
        response = self.client.post(
            "/api/basis-documents",
            data={
                "title": "OCR 필요 기준문서",
                "category": "ocr_required",
                "file": (io.BytesIO(make_blank_pdf_bytes()), "blank-basis.pdf"),
            },
            content_type="multipart/form-data",
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        self.assertEqual(payload["processing_status"], "needs_ocr_setup")
        self.assertEqual(payload["ocr_status"], "needs_ocr_setup")
        self.assertEqual(payload["chunk_count"], 0)
        self.assertEqual(payload["index_status"], "skipped")

    def test_phase2e_2f_basis_index_and_search_returns_candidates_only(self) -> None:
        basis = self.upload_basis_document(
            """
            Article 1 Scope
            This basis document is used when reviewing small business and women-owned business preferences.

            Article 2 Direct Production
            A direct production confirmation certificate is managed as a required document candidate
            for goods procurement bids.
            """
        )
        self.assertEqual(basis["index_status"], "completed")
        self.assertEqual(basis["vector_count"], basis["chunk_count"])

        search_response = self.client.post(
            "/api/basis-search",
            json={"query": "direct production certificate required document", "top_k": 3},
        )
        self.assertEqual(search_response.status_code, 200)
        payload = search_response.get_json()
        self.assertGreaterEqual(payload["result_count"], 1)
        self.assertEqual(payload["index_source"], "json_basis_index")
        self.assertEqual(payload["results"][0]["index_source"], "json_basis_index")
        self.assertTrue(payload["results"][0]["citation_candidate_id"].startswith("basis:"))
        self.assertIn("direct production confirmation certificate", payload["results"][0]["chunk"]["chunk_text"])
        self.assertNotIn("eligible", json.dumps(payload).lower())
        self.assertNotIn("지원 가능", json.dumps(payload, ensure_ascii=False))

    def test_basis_search_ranking_does_not_overvalue_repeated_single_token(self) -> None:
        repeated = self.upload_basis_document(
            "direct direct direct direct direct direct direct direct unrelated glossary",
            file_name="basis-repeated-token.pdf",
        )
        balanced = self.upload_basis_document(
            "direct production certificate is required for procurement review",
            file_name="basis-balanced-match.pdf",
        )

        search_response = self.client.post(
            "/api/basis-search",
            json={"query": "direct production certificate", "top_k": 2},
        )

        self.assertEqual(search_response.status_code, 200)
        payload = search_response.get_json()
        self.assertGreaterEqual(payload["result_count"], 2)
        first = payload["results"][0]
        self.assertEqual(first["document"]["id"], balanced["id"])
        self.assertNotEqual(first["document"]["id"], repeated["id"])
        self.assertLessEqual(first["score"], 1)

    def test_basis_search_excludes_failed_or_unindexed_chunks(self) -> None:
        basis = self.upload_basis_document(
            "Direct production confirmation certificate is required for goods procurement bids.",
            file_name="basis-search-active-filter.pdf",
        )

        search_response = self.client.post(
            "/api/basis-search",
            json={"query": "direct production confirmation certificate", "top_k": 3},
        )
        self.assertEqual(search_response.status_code, 200)
        self.assertGreaterEqual(search_response.get_json()["result_count"], 1)

        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE basis_documents
                SET processing_status='failed', index_status='failed'
                WHERE id=?
                """,
                (basis["id"],),
            )

        failed_search_response = self.client.post(
            "/api/basis-search",
            json={"query": "direct production confirmation certificate", "top_k": 3},
        )
        self.assertEqual(failed_search_response.status_code, 409)
        self.assertEqual(failed_search_response.get_json()["status"], "basis_index_unavailable")

        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE basis_documents
                SET processing_status='completed', index_status='completed'
                WHERE id=?
                """,
                (basis["id"],),
            )
            conn.execute(
                """
                UPDATE basis_document_chunks
                SET vector_status='pending', vector_id=''
                WHERE basis_document_id=?
                """,
                (basis["id"],),
            )

        unindexed_search_response = self.client.post(
            "/api/basis-search",
            json={"query": "direct production confirmation certificate", "top_k": 3},
        )
        self.assertEqual(unindexed_search_response.status_code, 409)
        self.assertEqual(unindexed_search_response.get_json()["status"], "basis_index_unavailable")

    def test_basis_index_status_detects_corruption_and_rebuild_restores_search(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        self.upload_basis_document(
            "Small business certificate submission is used for bidder qualification review.",
            file_name="basis-index-corrupt-rebuild.pdf",
        )
        status_response = self.client.get("/api/basis-index/status")
        self.assertEqual(status_response.status_code, 200)
        status_payload = status_response.get_json()
        self.assertTrue(status_payload["valid"])
        self.assertTrue(status_payload["can_search"])
        self.assertEqual(status_payload["status"], "ok")

        basis_pipeline.BASIS_INDEX_PATH.write_text("{not valid json", encoding="utf-8")
        corrupt_status_response = self.client.get("/api/basis-index/status")
        self.assertEqual(corrupt_status_response.status_code, 200)
        corrupt_payload = corrupt_status_response.get_json()
        self.assertEqual(corrupt_payload["status"], "corrupt")
        self.assertFalse(corrupt_payload["valid"])
        self.assertTrue(corrupt_payload["rebuild_required"])

        blocked_search_response = self.client.post(
            "/api/basis-search",
            json={"query": "small business certificate", "top_k": 3},
        )
        self.assertEqual(blocked_search_response.status_code, 409)
        self.assertEqual(blocked_search_response.get_json()["status"], "basis_index_unavailable")

        rebuild_response = self.client.post("/api/basis-index/rebuild", json={})
        self.assertEqual(rebuild_response.status_code, 200)
        rebuild_payload = rebuild_response.get_json()
        self.assertTrue(rebuild_payload["valid"])
        self.assertTrue(rebuild_payload["can_search"])
        self.assertGreaterEqual(rebuild_payload["rebuilt_chunk_count"], 1)
        self.assertTrue(Path(rebuild_payload["archived_path"]).exists())

        restored_search_response = self.client.post(
            "/api/basis-search",
            json={"query": "small business certificate", "top_k": 3},
        )
        self.assertEqual(restored_search_response.status_code, 200)
        self.assertGreaterEqual(restored_search_response.get_json()["result_count"], 1)

    def test_basis_document_delete_returns_409_when_basis_index_is_corrupt(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        basis = self.upload_basis_document(
            "Direct production confirmation certificate is required for goods procurement bids.",
            file_name="basis-delete-corrupt-index.pdf",
        )
        with runtime.db_conn() as conn:
            chunk_count_before = conn.execute(
                "SELECT COUNT(*) FROM basis_document_chunks WHERE basis_document_id=?",
                (basis["id"],),
            ).fetchone()[0]
        self.assertGreaterEqual(chunk_count_before, 1)

        basis_pipeline.BASIS_INDEX_PATH.write_text("{not valid json", encoding="utf-8")
        delete_response = self.client.delete(f"/api/basis-documents/{basis['id']}")

        self.assertEqual(delete_response.status_code, 409)
        delete_payload = delete_response.get_json()
        self.assertEqual(delete_payload["status"], "basis_index_unavailable")
        self.assertTrue(delete_payload["rebuild_required"])
        with runtime.db_conn() as conn:
            basis_row = conn.execute("SELECT id FROM basis_documents WHERE id=?", (basis["id"],)).fetchone()
            chunk_count_after = conn.execute(
                "SELECT COUNT(*) FROM basis_document_chunks WHERE basis_document_id=?",
                (basis["id"],),
            ).fetchone()[0]
        self.assertIsNotNone(basis_row)
        self.assertEqual(chunk_count_after, chunk_count_before)

        rebuild_response = self.client.post("/api/basis-index/rebuild", json={})
        self.assertEqual(rebuild_response.status_code, 200)
        self.assertTrue(rebuild_response.get_json()["valid"])
        retry_delete_response = self.client.delete(f"/api/basis-documents/{basis['id']}")
        self.assertEqual(retry_delete_response.status_code, 200)
        self.assertEqual(retry_delete_response.get_json()["status"], "deleted")

    def test_basis_document_delete_returns_409_when_basis_index_is_missing(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        first = self.upload_basis_document(
            "Small business confirmation is required for restricted procurement bids.",
            file_name="basis-delete-missing-index-1.pdf",
        )
        second = self.upload_basis_document(
            "Direct production confirmation is required for goods procurement bids.",
            file_name="basis-delete-missing-index-2.pdf",
        )
        with runtime.db_conn() as conn:
            indexed_chunk_count_before = conn.execute(
                "SELECT COUNT(*) FROM basis_document_chunks WHERE vector_status='indexed'"
            ).fetchone()[0]
        self.assertGreaterEqual(indexed_chunk_count_before, 2)

        basis_pipeline.BASIS_INDEX_PATH.unlink()
        delete_response = self.client.delete(f"/api/basis-documents/{first['id']}")

        self.assertEqual(delete_response.status_code, 409)
        delete_payload = delete_response.get_json()
        self.assertEqual(delete_payload["status"], "basis_index_unavailable")
        self.assertEqual(delete_payload["index_status"], "missing")
        self.assertTrue(delete_payload["rebuild_required"])
        with runtime.db_conn() as conn:
            remaining_documents = conn.execute("SELECT COUNT(*) FROM basis_documents").fetchone()[0]
            indexed_chunk_count_after = conn.execute(
                "SELECT COUNT(*) FROM basis_document_chunks WHERE vector_status='indexed'"
            ).fetchone()[0]
        self.assertEqual(remaining_documents, 2)
        self.assertEqual(indexed_chunk_count_after, indexed_chunk_count_before)

        rebuild_response = self.client.post("/api/basis-index/rebuild", json={})
        self.assertEqual(rebuild_response.status_code, 200)
        self.assertTrue(rebuild_response.get_json()["valid"])
        retry_delete_response = self.client.delete(f"/api/basis-documents/{first['id']}")
        self.assertEqual(retry_delete_response.status_code, 200)
        self.assertEqual(retry_delete_response.get_json()["status"], "deleted")
        search_response = self.client.post("/api/basis-search", json={"query": "direct production", "top_k": 3})
        self.assertEqual(search_response.status_code, 200)
        self.assertIn(str(second["id"]), search_response.get_json()["results"][0]["citation_candidate_id"])

    def test_basis_document_delete_returns_409_when_basis_index_is_inconsistent(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        first = self.upload_basis_document(
            "Small business confirmation is required for restricted procurement bids.",
            file_name="basis-delete-inconsistent-index-1.pdf",
        )
        second = self.upload_basis_document(
            "Direct production confirmation is required for goods procurement bids.",
            file_name="basis-delete-inconsistent-index-2.pdf",
        )
        with runtime.db_conn() as conn:
            indexed_chunk_count_before = conn.execute(
                "SELECT COUNT(*) FROM basis_document_chunks WHERE vector_status='indexed'"
            ).fetchone()[0]
        self.assertGreaterEqual(indexed_chunk_count_before, 2)

        payload = json.loads(basis_pipeline.BASIS_INDEX_PATH.read_text(encoding="utf-8"))
        chunks = payload["chunks"]
        removed_vector_id = next(
            vector_id
            for vector_id, item in chunks.items()
            if int(item["basis_document_id"]) == second["id"]
        )
        chunks.pop(removed_vector_id)
        basis_pipeline.save_basis_index(payload)

        delete_response = self.client.delete(f"/api/basis-documents/{first['id']}")

        self.assertEqual(delete_response.status_code, 409)
        delete_payload = delete_response.get_json()
        self.assertEqual(delete_payload["status"], "basis_index_unavailable")
        self.assertEqual(delete_payload["index_status"], "inconsistent")
        self.assertTrue(delete_payload["rebuild_required"])
        with runtime.db_conn() as conn:
            remaining_documents = conn.execute("SELECT COUNT(*) FROM basis_documents").fetchone()[0]
            indexed_chunk_count_after = conn.execute(
                "SELECT COUNT(*) FROM basis_document_chunks WHERE vector_status='indexed'"
            ).fetchone()[0]
        self.assertEqual(remaining_documents, 2)
        self.assertEqual(indexed_chunk_count_after, indexed_chunk_count_before)

        rebuild_response = self.client.post("/api/basis-index/rebuild", json={})
        self.assertEqual(rebuild_response.status_code, 200)
        self.assertTrue(rebuild_response.get_json()["valid"])
        retry_delete_response = self.client.delete(f"/api/basis-documents/{first['id']}")
        self.assertEqual(retry_delete_response.status_code, 200)
        self.assertEqual(retry_delete_response.get_json()["status"], "deleted")

    def test_phase2_closeout_summary_documents_baseline_and_known_issues(self) -> None:
        known_issue_path = _TMP_DIR / "qa-known-issues.json"
        known_issue_path.write_text(
            json.dumps(
                [
                    {
                        "type": "mupdf_internal_syntax_warning",
                        "severity": "low",
                        "sample_path": "backend/tests/nara-notice-pdf-samples/sample.pdf",
                    }
                ],
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        previous_path = os.environ.get("NARA_NOTICE_PDF_KNOWN_ISSUES_PATH")
        os.environ["NARA_NOTICE_PDF_KNOWN_ISSUES_PATH"] = str(known_issue_path)
        try:
            response = self.client.get("/api/qa/phase2-closeout")
            self.assertEqual(response.status_code, 200)
            payload = response.get_json()
        finally:
            if previous_path is None:
                os.environ.pop("NARA_NOTICE_PDF_KNOWN_ISSUES_PATH", None)
            else:
                os.environ["NARA_NOTICE_PDF_KNOWN_ISSUES_PATH"] = previous_path

        self.assertEqual(payload["status"], "phase2_mvp_complete")
        self.assertIn("default_backend", payload["test_baseline"])
        self.assertEqual(payload["sample_policy"]["notice_pdf_use"], "공고문 PDF 파이프라인 안정성 테스트용")
        self.assertEqual(payload["known_issues"][0]["type"], "mupdf_internal_syntax_warning")
        self.assertNotIn("eligible", json.dumps(payload).lower())
        self.assertNotIn("지원 가능", json.dumps(payload, ensure_ascii=False))

    def test_phase25a_basis_rule_candidate_extraction_keeps_review_status(self) -> None:
        basis = self.upload_basis_document(
            """
            Bidder qualification: a small business certificate must be submitted.
            License requirement: a landscape construction business license must be held.
            Required document: business registration and national tax payment certificate must be submitted.
            """,
            file_name="basis-rules.pdf",
        )

        response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()

        self.assertGreaterEqual(payload["candidate_count"], 3)
        self.assertTrue(all(item["status"] == "needs_review" for item in payload["candidates"]))
        self.assertTrue(all(item["citation_candidate_id"].startswith("basis:") for item in payload["candidates"]))
        rule_types = {item["rule_type"] for item in payload["candidates"]}
        self.assertIn("company_type", rule_types)
        self.assertIn("license", rule_types)
        self.assertIn("required_document", rule_types)

        list_response = self.client.get(f"/api/basis-documents/{basis['id']}/rule-candidates")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(list_response.get_json()["candidate_count"], payload["candidate_count"])
        self.assertEqual(list_response.get_json()["returned_count"], payload["candidate_count"])

        limited_response = self.client.get("/api/basis-rule-candidates?limit=1")
        self.assertEqual(limited_response.status_code, 200)
        limited_payload = limited_response.get_json()
        self.assertGreaterEqual(limited_payload["candidate_count"], payload["candidate_count"])
        self.assertEqual(limited_payload["returned_count"], 1)
        self.assertEqual(len(limited_payload["candidates"]), 1)
        self.assertEqual(limited_payload["limit"], 1)

    def test_basis_rule_candidate_reextract_preserves_reviewed_candidates(self) -> None:
        basis = self.upload_basis_document(
            """
            License requirement: forest business license must be held.
            Required document: national tax payment certificate must be submitted.
            """,
            file_name="basis-rules-reextract-preserve.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidates = extract_response.get_json()["candidates"]
        approved_target = next(item for item in candidates if item["rule_type"] == "license")
        rejected_target = next(item for item in candidates if item["id"] != approved_target["id"])

        manual_condition = "Manual reviewed forest business license requirement."
        approve_response = self.client.post(
            f"/api/basis-rule-candidates/{approved_target['id']}/approve",
            json={
                "condition_text": manual_condition,
                "required_evidence_types": ["수동 면허 확인서"],
                "related_profile_fields": ["license_summary"],
                "reviewer_name": "tester",
                "review_note": "면허 기준 승인",
            },
        )
        self.assertEqual(approve_response.status_code, 200)
        reject_response = self.client.post(
            f"/api/basis-rule-candidates/{rejected_target['id']}/reject",
            json={"reviewer_name": "tester", "review_note": "중복 후보 반려"},
        )
        self.assertEqual(reject_response.status_code, 200)

        reextract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(reextract_response.status_code, 200)
        payload = reextract_response.get_json()
        by_id = {item["id"]: item for item in payload["candidates"]}

        self.assertEqual(by_id[approved_target["id"]]["status"], "approved")
        self.assertEqual(by_id[approved_target["id"]]["condition_text"], manual_condition)
        self.assertEqual(by_id[approved_target["id"]]["required_evidence_types"], ["수동 면허 확인서"])
        self.assertEqual(by_id[approved_target["id"]]["related_profile_fields"], ["license_summary"])
        self.assertNotEqual(by_id[approved_target["id"]]["source_condition_text"], manual_condition)
        self.assertTrue(by_id[approved_target["id"]]["extraction_key"])
        self.assertEqual(by_id[rejected_target["id"]]["status"], "rejected")
        self.assertEqual(payload["new_candidate_count"], 0)
        self.assertGreaterEqual(payload["updated_candidate_count"], 2)

    def test_basis_rule_candidate_extraction_requires_completed_indexed_basis(self) -> None:
        basis = self.upload_basis_document(
            "Required document: national tax payment certificate must be submitted.",
            file_name="basis-rules-not-ready.pdf",
        )
        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE basis_documents
                SET index_status='pending'
                WHERE id=?
                """,
                (basis["id"],),
            )

        response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(response.status_code, 409)
        payload = response.get_json()
        self.assertEqual(payload["status"], "basis_not_ready")
        self.assertEqual(payload["candidate_count"], 0)

        runs_response = self.client.get("/api/operation-runs?operation_type=basis_rule_candidate_extraction")
        self.assertEqual(runs_response.status_code, 200)
        runs = runs_response.get_json()
        self.assertEqual(runs[0]["status"], "failed")
        self.assertEqual(runs[0]["error_code"], "basis_not_ready")

    def test_basis_rule_candidate_zero_reextract_preserves_existing_candidates(self) -> None:
        previous_extractor = runtime.extract_basis_rule_candidates_from_chunk
        basis = self.upload_basis_document(
            "Required document: local tax payment certificate must be submitted.",
            file_name="basis-rules-zero-reextract.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        extracted_payload = extract_response.get_json()
        original_candidate_count = extracted_payload["candidate_count"]
        candidate = extracted_payload["candidates"][0]
        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 200)

        def no_candidates(_chunk):
            return []

        try:
            runtime.extract_basis_rule_candidates_from_chunk = no_candidates
            reextract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        finally:
            runtime.extract_basis_rule_candidates_from_chunk = previous_extractor

        self.assertEqual(reextract_response.status_code, 200)
        payload = reextract_response.get_json()
        self.assertEqual(payload["status"], "no_candidates_extracted_existing_preserved")
        self.assertEqual(payload["candidate_count"], original_candidate_count)
        self.assertEqual(payload["candidates"][0]["id"], candidate["id"])
        self.assertEqual(payload["candidates"][0]["status"], "approved")

    def test_basis_rule_candidate_can_be_reviewed_updated_approved_and_rejected(self) -> None:
        basis = self.upload_basis_document(
            """
            Bidder qualification: a small business certificate must be submitted.
            License requirement: a landscape construction business license must be held.
            """,
            file_name="basis-rules-review.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidates = extract_response.get_json()["candidates"]
        first = candidates[0]

        update_response = self.client.patch(
            f"/api/basis-rule-candidates/{first['id']}",
            json={
                "condition_text": "Small business certificate must be submitted before bidding.",
                "required_evidence_types": ["중소기업확인서"],
                "related_profile_fields": ["company_size_classification"],
                "review_note": "문구 정리",
            },
        )
        self.assertEqual(update_response.status_code, 200)
        updated = update_response.get_json()
        self.assertEqual(updated["condition_text"], "Small business certificate must be submitted before bidding.")
        self.assertEqual(updated["status"], "needs_review")
        self.assertEqual(updated["review_note"], "문구 정리")

        approve_response = self.client.post(
            f"/api/basis-rule-candidates/{first['id']}/approve",
            json={"review_note": "근거 확인", "reviewer_name": "tester"},
        )
        self.assertEqual(approve_response.status_code, 200)
        approved = approve_response.get_json()
        self.assertEqual(approved["status"], "approved")
        self.assertEqual(approved["review_note"], "근거 확인")
        self.assertEqual(approved["reviewer_name"], "tester")
        self.assertTrue(approved["reviewed_at"])
        self.assertIsNotNone(approved["basis_document"])
        self.assertIsNotNone(approved["chunk"])

        approved_list_response = self.client.get("/api/basis-rule-candidates?status=approved")
        self.assertEqual(approved_list_response.status_code, 200)
        self.assertEqual(approved_list_response.get_json()["candidate_count"], 1)

        second = candidates[1]
        reject_response = self.client.post(
            f"/api/basis-rule-candidates/{second['id']}/reject",
            json={"review_note": "중복 후보"},
        )
        self.assertEqual(reject_response.status_code, 200)
        self.assertEqual(reject_response.get_json()["status"], "rejected")

    def test_basis_rule_candidate_approval_requires_citation(self) -> None:
        basis = self.upload_basis_document(
            "Required document: business registration certificate must be submitted.",
            file_name="basis-rules-no-citation.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidate = extract_response.get_json()["candidates"][0]

        update_response = self.client.patch(
            f"/api/basis-rule-candidates/{candidate['id']}",
            json={"citation_candidate_id": ""},
        )
        self.assertEqual(update_response.status_code, 200)

        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 400)
        self.assertIn("citation_candidate_id", approve_response.get_json()["detail"])

    def test_basis_rule_candidate_approval_rejects_mismatched_citation(self) -> None:
        basis = self.upload_basis_document(
            "License requirement: forest business license must be held.",
            file_name="basis-rules-bad-citation.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidate = extract_response.get_json()["candidates"][0]

        update_response = self.client.patch(
            f"/api/basis-rule-candidates/{candidate['id']}",
            json={"citation_candidate_id": f"basis:{basis['id'] + 999}:chunk:{candidate['basis_chunk_id']}"},
        )
        self.assertEqual(update_response.status_code, 200)

        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 400)
        self.assertIn("basis_document_id", approve_response.get_json()["detail"])

    def test_basis_rule_candidate_approval_rejects_unindexed_chunk(self) -> None:
        basis = self.upload_basis_document(
            "License requirement: forest business license must be held.",
            file_name="basis-rules-unindexed-citation.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidate = extract_response.get_json()["candidates"][0]

        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE basis_document_chunks
                SET vector_status='pending', vector_id=''
                WHERE id=?
                """,
                (candidate["basis_chunk_id"],),
            )

        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 400)
        self.assertIn("indexed", approve_response.get_json()["detail"])

    def test_basis_rule_candidate_approval_requires_valid_json_index(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        basis = self.upload_basis_document(
            "License requirement: forest business license must be held by the bidder.",
            file_name="rule-candidate-invalid-index.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidate = next(item for item in extract_response.get_json()["candidates"] if item["rule_type"] == "license")

        basis_pipeline.BASIS_INDEX_PATH.write_text("{not valid json", encoding="utf-8")
        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})

        self.assertEqual(approve_response.status_code, 409)
        payload = approve_response.get_json()
        self.assertEqual(payload["status"], "basis_index_unavailable")
        self.assertTrue(payload["rebuild_required"])

    def test_basis_rule_candidate_approval_requires_candidate_vector_in_json_index(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        basis = self.upload_basis_document(
            "License requirement: information communication construction license must be held.",
            file_name="rule-candidate-missing-vector.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidate = next(item for item in extract_response.get_json()["candidates"] if item["rule_type"] == "license")
        with runtime.db_conn() as conn:
            vector_id = conn.execute(
                "SELECT vector_id FROM basis_document_chunks WHERE id=?",
                (candidate["basis_chunk_id"],),
            ).fetchone()["vector_id"]
        index_payload = json.loads(basis_pipeline.BASIS_INDEX_PATH.read_text(encoding="utf-8"))
        index_payload["chunks"].pop(vector_id)
        basis_pipeline.save_basis_index(index_payload)

        approve_response = self.client.post(f"/api/basis-rule-candidates/{candidate['id']}/approve", json={})

        self.assertEqual(approve_response.status_code, 409)
        payload = approve_response.get_json()
        self.assertEqual(payload["status"], "basis_index_unavailable")
        self.assertTrue(payload["rebuild_required"])

    def test_basis_rule_candidate_reopen_clears_review_metadata(self) -> None:
        basis = self.upload_basis_document(
            "Required document: national tax payment certificate must be submitted.",
            file_name="basis-rules-reopen.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidate = extract_response.get_json()["candidates"][0]

        approve_response = self.client.post(
            f"/api/basis-rule-candidates/{candidate['id']}/approve",
            json={"reviewer_name": "tester", "review_note": "승인"},
        )
        self.assertEqual(approve_response.status_code, 200)
        approved = approve_response.get_json()
        self.assertTrue(approved["reviewed_at"])
        self.assertEqual(approved["reviewer_name"], "tester")

        reopen_response = self.client.patch(
            f"/api/basis-rule-candidates/{candidate['id']}",
            json={"status": "needs_review", "review_note": "재검토"},
        )
        self.assertEqual(reopen_response.status_code, 200)
        reopened = reopen_response.get_json()
        self.assertEqual(reopened["status"], "needs_review")
        self.assertEqual(reopened["reviewed_at"], "")
        self.assertEqual(reopened["reviewer_name"], "")

    def test_phase25b_structured_notice_requirements_are_phase3_inputs(self) -> None:
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510001",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "구조화 요구조건 공고",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "presmptPrce": "36190000",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "조경식재공사업",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/requirements/structured")
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()

        self.assertEqual(payload["contract_version"], "phase3_gap_judgment_contract_v1")
        self.assertGreaterEqual(payload["requirement_count"], 3)
        first = payload["requirements"][0]
        self.assertIn("requirement_input_id", first)
        self.assertIn("related_profile_fields", first)
        self.assertIn("comparison_strategy", first)
        self.assertNotIn("eligible", json.dumps(payload).lower())

    def test_phase25c_retrieval_evaluation_tracks_citation_coverage(self) -> None:
        basis = self.upload_basis_document(
            """
            Small business certificate submission is used for bidder qualification review.
            Landscape construction business license is used for license requirement review.
            """,
            file_name="retrieval-basis.pdf",
        )
        search_response = self.client.post(
            "/api/basis-search",
            json={"query": "small business certificate submission", "top_k": 3},
        )
        self.assertEqual(search_response.status_code, 200)
        expected_citation_id = search_response.get_json()["results"][0]["citation_candidate_id"]

        response = self.client.post(
            "/api/basis-retrieval-evaluations",
            json={
                "name": "Phase 2.5C 검색 평가",
                "queries": [
                    {
                        "query": "small business certificate submission",
                        "expected_citation_candidate_ids": [expected_citation_id],
                    },
                    {
                        "query": "landscape construction business license",
                        "expected_citation_candidate_ids": ["basis:missing:chunk:0"],
                    },
                    "landscape construction business license",
                ],
                "top_k": 3,
            },
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        self.assertEqual(payload["query_count"], 3)
        self.assertEqual(payload["citation_coverage"], 0.5)
        self.assertEqual(payload["result"]["metrics"]["result_coverage"], 1.0)
        self.assertEqual(payload["result"]["metrics"]["expected_citation_query_count"], 2)
        self.assertEqual(payload["result"]["metrics"]["expected_citation_coverage"], 0.5)
        self.assertGreater(payload["average_top_score"], 0)
        self.assertEqual(payload["result"]["index_source"], "json_basis_index")
        self.assertIn("JSON", payload["result"]["policy"])
        self.assertEqual(payload["result"]["query_results"][0]["result_count"], 1)
        self.assertEqual(payload["result"]["query_results"][0]["results"][0]["index_source"], "json_basis_index")
        self.assertTrue(payload["result"]["query_results"][0]["expected_citation_hit"])
        self.assertFalse(payload["result"]["query_results"][1]["expected_citation_hit"])
        self.assertEqual(payload["result"]["query_results"][2]["expected_citation_hit"], None)
        self.assertEqual(payload["result"]["query_results"][1]["missed_expected_citation_ids"], ["basis:missing:chunk:0"])
        self.assertIn(str(basis["id"]), payload["result"]["query_results"][0]["citation_candidate_ids"][0])

    def test_phase25d_judgment_contract_exposes_gap_first_contract(self) -> None:
        response = self.client.get("/api/judgment-contract")
        self.assertEqual(response.status_code, 200)
        payload = response.get_json()

        self.assertEqual(payload["contract_version"], "phase3_gap_judgment_contract_v1")
        self.assertIn("notice_requirement_inputs", payload["input_schema"])
        self.assertIn("matched", payload["output_schema"]["item_statuses"])
        self.assertIn("missing", payload["output_schema"]["item_statuses"])
        self.assertIn("preparation_guide", payload["output_schema"])
        self.assertNotIn("eligible", json.dumps(payload).lower())

    def test_phase3_judgment_run_citations_preparation_and_review_workflow(self) -> None:
        self.upload_basis_document(
            """
            Landscape construction business license is a citation candidate for bidder qualification review.
            Forest business license is a citation candidate for bidder qualification review.
            Small business certificate is a required document candidate for company type review.
            National tax payment certificate is a contract document candidate.
            """,
            file_name="judgment-basis.pdf",
        )
        corporation_response = self.client.post(
            "/api/corporations",
            json={
                "name": "판단 테스트 법인",
                "region": "경기도",
                "business_registration_number": "1428128387",
                "business_item": "landscape construction business license",
                "license_summary": "landscape construction business license",
                "company_size_classification": "중소기업",
            },
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510002",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "판단 실행 공고",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "presmptPrce": "36190000",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "forest business license",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["review_status"], "pending")
        self.assertGreaterEqual(payload["summary"]["matched_count"], 1)
        self.assertGreaterEqual(payload["summary"]["missing_count"], 1)
        self.assertIn("preparation_guide", payload["result"])
        self.assertEqual(payload["result"]["user_summary"]["generated_by"], "fallback")
        self.assertTrue(payload["result"]["user_summary"]["plain_summary"])
        self.assertTrue(any(link["type"] == "basis_chunk" for link in payload["result"]["user_summary"]["evidence_links"]))
        self.assertTrue(payload["input_snapshot"]["notice_requirements"])
        statuses = {item["match_status"] for item in payload["result"]["items"]}
        self.assertIn("matched", statuses)
        self.assertIn("missing", statuses)
        cited_items = [item for item in payload["result"]["items"] if item["citation_candidates"]]
        self.assertTrue(cited_items)
        self.assertNotIn("eligible", json.dumps(payload).lower())
        self.assertNotIn("지원 가능", json.dumps(payload, ensure_ascii=False))

        review_response = self.client.patch(
            f"/api/judgment-runs/{payload['id']}/review",
            json={"review_status": "needs_followup", "reviewer_note": "면허 증빙 추가 확인"},
        )
        self.assertEqual(review_response.status_code, 200)
        reviewed = review_response.get_json()
        self.assertEqual(reviewed["review_status"], "needs_followup")
        self.assertEqual(reviewed["reviewer_note"], "면허 증빙 추가 확인")

    def test_phase3_judgment_prefers_approved_rule_candidate_citations(self) -> None:
        basis = self.upload_basis_document(
            "License requirement: forest business license must be held by the bidder.",
            file_name="judgment-approved-rule-basis.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        candidates = extract_response.get_json()["candidates"]
        license_candidate = next(item for item in candidates if item["rule_type"] == "license")
        approve_response = self.client.post(
            f"/api/basis-rule-candidates/{license_candidate['id']}/approve",
            json={"reviewer_name": "tester", "review_note": "면허 기준 확인"},
        )
        self.assertEqual(approve_response.status_code, 200)
        approved = approve_response.get_json()

        corporation_response = self.client.post(
            "/api/corporations",
            json={"name": "승인 규칙 판단 법인", "region": "경기도", "license_summary": ""},
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510006",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "승인 규칙 citation 공고",
                    "lcnsLmtNm": "forest business license",
                    "bidClseDt": "2026-05-20 17:00",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        approved_items = [
            item
            for item in payload["result"]["items"]
            if approved["id"] in item.get("approved_rule_candidate_ids", [])
        ]
        self.assertTrue(approved_items)
        self.assertTrue(any(not item["basis_search_fallback_used"] for item in approved_items))
        self.assertTrue(
            any(
                citation.get("source_type") == "approved_rule_candidate"
                for item in approved_items
                for citation in item["citation_candidates"]
            )
        )
        self.assertIn("approved_rule_candidate_policy", payload["input_snapshot"])
        self.assertNotIn("eligible", json.dumps(payload).lower())

    def test_phase3_judgment_excludes_unhealthy_approved_rule_candidates(self) -> None:
        basis = self.upload_basis_document(
            "License requirement: forest business license must be held by the bidder.",
            file_name="judgment-unhealthy-approved-rule-basis.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        license_candidate = next(item for item in extract_response.get_json()["candidates"] if item["rule_type"] == "license")
        approve_response = self.client.post(
            f"/api/basis-rule-candidates/{license_candidate['id']}/approve",
            json={"reviewer_name": "tester", "review_note": "면허 기준 확인"},
        )
        self.assertEqual(approve_response.status_code, 200)

        with runtime.db_conn() as conn:
            conn.execute(
                """
                UPDATE basis_documents
                SET processing_status='failed', index_status='failed'
                WHERE id=?
                """,
                (basis["id"],),
            )

        corporation_response = self.client.post(
            "/api/corporations",
            json={"name": "비정상 기준문서 판단 법인", "region": "경기도", "license_summary": ""},
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510007",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "비정상 기준문서 citation 공고",
                    "lcnsLmtNm": "forest business license",
                    "bidClseDt": "2026-05-20 17:00",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()
        cited_candidate_ids = [
            candidate_id
            for item in payload["result"]["items"]
            for candidate_id in item.get("approved_rule_candidate_ids", [])
        ]
        self.assertNotIn(license_candidate["id"], cited_candidate_ids)

    def test_phase3_judgment_excludes_approved_rule_candidates_when_basis_index_invalid(self) -> None:
        from app.pipelines import basis_document as basis_pipeline

        basis = self.upload_basis_document(
            "License requirement: forest business license must be held by the bidder.",
            file_name="judgment-invalid-index-approved-rule-basis.pdf",
        )
        extract_response = self.client.post(f"/api/basis-documents/{basis['id']}/rule-candidates/extract")
        self.assertEqual(extract_response.status_code, 200)
        license_candidate = next(item for item in extract_response.get_json()["candidates"] if item["rule_type"] == "license")
        approve_response = self.client.post(
            f"/api/basis-rule-candidates/{license_candidate['id']}/approve",
            json={"reviewer_name": "tester", "review_note": "면허 기준 확인"},
        )
        self.assertEqual(approve_response.status_code, 200)

        basis_pipeline.BASIS_INDEX_PATH.write_text("{not valid json", encoding="utf-8")

        corporation_response = self.client.post(
            "/api/corporations",
            json={"name": "인덱스 오류 기준문서 판단 법인", "region": "경기도", "license_summary": ""},
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510008",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "인덱스 오류 citation 공고",
                    "lcnsLmtNm": "forest business license",
                    "bidClseDt": "2026-05-20 17:00",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
        )

        self.assertEqual(response.status_code, 201)
        payload = response.get_json()
        cited_candidate_ids = [
            candidate_id
            for item in payload["result"]["items"]
            for candidate_id in item.get("approved_rule_candidate_ids", [])
        ]
        self.assertNotIn(license_candidate["id"], cited_candidate_ids)
        self.assertTrue(
            any("인덱스 오류" in note for note in payload["result"]["preparation_guide"]["uncertainty_notes"])
        )

    def test_phase3_judgment_marks_low_score_citations_as_weak_not_ready(self) -> None:
        self.upload_basis_document(
            "Forest appears only as a glossary heading and is not a usable license requirement citation.",
            file_name="weak-citation-basis.pdf",
        )
        corporation_response = self.client.post(
            "/api/corporations",
            json={"name": "약한 citation 테스트 법인", "region": "서울"},
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510005",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "약한 citation 공고",
                    "lcnsLmtNm": "Forest business license specialized rare permit",
                    "bidClseDt": "2026-05-20 17:00",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"]},
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        weak_items = [item for item in payload["result"]["items"] if item["citation_status"] == "weak_candidate"]
        self.assertTrue(weak_items)
        self.assertEqual(payload["citation_coverage"], 0)
        self.assertTrue(all(not item["review_evidence_ready"] for item in weak_items))
        self.assertTrue(any(item["citation_candidates"] for item in weak_items))
        self.assertTrue(payload["result"]["preparation_guide"]["uncertainty_notes"])

    def test_phase3f_nara_collection_run_uses_injected_api_items_without_crawling(self) -> None:
        response = self.client.post(
            "/api/nara/collection-runs",
            json={
                "keyword": "조경",
                "start_date": "2026-05-01",
                "end_date": "2026-05-22",
                "business_type": "goods",
                "save": True,
                "notices": [
                    {
                        "bidNtceNo": "20260510003",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "자동 수집 테스트 공고",
                        "bidNtceDt": "2026-05-05 10:00",
                        "bidClseDt": "2026-05-20 17:00",
                    }
                ],
            },
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        self.assertEqual(payload["mode"], "injected")
        self.assertEqual(payload["searched_count"], 1)
        self.assertEqual(payload["saved_count"], 1)
        self.assertEqual(payload["criteria"]["business_type"], "goods")
        self.assertEqual(payload["result"]["business_type"], "goods")
        self.assertEqual(payload["result"]["items"][0]["business_type"], "goods")
        self.assertIn("HTML", payload["result"]["policy"])
        saved_response = self.client.get("/api/nara/saved-notices")
        self.assertEqual(saved_response.status_code, 200)
        self.assertEqual(saved_response.get_json()[0]["save_status"], "discovered")

        list_response = self.client.get("/api/nara/collection-runs?status=completed&keyword=조경")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(len(list_response.get_json()), 1)

        missing_key_response = self.client.post(
            "/api/nara/collection-runs",
            json={"keyword": "조경", "start_date": "2026-05-01", "end_date": "2026-05-22"},
        )
        self.assertEqual(missing_key_response.status_code, 201)
        self.assertEqual(missing_key_response.get_json()["status"], "not_configured")

        failed_list_response = self.client.get("/api/nara/collection-runs?status=not_configured")
        self.assertEqual(failed_list_response.status_code, 200)
        self.assertEqual(len(failed_list_response.get_json()), 1)

        partial_response = self.client.post(
            "/api/nara/collection-runs",
            json={
                "keyword": "부분실패",
                "start_date": "2026-05-01",
                "end_date": "2026-05-22",
                "save": True,
                "notices": [
                    {"bidNtceNo": "20260510007", "bidNtceOrd": "000", "bidNtceNm": "부분 성공 공고"},
                    {"bidNtceNm": "공고번호 없는 항목"},
                ],
            },
        )
        self.assertEqual(partial_response.status_code, 201)
        partial_payload = partial_response.get_json()
        self.assertEqual(partial_payload["status"], "partial_failed")
        self.assertTrue(partial_payload["result"]["retryable"])

        partial_list_response = self.client.get("/api/nara/collection-runs?status=partial_failed&keyword=Some")
        self.assertEqual(partial_list_response.status_code, 200)
        self.assertEqual(len(partial_list_response.get_json()), 1)

    def test_phase3g_judgment_marks_missing_citations_as_review_risk(self) -> None:
        corporation_response = self.client.post(
            "/api/corporations",
            json={"name": "Citation 없는 법인", "region": "서울"},
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260510004",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "Citation 없는 공고",
                    "prtcptPsblRgnNm": "부산",
                    "lcnsLmtNm": "정보통신공사업",
                    "bidClseDt": "2026-05-20 17:00",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        response = self.client.post(
            "/api/judgment-runs",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"]},
        )
        self.assertEqual(response.status_code, 201)
        payload = response.get_json()

        self.assertEqual(payload["citation_coverage"], 0)
        self.assertTrue(payload["result"]["preparation_guide"]["uncertainty_notes"])
        self.assertTrue(all(item["citation_status"] == "missing" for item in payload["result"]["items"]))
        self.assertNotIn("eligible", json.dumps(payload).lower())
        self.assertNotIn("지원 가능", json.dumps(payload, ensure_ascii=False))

    def test_ai_model_settings_masks_keys_and_defaults_to_gemini(self) -> None:
        previous_gemini_key = runtime.GEMINI_API_KEY
        previous_openai_key = runtime.OPENAI_API_KEY
        try:
            runtime.GEMINI_API_KEY = "gemini-test-key-123456"
            runtime.OPENAI_API_KEY = "openai-test-key-123456"

            response = self.client.get("/api/settings/ai-models")

            self.assertEqual(response.status_code, 200)
            payload = response.get_json()
            self.assertEqual(payload["default_provider"], "gemini")
            self.assertEqual(payload["default_model"], "gemini-2.5-flash")
            self.assertTrue(payload["providers"]["gemini"]["configured"])
            self.assertTrue(payload["providers"]["openai"]["configured"])
            self.assertNotIn("gemini-test-key-123456", json.dumps(payload))
            option_keys = {(item["provider"], item["model"]) for item in payload["options"]}
            self.assertIn(("gemini", "gemini-2.5-flash"), option_keys)
            gemini_option = next(item for item in payload["options"] if item["provider"] == "gemini")
            self.assertIn("gemini-2.5-flash", gemini_option["label"])
            self.assertNotIn("Flash-Lite", gemini_option["label"])
        finally:
            runtime.GEMINI_API_KEY = previous_gemini_key
            runtime.OPENAI_API_KEY = previous_openai_key

    def test_pdf_reader_status_exposes_engine_health_without_secrets(self) -> None:
        response = self.client.get("/api/settings/pdf-reader/status")

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertIn(payload["configured_engine"], {"auto", "opendataloader", "pymupdf"})
        self.assertIn("opendataloader", payload)
        self.assertIn("pymupdf", payload)
        self.assertIn("fallback_enabled", payload)
        self.assertEqual(payload["opendataloader"]["engine"], "opendataloader-pdf")
        self.assertEqual(payload["opendataloader"]["expected_version"], "2.4.7")
        self.assertEqual(payload["pymupdf"]["engine"], "PyMuPDF")
        self.assertNotIn("api_key", json.dumps(payload).lower())
        self.assertNotIn("service_key", json.dumps(payload).lower())

    def test_document_analysis_uses_selected_ai_model(self) -> None:
        corporation = self.create_corporation()
        project = self.create_project(corporation["id"])
        document = self.upload_document(project["id"])
        previous_summarizer = runtime.summarize_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY
        observed = {}

        def fake_summarizer(text: str, selection: dict):
            observed["selection"] = dict(selection)
            payload = {
                "document_summary": "선택 모델 테스트 요약",
                "key_dates": [],
                "requirements": ["요구사항 A"],
                "required_documents": [],
                "risks": [],
                "questions_to_check": [],
                "confidence_note": "테스트",
            }
            usage = {
                "provider": selection["provider"],
                "model": selection["model"],
                "input_chars": len(text),
            }
            return payload, runtime.render_summary_markdown(payload), usage

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.summarize_with_ai = fake_summarizer

            response = self.client.post(
                f"/api/documents/{document['id']}/analyze",
                json={
                    "model_provider": "gemini",
                    "model_name": "gemini-2.5-flash",
                },
            )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(observed["selection"]["provider"], "gemini")
            self.assertEqual(observed["selection"]["model"], "gemini-2.5-flash")

            latest = self.client.get(f"/api/analyses/latest/by-document/{document['id']}")
            self.assertEqual(latest.status_code, 200)
            analysis = latest.get_json()
            self.assertEqual(analysis["model_provider"], "gemini")
            self.assertEqual(analysis["model_name"], "gemini-2.5-flash")
        finally:
            runtime.summarize_with_ai = previous_summarizer
            runtime.GEMINI_API_KEY = previous_gemini_key

    def test_unknown_evidence_uses_ai_classification_fallback_as_review_candidates(self) -> None:
        previous_classifier = runtime.classify_unknown_evidence_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY
        observed = {}

        def fake_classifier(text: str, file_name: str, provider=None, model=None):
            observed["file_name"] = file_name
            observed["text"] = text
            return runtime.EvidenceExtractionResult(
                document_type="credit_rating_certificate",
                classification_confidence=0.73,
                classification_status="ai_suggested",
                candidates=[
                    runtime.EvidenceFieldCandidate(
                        field_key="certifications_json",
                        field_label="인증/확인서",
                        extracted_value='["기업신용평가등급 B+"]',
                        confidence=0.7,
                        source_text="기업신용평가등급 B+",
                    )
                ],
                warnings=["AI 분류 제안 테스트"],
            )

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.classify_unknown_evidence_with_ai = fake_classifier
            response = self.client.post(
                "/api/corporation-evidence-documents",
                data={
                    "document_type": "auto",
                    "management_group_name": "AI그룹",
                    "memo": "알 수 없는 증빙 AI 분류 테스트",
                    "file": (
                        io.BytesIO(
                            make_docx_bytes(
                                """
                                협력업체 내부 검토 메모
                                업체명 : 테스트 법인
                                참고 문구 : 담당자가 별도 확인해야 하는 임의 자료
                                비고 : 규칙 기반 분류 대상이 아닌 문서
                                """
                            )
                        ),
                        "unknown-evidence.docx",
                    ),
                },
                content_type="multipart/form-data",
            )

            self.assertEqual(response.status_code, 201)
            payload = response.get_json()
            self.assertEqual(payload["document_type"], "credit_rating_certificate")
            self.assertEqual(payload["classification_status"], "ai_suggested")
            self.assertEqual(payload["review_status"], "pending")
            self.assertEqual(payload["candidates"][0]["field_key"], "certifications_json")
            self.assertEqual(payload["candidates"][0]["status"], "pending")
            self.assertEqual(observed["file_name"], "unknown-evidence.docx")
        finally:
            runtime.classify_unknown_evidence_with_ai = previous_classifier
            runtime.GEMINI_API_KEY = previous_gemini_key

    def test_business_registration_kind_fields_use_ai_cleanup_when_configured(self) -> None:
        previous_generator = runtime.generate_json_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY
        observed = {}

        def fake_json_generator(prompt: str, selection: dict):
            observed["prompt"] = prompt
            observed["selection"] = dict(selection)
            return (
                {
                    "business_type": ["건설업", "도소매", "도매 및 소매업"],
                    "business_item": [
                        "전기공사",
                        "신재생에너지설비설치전문기업",
                        "일반건축공사업",
                        "전문소방시설공사",
                        "안전시설물설치",
                        "정보통신공사업",
                        "토목공사업",
                        "전기",
                        "소방",
                        "통신자재",
                        "컴퓨터 관련 주변기기",
                        "산업안전용품",
                    ],
                    "business_category": "업태: 건설업, 도소매, 도매 및 소매업 / 종목: 전기공사, 신재생에너지설비설치전문기업, 일반건축공사업, 전문소방시설공사, 안전시설물설치, 정보통신공사업, 토목공사업, 전기, 소방, 통신자재, 컴퓨터 관련 주변기기, 산업안전용품",
                    "warnings": [],
                },
                {"provider": "gemini", "model": selection["model"]},
            )

        ocr_text = """
        사업자등록증
        등록번호 : 142-81-28387
        법인명(단체명) : 주식회사 온세이엔씨
        대표자 : 안영식
        사업의 종류:
        업태
        건설업
        종목
        전기공사,신재생에너지설비설치전문기
        업
        건설업
        일반건축공사업
        도매 및
        소매업
        컴퓨터 관련 주변기기
        """

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.generate_json_with_ai = fake_json_generator
            result = runtime.analyze_corporation_evidence_text(
                ocr_text,
                "사업자등록증.png",
                "auto",
            )

            fields = {candidate.field_key: candidate.extracted_value for candidate in result.candidates}
            self.assertEqual(observed["selection"]["provider"], "gemini")
            self.assertIn("사업의 종류", observed["prompt"])
            self.assertIn("신재생에너지설비설치전문기업", fields["business_item"])
            self.assertIn("컴퓨터 관련 주변기기", fields["business_item"])
            self.assertIn("업태: 건설업", fields["business_category"])
            self.assertTrue(any("AI 업태/종목 정리" in warning for warning in result.warnings))
        finally:
            runtime.generate_json_with_ai = previous_generator
            runtime.GEMINI_API_KEY = previous_gemini_key

    def test_business_registration_ai_cleanup_sanitizes_sloppy_label_output(self) -> None:
        previous_generator = runtime.generate_json_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY

        def fake_json_generator(prompt: str, selection: dict):
            return (
                {
                    "business_type": ["사업의 종류", "업태 건설업", "도매 및\n소매업"],
                    "business_item": [
                        "종목 전기공사,신재생에너지설비설치전문기\n업",
                        "건설업 정보통신공사업, 토목공사업",
                    ],
                    "business_category": "사업의 종류: 업태 건설업 종목 전기공사",
                    "warnings": [],
                },
                {"provider": "gemini", "model": selection["model"]},
            )

        ocr_text = """
        사업자등록증
        등록번호 : 142-81-28387
        법인명(단체명) : 주식회사 온세이엔씨
        대표자 : 안영식
        사업의 종류:
        업태
        건설업
        종목
        전기공사,신재생에너지설비설치전문기
        업
        건설업
        정보통신공사업, 토목공사업
        도매 및
        소매업
        컴퓨터 관련 주변기기
        """

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.generate_json_with_ai = fake_json_generator
            result = runtime.analyze_corporation_evidence_text(
                ocr_text,
                "사업자등록증.png",
                "auto",
            )

            fields = {candidate.field_key: candidate.extracted_value for candidate in result.candidates}
            self.assertEqual(fields["business_type"], "건설업, 도매 및 소매업")
            self.assertIn("전기공사", fields["business_item"])
            self.assertIn("신재생에너지설비설치전문기업", fields["business_item"])
            self.assertIn("정보통신공사업", fields["business_item"])
            self.assertNotIn("사업의 종류", fields["business_category"])
            self.assertNotIn("종목", fields["business_item"])
        finally:
            runtime.generate_json_with_ai = previous_generator
            runtime.GEMINI_API_KEY = previous_gemini_key

    def test_evidence_upload_extracts_candidates_and_creates_corporation_after_approval(self) -> None:
        evidence = self.upload_business_registration_evidence()

        self.assertEqual(evidence["document_type"], "business_registration_certificate")
        self.assertEqual(evidence["classification_status"], "classified")
        field_values = {candidate["field_key"]: candidate["extracted_value"] for candidate in evidence["candidates"]}
        self.assertEqual(field_values["business_registration_number"], "142-81-28387")
        self.assertEqual(field_values["name"], "주식회사 온세이엔씨")
        self.assertEqual(field_values["representative_name"], "안영식")

        response = self.client.post(f"/api/corporation-evidence-documents/{evidence['id']}/approve", json={})
        self.assertEqual(response.status_code, 200)

        payload = response.get_json()
        corporation = payload["corporation"]
        self.assertEqual(corporation["name"], "주식회사 온세이엔씨")
        self.assertEqual(corporation["business_registration_number"], "142-81-28387")
        self.assertEqual(corporation["representative_name"], "안영식")
        self.assertEqual(corporation["evidence_verification_status"], "evidence_reviewed")
        self.assertIn("name", payload["applied_fields"])

    def test_evidence_approval_updates_existing_corporation(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_business_registration_evidence(corporation["id"])

        response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/approve",
            json={"field_values": {"name": corporation["name"]}},
        )
        self.assertEqual(response.status_code, 200)

        updated = response.get_json()["corporation"]
        self.assertEqual(updated["id"], corporation["id"])
        self.assertEqual(updated["business_registration_number"], "142-81-28387")
        self.assertEqual(updated["representative_name"], "안영식")

        list_response = self.client.get(f"/api/corporations/{corporation['id']}/evidence-documents")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(len(list_response.get_json()), 1)

    def test_empty_candidate_ids_does_not_approve_all_candidates_when_field_values_exist(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_business_registration_evidence(corporation["id"])

        response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/approve",
            json={
                "candidate_ids": [],
                "field_values": {"name": corporation["name"]},
            },
        )

        self.assertEqual(response.status_code, 200)
        updated = response.get_json()["corporation"]
        self.assertEqual(updated["name"], corporation["name"])
        self.assertEqual(updated["business_registration_number"], "")

        detail = self.client.get(f"/api/corporation-evidence-documents/{evidence['id']}")
        self.assertEqual(detail.status_code, 200)
        self.assertTrue(all(candidate["status"] == "pending" for candidate in detail.get_json()["candidates"]))

    def test_selected_evidence_candidates_only_apply_selected_and_edited_values(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_business_registration_evidence(corporation["id"])
        candidates = {candidate["field_key"]: candidate for candidate in evidence["candidates"]}

        response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/approve",
            json={
                "candidate_ids": [
                    candidates["business_registration_number"]["id"],
                    candidates["representative_name"]["id"],
                ],
                "field_values": {"representative_name": "수정대표"},
            },
        )

        self.assertEqual(response.status_code, 200)
        updated = response.get_json()["corporation"]
        self.assertEqual(updated["name"], corporation["name"])
        self.assertEqual(updated["business_registration_number"], "142-81-28387")
        self.assertEqual(updated["representative_name"], "수정대표")

        detail = self.client.get(f"/api/corporation-evidence-documents/{evidence['id']}")
        self.assertEqual(detail.status_code, 200)
        statuses = {candidate["field_key"]: candidate["status"] for candidate in detail.get_json()["candidates"]}
        self.assertEqual(statuses["business_registration_number"], "approved")
        self.assertEqual(statuses["representative_name"], "approved")
        self.assertEqual(statuses["name"], "pending")

    def test_core_evidence_approval_merges_certifications_and_preference_tags(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_small_business_evidence(corporation["id"])

        self.assertEqual(evidence["document_type"], "small_business_confirmation")
        fields = {candidate["field_key"]: candidate["extracted_value"] for candidate in evidence["candidates"]}
        self.assertEqual(fields["company_size_classification"], "소기업")

        response = self.client.post(f"/api/corporation-evidence-documents/{evidence['id']}/approve", json={})
        self.assertEqual(response.status_code, 200)

        updated = response.get_json()["corporation"]
        certifications = json.loads(updated["certifications_json"])
        preference_tags = json.loads(updated["preference_tags_json"])
        self.assertIn("ISO9001", certifications)
        self.assertIn("중소기업확인서", certifications)
        self.assertIn("중소기업", preference_tags)
        self.assertEqual(updated["company_size_classification"], "소기업")

    def test_corporation_readiness_reports_missing_items_and_evidence_progress(self) -> None:
        corporation = self.create_corporation()

        initial_response = self.client.get(f"/api/corporations/{corporation['id']}/readiness")
        self.assertEqual(initial_response.status_code, 200)
        initial = initial_response.get_json()
        self.assertEqual(initial["status"], "needs_evidence")
        self.assertIn("사업자등록번호", initial["missing_items"])
        self.assertEqual(initial["evidence_count"], 0)

        evidence = self.upload_business_registration_evidence(corporation["id"])
        approve_response = self.client.post(f"/api/corporation-evidence-documents/{evidence['id']}/approve", json={})
        self.assertEqual(approve_response.status_code, 200)

        list_response = self.client.get("/api/corporations/readiness")
        self.assertEqual(list_response.status_code, 200)
        readiness = list_response.get_json()[0]
        self.assertEqual(readiness["corporation_id"], corporation["id"])
        self.assertGreater(readiness["score"], initial["score"])
        self.assertGreaterEqual(readiness["approved_candidate_count"], 1)
        self.assertNotIn("사업자등록번호", readiness["missing_items"])

    def test_evidence_documents_can_be_listed_updated_and_reprocessed(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_small_business_evidence(corporation["id"])

        list_response = self.client.get("/api/corporation-evidence-documents")
        self.assertEqual(list_response.status_code, 200)
        rows = list_response.get_json()
        self.assertEqual(rows[0]["id"], evidence["id"])
        self.assertEqual(rows[0]["corporation_name"], corporation["name"])
        self.assertGreater(rows[0]["candidate_count"], 0)
        self.assertEqual(rows[0]["pending_candidate_count"], rows[0]["candidate_count"])
        self.assertEqual(rows[0]["candidates"], [])

        patch_response = self.client.patch(
            f"/api/corporation-evidence-documents/{evidence['id']}",
            json={
                "memo": "메타데이터 수정 테스트",
                "review_status": "needs_review",
                "document_type": "small_business_confirmation",
            },
        )
        self.assertEqual(patch_response.status_code, 200)
        patched = patch_response.get_json()
        self.assertEqual(patched["memo"], "메타데이터 수정 테스트")
        self.assertEqual(patched["review_status"], "needs_review")

        reprocess_response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/reprocess",
            json={"document_type": "small_business_confirmation"},
        )
        self.assertEqual(reprocess_response.status_code, 200)
        reprocessed = reprocess_response.get_json()
        self.assertEqual(reprocessed["document_type"], "small_business_confirmation")
        self.assertEqual(reprocessed["review_status"], "pending")
        self.assertTrue(reprocessed["candidates"])
        self.assertIn("중소기업확인서", reprocessed["extracted_text_preview"])

    def test_evidence_reprocess_preserves_approved_candidates_without_duplicate_pending_values(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_business_registration_evidence(corporation["id"])
        candidates = {candidate["field_key"]: candidate for candidate in evidence["candidates"]}

        approve_response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/approve",
            json={"candidate_ids": [candidates["business_registration_number"]["id"]]},
        )
        self.assertEqual(approve_response.status_code, 200)

        reprocess_response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/reprocess",
            json={"document_type": "business_registration_certificate"},
        )
        self.assertEqual(reprocess_response.status_code, 200)
        reprocessed = reprocess_response.get_json()
        business_number_candidates = [
            candidate
            for candidate in reprocessed["candidates"]
            if candidate["field_key"] == "business_registration_number"
            and candidate["extracted_value"] == "142-81-28387"
        ]

        self.assertEqual(len(business_number_candidates), 1)
        self.assertEqual(business_number_candidates[0]["status"], "approved")
        self.assertTrue(any(candidate["status"] == "pending" for candidate in reprocessed["candidates"]))

    def test_evidence_corrected_text_can_be_reanalyzed(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_business_registration_evidence(corporation["id"])

        corrected_text = """
        기업신용평가등급확인서
        업체명 : 테스트 법인
        사업자등록번호 : 142-81-28387
        평가등급 : A-
        유효기간 : 2026.01.01 ~ 2026.12.31
        """
        response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/reanalyze-text",
            json={
                "document_type": "auto",
                "extracted_text": corrected_text,
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.get_json()
        self.assertEqual(payload["document_type"], "credit_rating_certificate")
        self.assertEqual(payload["ocr_status"], "corrected")
        self.assertIn("A-", payload["extracted_text"])
        fields = {candidate["field_key"]: candidate["extracted_value"] for candidate in payload["candidates"]}
        self.assertIn("기업신용평가등급 A-", fields["certifications_json"])

    def test_evidence_corrected_text_requires_non_empty_text(self) -> None:
        corporation = self.create_corporation()
        evidence = self.upload_business_registration_evidence(corporation["id"])

        response = self.client.post(
            f"/api/corporation-evidence-documents/{evidence['id']}/reanalyze-text",
            json={"document_type": "auto", "extracted_text": "   "},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("extracted_text", response.get_json()["detail"])

    def test_business_registration_duplicate_is_blocked_inside_same_management_group(self) -> None:
        first = self.client.post(
            "/api/corporations",
            json={
                "name": "기존 법인",
                "management_group_name": "A그룹",
                "business_registration_number": "142-81-28387",
            },
        )
        self.assertEqual(first.status_code, 201)

        duplicate = self.client.post(
            "/api/corporations",
            json={
                "name": "같은 그룹 중복 법인",
                "management_group_name": "A그룹",
                "business_registration_number": "1428128387",
            },
        )

        self.assertEqual(duplicate.status_code, 409)
        self.assertIn("동일한 사업자등록번호", duplicate.get_json()["detail"])

    def test_business_registration_duplicate_is_allowed_across_management_groups_with_warning(self) -> None:
        first = self.client.post(
            "/api/corporations",
            json={
                "name": "기존 법인",
                "management_group_name": "A그룹",
                "business_registration_number": "142-81-28387",
            },
        )
        self.assertEqual(first.status_code, 201)

        second = self.client.post(
            "/api/corporations",
            json={
                "name": "다른 그룹 법인",
                "management_group_name": "B그룹",
                "business_registration_number": "1428128387",
            },
        )

        self.assertEqual(second.status_code, 201)
        payload = second.get_json()
        self.assertEqual(payload["business_registration_number"], "142-81-28387")
        self.assertIn("warnings", payload)
        self.assertEqual(payload["duplicate_corporations"][0]["management_group_name"], "A그룹")

    def test_evidence_created_corporation_uses_management_group_duplicate_policy(self) -> None:
        existing = self.client.post(
            "/api/corporations",
            json={
                "name": "기존 법인",
                "management_group_name": "A그룹",
                "business_registration_number": "142-81-28387",
            },
        )
        self.assertEqual(existing.status_code, 201)

        blocked_evidence = self.upload_business_registration_evidence(management_group_name="A그룹")
        blocked = self.client.post(f"/api/corporation-evidence-documents/{blocked_evidence['id']}/approve", json={})
        self.assertEqual(blocked.status_code, 409)

        allowed_evidence = self.upload_business_registration_evidence(management_group_name="B그룹")
        allowed = self.client.post(f"/api/corporation-evidence-documents/{allowed_evidence['id']}/approve", json={})
        self.assertEqual(allowed.status_code, 200)
        payload = allowed.get_json()
        self.assertEqual(payload["corporation"]["management_group_name"], "B그룹")
        self.assertTrue(payload["warnings"])

    def test_nara_settings_status_does_not_expose_full_key_when_unconfigured(self) -> None:
        response = self.client.get("/api/settings/integrations/nara/status")
        self.assertEqual(response.status_code, 200)

        payload = response.get_json()
        self.assertFalse(payload["configured"])
        self.assertEqual(payload["masked_key"], "")
        self.assertEqual(payload["last_test_status"], "not_run")
        self.assertNotIn("service_key", payload)

    def test_nara_settings_status_returns_latest_saved_test_result(self) -> None:
        runtime.save_integration_test_result("nara", "ok", 200, "00", "NORMAL SERVICE", 7)

        response = self.client.get("/api/settings/integrations/nara/status")
        self.assertEqual(response.status_code, 200)

        payload = response.get_json()
        self.assertEqual(payload["last_test_status"], "ok")
        self.assertEqual(payload["last_test_http_status"], 200)
        self.assertEqual(payload["last_test_result_code"], "00")
        self.assertEqual(payload["last_test_total_count"], 7)
        self.assertTrue(payload["last_tested_at"])

    def test_nara_search_requires_api_key(self) -> None:
        response = self.client.get("/api/nara/notices/search")
        self.assertEqual(response.status_code, 400)
        self.assertIn("NARA_API_SERVICE_KEY", response.get_json()["detail"])

    def test_nara_search_uses_selected_business_type_operation(self) -> None:
        original_key = runtime.NARA_API_SERVICE_KEY
        original_request = runtime.request_nara_operation
        calls: list[str] = []

        def fake_request_nara_operation(operation: str, params: dict[str, str]) -> dict:
            calls.append(operation)
            business_label = {
                "getBidPblancListInfoServcPPSSrch": "용역",
                "getBidPblancListInfoThngPPSSrch": "물품",
            }[operation]
            return {
                "http_status": 200,
                "header": {"resultCode": "00", "resultMsg": "NORMAL SERVICE"},
                "total_count": 1,
                "items": [
                    {
                        "bidNtceNo": f"20260614-{business_label}",
                        "bidNtceOrd": "000",
                        "bidNtceNm": f"{business_label} 테스트 공고",
                        "bidNtceDt": "202606141000",
                    }
                ],
            }

        runtime.NARA_API_SERVICE_KEY = "TEST_NARA_KEY"
        runtime.request_nara_operation = fake_request_nara_operation
        try:
            service_response = self.client.get("/api/nara/notices/search", query_string={"business_type": "service"})
            self.assertEqual(service_response.status_code, 200)
            service_payload = service_response.get_json()
            self.assertEqual(calls, ["getBidPblancListInfoServcPPSSrch"])
            self.assertEqual(service_payload["items"][0]["business_type"], "service")
            self.assertEqual(service_payload["items"][0]["business_type_label"], "용역")

            calls.clear()
            goods_response = self.client.get("/api/nara/notices/search", query_string={"business_type": "goods"})
            self.assertEqual(goods_response.status_code, 200)
            goods_payload = goods_response.get_json()
            self.assertEqual(calls, ["getBidPblancListInfoThngPPSSrch"])
            self.assertEqual(goods_payload["items"][0]["business_type"], "goods")
            self.assertEqual(goods_payload["items"][0]["business_type_label"], "물품")
        finally:
            runtime.NARA_API_SERVICE_KEY = original_key
            runtime.request_nara_operation = original_request

    def test_nara_search_all_merges_business_types_and_deduplicates_notices(self) -> None:
        original_key = runtime.NARA_API_SERVICE_KEY
        original_request = runtime.request_nara_operation
        calls: list[str] = []

        item_map = {
            "getBidPblancListInfoCnstwkPPSSrch": [
                {"bidNtceNo": "202606140001", "bidNtceOrd": "000", "bidNtceNm": "공사 공고", "bidNtceDt": "202606141000"}
            ],
            "getBidPblancListInfoServcPPSSrch": [
                {"bidNtceNo": "202606140002", "bidNtceOrd": "000", "bidNtceNm": "용역 공고", "bidNtceDt": "202606141200"}
            ],
            "getBidPblancListInfoThngPPSSrch": [
                {"bidNtceNo": "202606140002", "bidNtceOrd": "000", "bidNtceNm": "중복 물품 공고", "bidNtceDt": "202606141300"}
            ],
            "getBidPblancListInfoEtcPPSSrch": [
                {"bidNtceNo": "202606140004", "bidNtceOrd": "000", "bidNtceNm": "기타 공고", "bidNtceDt": "202606140900"}
            ],
        }

        def fake_request_nara_operation(operation: str, params: dict[str, str]) -> dict:
            calls.append(operation)
            return {
                "http_status": 200,
                "header": {"resultCode": "00", "resultMsg": "NORMAL SERVICE"},
                "total_count": len(item_map.get(operation, [])),
                "items": item_map.get(operation, []),
            }

        runtime.NARA_API_SERVICE_KEY = "TEST_NARA_KEY"
        runtime.request_nara_operation = fake_request_nara_operation
        try:
            response = self.client.get("/api/nara/notices/search", query_string={"business_type": "all", "page_size": 20})
            self.assertEqual(response.status_code, 200)
            payload = response.get_json()

            self.assertEqual(
                calls,
                [
                    "getBidPblancListInfoCnstwkPPSSrch",
                    "getBidPblancListInfoServcPPSSrch",
                    "getBidPblancListInfoThngPPSSrch",
                    "getBidPblancListInfoEtcPPSSrch",
                ],
            )
            self.assertEqual(payload["business_type"], "all")
            self.assertEqual(payload["queried_business_types"], ["construction", "service", "goods", "etc"])
            self.assertEqual(payload["total_count"], 4)
            self.assertEqual(len(payload["items"]), 3)
            self.assertEqual(len({(item["bid_ntce_no"], item["bid_ntce_ord"]) for item in payload["items"]}), 3)
            self.assertEqual({item["business_type"] for item in payload["items"]}, {"construction", "service", "etc"})
        finally:
            runtime.NARA_API_SERVICE_KEY = original_key
            runtime.request_nara_operation = original_request

    def test_nara_search_all_uses_global_merged_pagination_without_duplicates(self) -> None:
        original_key = runtime.NARA_API_SERVICE_KEY
        original_request = runtime.request_nara_operation

        item_map = {
            "getBidPblancListInfoCnstwkPPSSrch": [
                {"bidNtceNo": "C-1", "bidNtceOrd": "000", "bidNtceNm": "공사 1", "bidNtceDt": "202606141400"},
                {"bidNtceNo": "C-2", "bidNtceOrd": "000", "bidNtceNm": "공사 2", "bidNtceDt": "202606141000"},
            ],
            "getBidPblancListInfoServcPPSSrch": [
                {"bidNtceNo": "S-1", "bidNtceOrd": "000", "bidNtceNm": "용역 1", "bidNtceDt": "202606141300"},
                {"bidNtceNo": "S-2", "bidNtceOrd": "000", "bidNtceNm": "용역 2", "bidNtceDt": "202606140900"},
            ],
            "getBidPblancListInfoThngPPSSrch": [
                {"bidNtceNo": "G-1", "bidNtceOrd": "000", "bidNtceNm": "물품 1", "bidNtceDt": "202606141200"},
            ],
            "getBidPblancListInfoEtcPPSSrch": [
                {"bidNtceNo": "E-1", "bidNtceOrd": "000", "bidNtceNm": "기타 1", "bidNtceDt": "202606141100"},
            ],
        }

        def fake_request_nara_operation(operation: str, params: dict[str, str]) -> dict:
            rows = item_map[operation]
            page_no = int(params.get("pageNo", "1"))
            page_size = int(params.get("numOfRows", "20"))
            start = (page_no - 1) * page_size
            return {
                "http_status": 200,
                "header": {"resultCode": "00", "resultMsg": "NORMAL SERVICE"},
                "total_count": len(rows),
                "items": rows[start : start + page_size],
            }

        runtime.NARA_API_SERVICE_KEY = "TEST_NARA_KEY"
        runtime.request_nara_operation = fake_request_nara_operation
        try:
            first = self.client.get("/api/nara/notices/search", query_string={"business_type": "all", "page_size": 2, "page_no": 1})
            second = self.client.get("/api/nara/notices/search", query_string={"business_type": "all", "page_size": 2, "page_no": 2})

            self.assertEqual(first.status_code, 200)
            self.assertEqual(second.status_code, 200)
            first_payload = first.get_json()
            second_payload = second.get_json()
            first_keys = [(item["bid_ntce_no"], item["bid_ntce_ord"]) for item in first_payload["items"]]
            second_keys = [(item["bid_ntce_no"], item["bid_ntce_ord"]) for item in second_payload["items"]]

            self.assertEqual(first_payload["pagination_mode"], "merged_all")
            self.assertTrue(first_payload["total_count_is_estimated"])
            self.assertTrue(first_payload["has_next_page"])
            self.assertEqual(first_keys, [("C-1", "000"), ("S-1", "000")])
            self.assertEqual(second_keys, [("G-1", "000"), ("E-1", "000")])
            self.assertFalse(set(first_keys) & set(second_keys))
        finally:
            runtime.NARA_API_SERVICE_KEY = original_key
            runtime.request_nara_operation = original_request

    def test_nara_search_all_returns_partial_results_when_one_business_type_fails(self) -> None:
        original_key = runtime.NARA_API_SERVICE_KEY
        original_request = runtime.request_nara_operation

        def fake_request_nara_operation(operation: str, params: dict[str, str]) -> dict:
            if operation == "getBidPblancListInfoThngPPSSrch":
                raise RuntimeError("goods endpoint timeout")
            return {
                "http_status": 200,
                "header": {"resultCode": "00", "resultMsg": "NORMAL SERVICE"},
                "total_count": 1,
                "items": [
                    {
                        "bidNtceNo": f"{operation}-1",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "부분 성공 공고",
                        "bidNtceDt": "202606141000",
                    }
                ],
            }

        runtime.NARA_API_SERVICE_KEY = "TEST_NARA_KEY"
        runtime.request_nara_operation = fake_request_nara_operation
        try:
            response = self.client.get("/api/nara/notices/search", query_string={"business_type": "all"})
            self.assertEqual(response.status_code, 200)
            payload = response.get_json()

            self.assertEqual(payload["result_code"], "partial_failed")
            self.assertEqual(payload["result_msg"], "일부 업무유형 조회 실패")
            self.assertEqual(payload["partial_errors"][0]["business_type"], "goods")
            self.assertTrue(payload["items"])
        finally:
            runtime.NARA_API_SERVICE_KEY = original_key
            runtime.request_nara_operation = original_request

    def test_nara_search_all_returns_502_when_all_business_types_fail(self) -> None:
        original_key = runtime.NARA_API_SERVICE_KEY
        original_request = runtime.request_nara_operation

        def fake_request_nara_operation(operation: str, params: dict[str, str]) -> dict:
            raise RuntimeError(f"{operation} failed")

        runtime.NARA_API_SERVICE_KEY = "TEST_NARA_KEY"
        runtime.request_nara_operation = fake_request_nara_operation
        try:
            response = self.client.get("/api/nara/notices/search", query_string={"business_type": "all"})
            self.assertEqual(response.status_code, 502)
            payload = response.get_json()
            self.assertEqual(len(payload["partial_errors"]), 4)
        finally:
            runtime.NARA_API_SERVICE_KEY = original_key
            runtime.request_nara_operation = original_request

    def test_nara_notice_business_type_backfill_uses_raw_json_when_default_is_stale(self) -> None:
        now = runtime.now_iso()
        with runtime.db_conn() as conn:
            conn.execute(
                """
                INSERT INTO nara_notices (
                  bid_ntce_no, bid_ntce_ord, business_type, bid_ntce_nm,
                  raw_json, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    "BACKFILL-SERVICE",
                    "000",
                    "construction",
                    "과거 용역 공고",
                    json.dumps({"bsnsDivNm": "용역", "bidNtceNm": "과거 용역 공고"}, ensure_ascii=False),
                    now,
                    now,
                ),
            )
            updated_count = runtime.backfill_nara_notice_business_type(conn)
            row = conn.execute("SELECT business_type FROM nara_notices WHERE bid_ntce_no=?", ("BACKFILL-SERVICE",)).fetchone()

        self.assertEqual(updated_count, 1)
        self.assertEqual(row["business_type"], "service")

    def test_nara_save_analysis_uses_business_type_specific_detail_operations(self) -> None:
        original_key = runtime.NARA_API_SERVICE_KEY
        original_request = runtime.request_nara_operation
        calls: list[str] = []

        def fake_request_nara_operation(operation: str, params: dict[str, str]) -> dict:
            calls.append(operation)
            return {
                "http_status": 200,
                "header": {"resultCode": "00", "resultMsg": "NORMAL SERVICE"},
                "total_count": 1 if operation.endswith(("Servc", "Thng")) else 0,
                "items": [
                    {
                        "bidNtceNo": params.get("bidNtceNo", ""),
                        "bidNtceOrd": params.get("bidNtceOrd", "000"),
                        "bidNtceNm": "업무유형 상세 공고",
                        "bidNtceDt": "202606141000",
                    }
                ]
                if operation.endswith(("Servc", "Thng"))
                else [],
            }

        runtime.NARA_API_SERVICE_KEY = "TEST_NARA_KEY"
        runtime.request_nara_operation = fake_request_nara_operation
        try:
            service_payload, service_code = runtime.save_and_analyze_nara_notice_item(
                {
                    "bidNtceNo": "202606140101",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "용역 저장 공고",
                    "business_type": "service",
                }
            )
            self.assertEqual(service_code, 200)
            self.assertEqual(service_payload["notice"]["business_type"], "service")
            self.assertIn("getBidPblancListInfoServc", calls)
            self.assertIn("getBidPblancListInfoServcBsisAmount", calls)
            self.assertNotIn("getBidPblancListInfoCnstwk", calls)

            calls.clear()
            goods_payload, goods_code = runtime.save_and_analyze_nara_notice_item(
                {
                    "bidNtceNo": "202606140102",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "물품 저장 공고",
                    "business_type": "goods",
                }
            )
            self.assertEqual(goods_code, 200)
            self.assertEqual(goods_payload["notice"]["business_type"], "goods")
            self.assertIn("getBidPblancListInfoThng", calls)
            self.assertIn("getBidPblancListInfoThngBsisAmount", calls)
        finally:
            runtime.NARA_API_SERVICE_KEY = original_key
            runtime.request_nara_operation = original_request

    def test_attachment_preview_rejects_missing_url(self) -> None:
        response = self.client.get("/api/nara/attachments/preview")
        self.assertEqual(response.status_code, 400)
        self.assertIn("url is required", response.get_json()["detail"])

    def test_attachment_preview_rejects_private_network_url(self) -> None:
        response = self.client.get(
            "/api/nara/attachments/preview",
            query_string={"url": "http://127.0.0.1/private.pdf", "name": "private.pdf"},
        )
        self.assertEqual(response.status_code, 400)

    def test_nara_notice_attachment_download_rejects_private_network_url(self) -> None:
        response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260509991",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "내부망 URL 차단 공고",
                    "ntceSpecFileNm1": "private.pdf",
                    "ntceSpecDocUrl1": "http://127.0.0.1/private.pdf",
                }
            },
        )
        self.assertEqual(response.status_code, 202)

        saved = self.wait_for_saved_nara_notice(response.get_json()["notice"]["id"])
        self.assertEqual(saved["download_status"], "partial_failed")
        self.assertEqual(saved["analysis_status"], "partial_failed")
        self.assertEqual(saved["attachments"][0]["download_status"], "failed")
        self.assertIn("Unsafe attachment URL", saved["attachments"][0]["error_message"])

    def test_nara_reanalysis_partial_failure_preserves_existing_requirements(self) -> None:
        original_request_binary = runtime.request_binary

        def fake_request_binary(url, timeout=30, max_bytes=50 * 1024 * 1024):
            if url == "https://example.com/valid.pdf":
                return (
                    200,
                    {"content-type": "application/pdf"},
                    make_pdf_bytes("제출서류: 국세 납세증명서. 면허: 조경식재공사업."),
                )
            return original_request_binary(url, timeout=timeout, max_bytes=max_bytes)

        runtime.request_binary = fake_request_binary
        try:
            initial_response = self.client.post(
                "/api/nara/notices/save-and-analyze",
                json={
                    "notice": {
                        "bidNtceNo": "20260509992",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "재분석 보존 공고",
                        "prtcptPsblRgnNm": "경기도",
                        "lcnsLmtNm": "조경식재공사업",
                        "ntceSpecFileNm1": "valid.pdf",
                        "ntceSpecDocUrl1": "https://example.com/valid.pdf",
                    }
                },
            )
            self.assertEqual(initial_response.status_code, 202)
            notice_id = initial_response.get_json()["notice"]["id"]
            initial_saved = self.wait_for_saved_nara_notice(notice_id)
            self.assertEqual(initial_saved["analysis_status"], "completed")
            self.assertTrue(initial_saved["attachments"])
            self.assertTrue(all(item["download_status"] == "completed" for item in initial_saved["attachments"]))
            initial_attachment_ids = {item["id"] for item in initial_saved["attachments"]}
            initial_requirements = self.client.get(f"/api/nara/saved-notices/{notice_id}/requirements").get_json()
            self.assertGreater(len(initial_requirements["requirements"]), 0)

            retry_response = self.client.post(
                "/api/nara/notices/save-and-analyze",
                json={
                    "notice": {
                        "bidNtceNo": "20260509992",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "재분석 보존 공고",
                        "prtcptPsblRgnNm": "경기도",
                        "lcnsLmtNm": "조경식재공사업",
                        "ntceSpecFileNm1": "private.pdf",
                        "ntceSpecDocUrl1": "http://127.0.0.1/private.pdf",
                    }
                },
            )
            self.assertEqual(retry_response.status_code, 202)

            deadline = time.time() + 5
            latest = {}
            while time.time() < deadline:
                detail_response = self.client.get(f"/api/nara/saved-notices/{notice_id}")
                self.assertEqual(detail_response.status_code, 200)
                latest = detail_response.get_json()
                if "기존 분석 결과를 유지" in latest.get("error_message", ""):
                    break
                time.sleep(0.05)

            self.assertIn("기존 분석 결과를 유지", latest.get("error_message", ""))
            requirements_after = self.client.get(f"/api/nara/saved-notices/{notice_id}/requirements").get_json()
            self.assertEqual(len(requirements_after["requirements"]), len(initial_requirements["requirements"]))
            self.assertEqual({item["id"] for item in latest["attachments"]}, initial_attachment_ids)
        finally:
            runtime.request_binary = original_request_binary

    def test_nara_reanalysis_without_supported_attachments_preserves_existing_results(self) -> None:
        original_request_binary = runtime.request_binary

        def fake_request_binary(url, timeout=30, max_bytes=50 * 1024 * 1024):
            if url == "https://example.com/valid-preserve.pdf":
                return (
                    200,
                    {"content-type": "application/pdf"},
                    make_pdf_bytes("제출서류: 지방세 납세증명서. 면허: 정보통신공사업."),
                )
            return original_request_binary(url, timeout=timeout, max_bytes=max_bytes)

        runtime.request_binary = fake_request_binary
        try:
            initial_response = self.client.post(
                "/api/nara/notices/save-and-analyze",
                json={
                    "notice": {
                        "bidNtceNo": "20260509993",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "첨부 누락 재분석 보존 공고",
                        "prtcptPsblRgnNm": "서울",
                        "lcnsLmtNm": "정보통신공사업",
                        "ntceSpecFileNm1": "valid-preserve.pdf",
                        "ntceSpecDocUrl1": "https://example.com/valid-preserve.pdf",
                    }
                },
            )
            self.assertEqual(initial_response.status_code, 202)
            notice_id = initial_response.get_json()["notice"]["id"]
            initial_saved = self.wait_for_saved_nara_notice(notice_id)
            self.assertEqual(initial_saved["analysis_status"], "completed")
            initial_attachment_ids = {item["id"] for item in initial_saved["attachments"]}
            initial_requirements = self.client.get(f"/api/nara/saved-notices/{notice_id}/requirements").get_json()
            self.assertGreater(len(initial_requirements["requirements"]), 0)

            retry_response = self.client.post(
                "/api/nara/notices/save-and-analyze",
                json={
                    "notice": {
                        "bidNtceNo": "20260509993",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "첨부 누락 재분석 보존 공고",
                        "prtcptPsblRgnNm": "서울",
                        "lcnsLmtNm": "정보통신공사업",
                    }
                },
            )
            self.assertEqual(retry_response.status_code, 202)

            deadline = time.time() + 5
            latest = {}
            while time.time() < deadline:
                detail_response = self.client.get(f"/api/nara/saved-notices/{notice_id}")
                self.assertEqual(detail_response.status_code, 200)
                latest = detail_response.get_json()
                if "지원 가능한 첨부" in latest.get("error_message", ""):
                    break
                time.sleep(0.05)

            self.assertEqual(latest["download_status"], "no_supported_attachments")
            self.assertEqual(latest["analysis_status"], "partial_failed")
            self.assertIn("지원 가능한 첨부", latest.get("error_message", ""))
            self.assertEqual({item["id"] for item in latest["attachments"]}, initial_attachment_ids)
            requirements_after = self.client.get(f"/api/nara/saved-notices/{notice_id}/requirements").get_json()
            self.assertEqual(len(requirements_after["requirements"]), len(initial_requirements["requirements"]))

            unsupported_response = self.client.post(
                "/api/nara/notices/save-and-analyze",
                json={
                    "notice": {
                        "bidNtceNo": "20260509993",
                        "bidNtceOrd": "000",
                        "bidNtceNm": "첨부 누락 재분석 보존 공고",
                        "prtcptPsblRgnNm": "서울",
                        "lcnsLmtNm": "정보통신공사업",
                        "ntceSpecFileNm1": "unsupported.hwp",
                        "ntceSpecDocUrl1": "https://example.com/unsupported.hwp",
                    }
                },
            )
            self.assertEqual(unsupported_response.status_code, 202)

            deadline = time.time() + 5
            latest = {}
            while time.time() < deadline:
                detail_response = self.client.get(f"/api/nara/saved-notices/{notice_id}")
                self.assertEqual(detail_response.status_code, 200)
                latest = detail_response.get_json()
                if "unsupported.hwp" in latest.get("raw_json", "") and "지원 가능한 첨부" in latest.get("error_message", ""):
                    break
                time.sleep(0.05)

            self.assertIn("지원 가능한 첨부", latest.get("error_message", ""))
            self.assertEqual({item["id"] for item in latest["attachments"]}, initial_attachment_ids)
            requirements_after_unsupported = self.client.get(f"/api/nara/saved-notices/{notice_id}/requirements").get_json()
            self.assertEqual(len(requirements_after_unsupported["requirements"]), len(initial_requirements["requirements"]))
        finally:
            runtime.request_binary = original_request_binary

    def test_nara_notice_can_be_saved_and_summarized_from_selected_raw_item(self) -> None:
        response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260500001",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "테스트 나라장터 공고",
                    "ntceInsttNm": "테스트 공고기관",
                    "dminsttNm": "테스트 수요기관",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidBeginDt": "2026-05-10 09:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "opengDt": "2026-05-21 11:00",
                    "presmptPrce": "1000000",
                    "bssamt": "1100000",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "조경식재공사업",
                    "bidNtceDtlUrl": "https://example.go.kr/notice/20260500001",
                    "ntceSpecFileNm1": "sample.hwp",
                    "ntceSpecDocUrl1": "",
                }
            },
        )
        self.assertEqual(response.status_code, 202)

        payload = response.get_json()
        self.assertEqual(payload["status"], "queued")
        self.assertEqual(payload["notice"]["bid_ntce_no"], "20260500001")
        self.assertEqual(payload["notice"]["analysis_status"], "pending")

        saved = self.wait_for_saved_nara_notice(payload["notice"]["id"])
        self.assertEqual(saved["download_status"], "no_supported_attachments")
        self.assertIn("문서 요약", saved["analysis_summary_markdown"])
        self.assertIn("공고 요구조건 구조화 후보", saved["analysis_summary_markdown"])
        summary_json = json.loads(saved["analysis_summary_json"])
        requirements = summary_json["notice_requirements"]
        self.assertIn("경기도", requirements["regions"])
        self.assertIn("조경식재공사업", requirements["licenses"])
        self.assertEqual(requirements["money"]["presmpt_prce"], "1000000")
        self.assertEqual(requirements["dates"]["bid_clse_dt"], "2026-05-20 17:00")
        self.assertFalse(
            any(attachment["source_field"] == "stdNtceDocUrl" for attachment in saved["attachments"])
        )

        saved_response = self.client.get(f"/api/nara/saved-notices/{payload['notice']['id']}")
        self.assertEqual(saved_response.status_code, 200)
        self.assertEqual(saved_response.get_json()["bid_ntce_nm"], "테스트 나라장터 공고")

    def test_saved_notice_requirements_api_returns_candidate_rows(self) -> None:
        response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260590013",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "요구조건 후보 공고",
                    "ntceInsttNm": "테스트 공고기관",
                    "dminsttNm": "테스트 수요기관",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidBeginDt": "2026-05-10 09:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "opengDt": "2026-05-21 11:00",
                    "presmptPrce": "36190000",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "조경식재공사업",
                }
            },
        )
        self.assertEqual(response.status_code, 202)
        saved = self.wait_for_saved_nara_notice(response.get_json()["notice"]["id"])

        requirements_response = self.client.get(f"/api/nara/saved-notices/{saved['id']}/requirements")
        self.assertEqual(requirements_response.status_code, 200)
        payload = requirements_response.get_json()

        self.assertEqual(payload["notice_id"], saved["id"])
        self.assertGreaterEqual(payload["summary"]["total_count"], 4)
        self.assertEqual(payload["summary"]["status"], "candidate_only")
        requirement_types = {item["requirement_type"] for item in payload["requirements"]}
        self.assertIn("region", requirement_types)
        self.assertIn("license", requirement_types)
        self.assertIn("date", requirement_types)
        money_values = [item["required_value"] for item in payload["requirements"] if item["requirement_type"] == "money"]
        self.assertIn("추정가격: 36,190,000원", money_values)
        self.assertNotIn("추정가격: 36190000", money_values)
        self.assertNotIn("eligible", json.dumps(payload).lower())

    def test_corporation_comparison_profile_normalizes_profile_and_evidence(self) -> None:
        response = self.client.post(
            "/api/corporations",
            json={
                "name": "비교 테스트 법인",
                "region": "경기도",
                "business_registration_number": "1428128387",
                "business_type": "건설업",
                "business_item": "조경식재공사업",
                "company_size_classification": "소기업",
                "certifications_json": ["여성기업"],
                "direct_production_items_json": ["조경식재"],
                "license_summary": "조경식재공사업",
            },
        )
        self.assertEqual(response.status_code, 201)
        corporation = response.get_json()

        profile_response = self.client.get(f"/api/corporations/{corporation['id']}/comparison-profile")
        self.assertEqual(profile_response.status_code, 200)
        profile = profile_response.get_json()

        self.assertIn("경기도", profile["regions"])
        self.assertIn("조경식재공사업", profile["licenses"])
        self.assertIn("소기업", profile["company_types"])
        self.assertIn("사업자등록증", profile["required_documents"])
        self.assertIn("직접생산확인증명서", profile["required_documents"])

    def test_notice_comparison_preview_persists_missing_items_without_final_verdict(self) -> None:
        corporation_response = self.client.post(
            "/api/corporations",
            json={
                "name": "부족조건 비교 법인",
                "region": "경기도",
                "business_registration_number": "1428128387",
                "business_item": "조경식재공사업",
                "company_size_classification": "중소기업",
                "license_summary": "조경식재공사업",
            },
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()

        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260500012",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "부족조건 비교 공고",
                    "ntceInsttNm": "테스트 공고기관",
                    "dminsttNm": "테스트 수요기관",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidBeginDt": "2026-05-10 09:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "opengDt": "2026-05-21 11:00",
                    "presmptPrce": "1000000",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "산림사업법인",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        comparison_response = self.client.post(
            "/api/notice-comparisons",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"]},
        )
        self.assertEqual(comparison_response.status_code, 201)
        comparison = comparison_response.get_json()

        self.assertEqual(comparison["status"], "preview")
        self.assertGreaterEqual(comparison["summary"]["prepared_count"], 1)
        self.assertGreaterEqual(comparison["summary"]["possibly_missing_count"], 1)
        self.assertGreaterEqual(comparison["summary"]["needs_review_count"], 1)
        self.assertEqual(comparison["user_summary"]["generated_by"], "fallback")
        self.assertTrue(comparison["user_summary"]["plain_summary"])
        self.assertTrue(any(link["type"] == "notice_requirement" for link in comparison["user_summary"]["evidence_links"]))
        self.assertNotIn("eligible", json.dumps(comparison).lower())

        requirement_link = next(link for link in comparison["user_summary"]["evidence_links"] if link["type"] == "notice_requirement")
        requirement_detail_response = self.client.get(f"/api/notice-requirements/{requirement_link['requirement_candidate_id']}")
        self.assertEqual(requirement_detail_response.status_code, 200)
        requirement_detail = requirement_detail_response.get_json()
        self.assertEqual(requirement_detail["detail_type"], "notice_requirement")
        self.assertEqual(requirement_detail["notice"]["id"], notice["id"])

        missing_requirement_response = self.client.get("/api/notice-requirements/999999")
        self.assertEqual(missing_requirement_response.status_code, 404)

        list_response = self.client.get("/api/notice-comparisons")
        self.assertEqual(list_response.status_code, 200)
        self.assertEqual(list_response.get_json()[0]["id"], comparison["id"])

        notice_history_response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/comparisons")
        self.assertEqual(notice_history_response.status_code, 200)
        self.assertEqual(len(notice_history_response.get_json()), 1)

        detail_response = self.client.get(f"/api/notice-comparisons/{comparison['id']}")
        self.assertEqual(detail_response.status_code, 200)
        self.assertEqual(detail_response.get_json()["summary"]["status"], "preview_only")

        internal_note_response = self.client.patch(
            f"/api/corporations/{corporation['id']}",
            json={"internal_notes": "비교와 무관한 메모 수정"},
        )
        self.assertEqual(internal_note_response.status_code, 200)
        unchanged_history_response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/comparisons")
        self.assertEqual(len(unchanged_history_response.get_json()), 1)

        profile_patch_response = self.client.patch(
            f"/api/corporations/{corporation['id']}",
            json={"license_summary": "산림사업법인"},
        )
        self.assertEqual(profile_patch_response.status_code, 200)
        invalidated_history_response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/comparisons")
        self.assertEqual(invalidated_history_response.status_code, 200)
        self.assertEqual(invalidated_history_response.get_json(), [])

        comparison_response = self.client.post(
            "/api/notice-comparisons",
            json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"]},
        )
        self.assertEqual(comparison_response.status_code, 201)

        extract_response = self.client.post(f"/api/nara/saved-notices/{notice['id']}/requirements/extract")
        self.assertEqual(extract_response.status_code, 200)

        stale_history_response = self.client.get(f"/api/nara/saved-notices/{notice['id']}/comparisons")
        self.assertEqual(stale_history_response.status_code, 200)
        self.assertEqual(stale_history_response.get_json(), [])

    def test_notice_comparison_user_summary_uses_gemini_payload_when_configured(self) -> None:
        corporation_response = self.client.post(
            "/api/corporations",
            json={
                "name": "Gemini 요약 테스트 법인",
                "region": "경기도",
                "business_registration_number": "1428128387",
                "business_item": "조경식재공사업",
                "company_size_classification": "중소기업",
                "license_summary": "조경식재공사업",
            },
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()

        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260500013",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "Gemini 요약 비교 공고",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "산림사업법인",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        previous_generator = runtime.generate_json_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY
        observed = {}

        def fake_json_generator(prompt: str, selection: dict):
            observed["prompt"] = prompt
            observed["selection"] = dict(selection)
            return (
                {
                    "headline_status": "보강 필요",
                    "plain_summary": "AI가 기존 비교 결과만 쉬운 말로 정리했습니다.",
                    "top_priority_actions": [
                        {
                            "title": "면허 증빙 확인",
                            "reason": "공고 요구조건과 법인 승인 정보가 다릅니다.",
                            "next_step": "면허증을 업로드하고 승인 후보를 검토하세요.",
                            "related_requirement_ids": [],
                            "documents": ["면허증"],
                        }
                    ],
                    "missing_groups": [{"group": "면허", "count": 1, "summary": "면허 보강이 필요합니다."}],
                    "item_explanations": {},
                    "risk_notes": ["최종 판정이 아니라 검토용 요약입니다."],
                },
                {"provider": "gemini", "model": "gemini-test-model"},
            )

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.generate_json_with_ai = fake_json_generator
            comparison_response = self.client.post(
                "/api/notice-comparisons",
                json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"]},
            )
        finally:
            runtime.generate_json_with_ai = previous_generator
            runtime.GEMINI_API_KEY = previous_gemini_key

        self.assertEqual(comparison_response.status_code, 201)
        comparison = comparison_response.get_json()
        self.assertEqual(observed["selection"]["provider"], "gemini")
        self.assertIn("새 사실을 만들지 말고", observed["prompt"])
        self.assertNotIn("보강 필요", observed["prompt"])
        self.assertEqual(comparison["user_summary"]["generated_by"], "gemini")
        self.assertEqual(comparison["user_summary"]["model"], "gemini-test-model")
        self.assertNotIn("보강 필요", json.dumps(comparison["user_summary"], ensure_ascii=False))
        self.assertEqual(comparison["user_summary"]["plain_summary"], "AI가 기존 비교 결과만 쉬운 말로 정리했습니다.")

    def test_judgment_run_user_summary_uses_gemini_payload_when_configured(self) -> None:
        self.upload_basis_document(
            "Forest business license is a citation candidate for bidder qualification review.",
            file_name="judgment-gemini-basis.pdf",
        )
        corporation_response = self.client.post(
            "/api/corporations",
            json={
                "name": "Gemini 판단 테스트 법인",
                "region": "경기도",
                "business_registration_number": "1428128387",
                "business_item": "조경식재공사업",
                "company_size_classification": "중소기업",
                "license_summary": "조경식재공사업",
            },
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()

        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260500014",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "Gemini 판단 검토 공고",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "산림사업법인",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        previous_generator = runtime.generate_json_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY
        observed = {}

        def fake_json_generator(prompt: str, selection: dict):
            observed["prompt"] = prompt
            observed["selection"] = dict(selection)
            return (
                {
                    "headline_status": "보강 필요",
                    "plain_summary": "Gemini가 판단 검토 결과를 사람이 이해하기 쉽게 정리했습니다.",
                    "top_priority_actions": [
                        {
                            "title": "면허 조건 확인",
                            "reason": "공고가 요구한 면허와 법인 보유 정보가 일치하지 않습니다.",
                            "next_step": "면허 증빙을 업로드하고 기준문서 근거를 함께 검토하세요.",
                            "related_requirement_ids": [],
                            "documents": ["면허증"],
                        }
                    ],
                    "missing_groups": [{"group": "면허", "count": 1, "summary": "면허 조건 확인이 필요합니다."}],
                    "item_explanations": {},
                    "risk_notes": ["최종 판정이 아니라 검토용 판단 정리입니다."],
                },
                {"provider": "gemini", "model": "gemini-test-model"},
            )

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.generate_json_with_ai = fake_json_generator
            judgment_response = self.client.post(
                "/api/judgment-runs",
                json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
            )
        finally:
            runtime.generate_json_with_ai = previous_generator
            runtime.GEMINI_API_KEY = previous_gemini_key

        self.assertEqual(judgment_response.status_code, 201)
        judgment = judgment_response.get_json()
        self.assertEqual(observed["selection"]["provider"], "gemini")
        self.assertIn('"mode": "judgment"', observed["prompt"])
        self.assertIn("새 사실을 만들지 말고", observed["prompt"])
        self.assertNotIn("보강 필요", observed["prompt"])
        self.assertEqual(judgment["result"]["user_summary"]["generated_by"], "gemini")
        self.assertEqual(judgment["result"]["user_summary"]["model"], "gemini-test-model")
        self.assertNotIn("보강 필요", json.dumps(judgment["result"]["user_summary"], ensure_ascii=False))
        self.assertEqual(
            judgment["result"]["user_summary"]["plain_summary"],
            "Gemini가 판단 검토 결과를 사람이 이해하기 쉽게 정리했습니다.",
        )

    def test_judgment_run_uses_gemini_weighted_merge_for_item_statuses(self) -> None:
        self.upload_basis_document(
            "Forest business license is a basis evidence candidate for bidder qualification review.",
            file_name="judgment-gemini-merge-basis.pdf",
        )
        corporation_response = self.client.post(
            "/api/corporations",
            json={
                "name": "Gemini 병합 테스트 법인",
                "region": "경기도",
                "business_registration_number": "1428128387",
                "license_summary": "",
            },
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260500015",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "Gemini 보수 병합 공고",
                    "bidNtceDt": "2026-05-05 10:00",
                    "bidClseDt": "2026-05-20 17:00",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "forest business license",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        previous_generator = runtime.generate_json_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY
        observed = {}

        def fake_json_generator(prompt: str, selection: dict):
            if "judgment_assistance" in prompt:
                observed["assistance_prompt"] = prompt
                context = json.loads(prompt.split("[context]\n", 1)[1])
                response_items = []
                for item in context["items"]:
                    if item["requirement_type"] == "region":
                        response_items.append(
                            {
                                "id": item["id"],
                                "match_status": "needs_review",
                                "reason": "지역 조건은 공고 원문 확인이 필요합니다.",
                                "recommended_action": "공고 원문과 법인 주소를 함께 확인하세요.",
                                "confidence": 0.72,
                            }
                        )
                    if item["requirement_type"] == "license":
                        response_items.append(
                            {
                                "id": item["id"],
                                "match_status": "matched",
                                "reason": "Gemini가 면허 보유처럼 보인다고 제안했습니다.",
                                "recommended_action": "면허 증빙을 검토하세요.",
                                "confidence": 0.91,
                            }
                        )
                return ({"items": response_items}, {"provider": "gemini", "model": "gemini-test-model"})
            observed["summary_prompt"] = prompt
            return (
                {
                    "headline_status": "준비 필요",
                    "plain_summary": "Gemini 요약",
                    "top_priority_actions": [],
                    "missing_groups": [],
                    "item_explanations": {},
                    "risk_notes": [],
                },
                {"provider": "gemini", "model": "gemini-test-model"},
            )

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.generate_json_with_ai = fake_json_generator
            judgment_response = self.client.post(
                "/api/judgment-runs",
                json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
            )
        finally:
            runtime.generate_json_with_ai = previous_generator
            runtime.GEMINI_API_KEY = previous_gemini_key

        self.assertEqual(judgment_response.status_code, 201)
        payload = judgment_response.get_json()
        self.assertIn("assistance_prompt", observed)
        self.assertNotIn("citation", observed["assistance_prompt"])
        ai_judgment = payload["result"]["ai_judgment"]
        self.assertEqual(ai_judgment["generated_by"], "gemini")
        self.assertEqual(ai_judgment["policy"], "gemini_weighted_70_conservative_merge")
        self.assertEqual(ai_judgment["ai_weight"], 0.7)
        self.assertEqual(ai_judgment["minimum_ai_confidence"], 0.7)
        region_item = next(item for item in payload["result"]["items"] if item["requirement_type"] == "region")
        license_item = next(item for item in payload["result"]["items"] if item["requirement_type"] == "license")
        self.assertEqual(region_item["deterministic_match_status"], "matched")
        self.assertEqual(region_item["ai_match_status"], "needs_review")
        self.assertEqual(region_item["match_status"], "needs_review")
        self.assertEqual(region_item["status_source"], "gemini_weighted")
        self.assertEqual(license_item["deterministic_match_status"], "missing")
        self.assertEqual(license_item["ai_match_status"], "matched")
        self.assertEqual(license_item["match_status"], "matched")
        self.assertEqual(license_item["status_source"], "gemini_weighted")
        self.assertNotIn("citation", json.dumps(payload["result"]["ai_judgment"], ensure_ascii=False))

    def test_judgment_run_invalid_gemini_assistance_falls_back_safely(self) -> None:
        self.upload_basis_document(
            "Forest business license is a basis evidence candidate for bidder qualification review.",
            file_name="judgment-gemini-invalid-basis.pdf",
        )
        corporation_response = self.client.post(
            "/api/corporations",
            json={"name": "Gemini invalid 테스트 법인", "region": "경기도", "license_summary": ""},
        )
        self.assertEqual(corporation_response.status_code, 201)
        corporation = corporation_response.get_json()
        notice_response = self.client.post(
            "/api/nara/notices/save-and-analyze",
            json={
                "notice": {
                    "bidNtceNo": "20260500016",
                    "bidNtceOrd": "000",
                    "bidNtceNm": "Gemini invalid 판단 공고",
                    "bidClseDt": "2026-05-20 17:00",
                    "prtcptPsblRgnNm": "경기도",
                    "lcnsLmtNm": "forest business license",
                }
            },
        )
        self.assertEqual(notice_response.status_code, 202)
        notice = self.wait_for_saved_nara_notice(notice_response.get_json()["notice"]["id"])

        previous_generator = runtime.generate_json_with_ai
        previous_gemini_key = runtime.GEMINI_API_KEY

        def fake_json_generator(prompt: str, selection: dict):
            if "judgment_assistance" in prompt:
                return ({"items": [None, {"id": 123, "match_status": "eligible", "reason": 456}]}, {"provider": "gemini", "model": "gemini-test-model"})
            return (
                {
                    "headline_status": "준비 필요",
                    "plain_summary": "Gemini 요약",
                    "top_priority_actions": [None, {"title": 123}],
                    "missing_groups": [],
                    "item_explanations": {},
                    "risk_notes": [],
                },
                {"provider": "gemini", "model": "gemini-test-model"},
            )

        try:
            runtime.GEMINI_API_KEY = "gemini-test-key"
            runtime.generate_json_with_ai = fake_json_generator
            judgment_response = self.client.post(
                "/api/judgment-runs",
                json={"nara_notice_id": notice["id"], "corporation_id": corporation["id"], "top_k": 3},
            )
        finally:
            runtime.generate_json_with_ai = previous_generator
            runtime.GEMINI_API_KEY = previous_gemini_key

        self.assertEqual(judgment_response.status_code, 201)
        payload = judgment_response.get_json()
        self.assertEqual(payload["result"]["ai_judgment"]["generated_by"], "fallback")
        self.assertEqual(payload["result"]["ai_judgment"]["fallback_reason"], "invalid_ai_items")
        self.assertTrue(all(item["ai_match_status"] == "" for item in payload["result"]["items"]))
        self.assertTrue(all(item["final_match_status"] == item["deterministic_match_status"] for item in payload["result"]["items"]))
        self.assertNotIn("보강", json.dumps(payload["result"], ensure_ascii=False))

    def test_notice_requirement_extraction_reads_text_candidates_without_verdict(self) -> None:
        requirements = runtime.extract_notice_requirements(
            {
                "region_text": "전라남도 해남군",
                "license_text": "산림사업법인",
                "presmpt_prce": "103980909",
                "bid_clse_dt": "2026-05-08 10:00",
            },
            """
            입찰참가자격: 중소기업확인서를 보유한 업체
            면허: 산림사업법인 또는 조경식재공사업 등록 업체
            제출서류: 사업자등록증, 국세납세증명서, 직접생산확인증명서
            우대조건: 여성기업
            """,
        )

        self.assertIn("전라남도 해남군", requirements["regions"])
        self.assertIn("산림사업법인", requirements["licenses"])
        self.assertIn("조경식재", requirements["licenses"])
        self.assertIn("중소기업", requirements["company_types"])
        self.assertIn("여성기업", requirements["company_types"])
        self.assertIn("사업자등록증", requirements["required_documents"])
        self.assertIn("국세 납세증명서", requirements["required_documents"])
        self.assertTrue(requirements["uncertainty_notes"])
        self.assertNotIn("eligible", json.dumps(requirements).lower())
        self.assertNotIn("지원 가능", json.dumps(requirements, ensure_ascii=False))

    def test_notice_requirement_extraction_handles_aliases_and_no_space_documents(self) -> None:
        requirements = runtime.extract_notice_requirements(
            {
                "region_text": "전라남도 해남군",
                "license_text": "",
                "presmpt_prce": "103980909",
                "bid_clse_dt": "2026-05-08 10:00",
            },
            """
            입찰참가자격: 주된 영업소가 전남 해남군에 소재한 업체
            등록요건: 조경식재·시설물공사업 또는 전기공사업 등록 업체
            제출서류: 국세납세증명서, 지방세납세증명서, 경쟁입찰참가자격 등록증
            자격요건: 소기업 또는 소상공인 확인서를 제출할 것
            """,
        )

        self.assertIn("전라남도 해남군", requirements["regions"])
        self.assertIn("전남", requirements["regions"])
        self.assertIn("해남군", requirements["regions"])
        self.assertIn("조경식재", requirements["licenses"])
        self.assertIn("조경시설물", requirements["licenses"])
        self.assertIn("전기공사업", requirements["licenses"])
        self.assertIn("소기업", requirements["company_types"])
        self.assertIn("소상공인", requirements["company_types"])
        self.assertIn("국세 납세증명서", requirements["required_documents"])
        self.assertIn("지방세 납세증명서", requirements["required_documents"])
        self.assertIn("나라장터 경쟁입찰참가자격 등록증", requirements["required_documents"])

        broad_company_requirements = runtime.extract_notice_requirements(
            {"region_text": "", "license_text": ""},
            "입찰참가자격: 중소기업확인서를 제출한 중소기업자",
        )
        self.assertIn("중소기업", broad_company_requirements["company_types"])
        self.assertNotIn("소기업", broad_company_requirements["company_types"])

    def test_comparison_matching_is_conservative_for_license_like_text(self) -> None:
        prepared = runtime.compare_requirement_candidate(
            {
                "id": 1,
                "requirement_type": "license",
                "label": "면허/업종",
                "required_value": "조경식재",
                "normalized_value": "조경식재",
                "source_text": "면허",
                "confidence": 0.8,
            },
            {"licenses": ["조경식재공사업"], "business_types": []},
        )
        self.assertEqual(prepared["status"], "prepared")

        missing = runtime.compare_requirement_candidate(
            {
                "id": 2,
                "requirement_type": "license",
                "label": "면허/업종",
                "required_value": "전기공사업",
                "normalized_value": "전기공사업",
                "source_text": "면허",
                "confidence": 0.8,
            },
            {"licenses": ["전기, 소방, 통신자재"], "business_types": []},
        )
        self.assertEqual(missing["status"], "possibly_missing")

        document_missing = runtime.compare_requirement_candidate(
            {
                "id": 3,
                "requirement_type": "required_document",
                "label": "필요서류",
                "required_value": "지방세 납세증명서",
                "normalized_value": "지방세 납세증명서",
                "source_text": "제출서류",
                "confidence": 0.8,
            },
            {"required_documents": ["국세 납세증명서"], "business_types": []},
        )
        self.assertEqual(document_missing["status"], "possibly_missing")

    def test_standard_notice_url_is_named_in_korean_when_present(self) -> None:
        attachments = runtime.collect_nara_attachments(
            [
                {
                    "stdNtceDocUrl": "https://example.go.kr/notice/standard",
                }
            ]
        )

        self.assertEqual(len(attachments), 1)
        self.assertEqual(attachments[0]["file_name"], "표준공고문.pdf")
        self.assertEqual(attachments[0]["support_status"], "supported")


if __name__ == "__main__":
    unittest.main()
