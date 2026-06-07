from __future__ import annotations

import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[2]
SCRIPT_PATH = ROOT_DIR / "scripts" / "compare-real-basis-document-txt.py"


def load_compare_module():
    spec = importlib.util.spec_from_file_location("real_basis_reference_compare", SCRIPT_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load compare script: {SCRIPT_PATH}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class RealBasisReferenceCompareTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.compare = load_compare_module()

    def test_docx_reference_extracts_header_and_table_cells_from_xml(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "reference.docx"
            document = Document()
            document.sections[0].header.paragraphs[0].text = "머리글 직접생산"
            document.add_paragraph("본문 확인기준")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "세부품명"
            table.cell(0, 1).text = "생산시설"
            table.cell(1, 0).text = "검사설비"
            table.cell(1, 1).text = "경쟁제품"
            document.save(docx_path)

            text, metadata = self.compare.read_reference_file(docx_path)

        self.assertEqual(metadata["type"], "docx")
        self.assertEqual(metadata["engine"], "docx-package-xml-wt")
        self.assertEqual(metadata["table_count"], 1)
        self.assertEqual(metadata["table_cell_count"], 4)
        self.assertGreater(metadata["xml_text_part_count"], 0)
        self.assertGreater(metadata["xml_char_count"], 0)
        self.assertGreater(metadata["python_docx_char_count"], 0)
        self.assertIn("머리글 직접생산", text)
        self.assertIn("본문 확인기준", text)
        self.assertIn("세부품명", text)
        self.assertIn("생산시설", text)
        self.assertIn("검사설비", text)
        self.assertIn("경쟁제품", text)

    def test_txt_reference_detects_encoding_and_compares_korean_tokens(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            txt_path = Path(tmp) / "reference.txt"
            txt_path.write_text("직접생산 확인기준\n세부품명 생산시설 검사설비\n", encoding="utf-8-sig")

            text, metadata = self.compare.read_reference_file(txt_path)
            metrics = self.compare.compare_texts(
                "직접생산 확인기준 세부품명 생산시설 검사설비",
                text,
                ngram_size=3,
            )

        self.assertEqual(metadata["type"], "txt")
        self.assertEqual(metadata["encoding"], "utf-8-sig")
        self.assertEqual(metrics["service_token_multiset_recall_in_reference"], 1.0)
        self.assertEqual(metrics["reference_token_multiset_recall_in_service"], 1.0)

    def test_md_reference_is_read_as_markdown_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            md_path = Path(tmp) / "reference.md"
            md_path.write_text(
                "# 직접생산 확인기준\n\n| 세부품명 | 생산시설 |\n| --- | --- |\n| 검사설비 | 경쟁제품 |\n",
                encoding="utf-8",
            )

            text, metadata = self.compare.read_reference_file(md_path)
            metrics = self.compare.compare_texts(
                "직접생산 확인기준 세부품명 생산시설 검사설비 경쟁제품",
                text,
                ngram_size=3,
            )

        self.assertEqual(metadata["type"], "md")
        self.assertEqual(metadata["encoding"], "utf-8-sig")
        self.assertIn("| 세부품명 | 생산시설 |", text)
        self.assertEqual(metrics["service_token_multiset_recall_in_reference"], 1.0)

    def test_unsupported_reference_extension_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            reference_path = Path(tmp) / "reference.hwp"
            reference_path.write_text("unsupported", encoding="utf-8")

            with self.assertRaises(ValueError):
                self.compare.read_reference_file(reference_path)


if __name__ == "__main__":
    unittest.main()
